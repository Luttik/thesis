"""Count how often each interview participant is quoted in the thesis."""

from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

import typer
from docx import Document
from docx.text.paragraph import Paragraph

app = typer.Typer(add_completion=False, no_args_is_help=True)

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOCX = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
INTERVIEWEES_XLSX = ROOT / "interviewees.xlsx"

OPEN_DOUBLE = {'"', "\u201c"}
CLOSE_DOUBLE = {'"', "\u201d"}
QUOTE_STYLES = {"Quote", "Block Text"}

INTERVIEWEE_RE = re.compile(r"\bInterviewee\s+(\d+)\b", re.I)
ATTR_RE = re.compile(r"^[\u2014\u2013-]\s*Interviewee\s+(\d+)\s*$", re.I)
CH4_START = re.compile(r"^4\.\s")
CH5_START = re.compile(r"^5\.\s")
ANALYSIS_START = re.compile(
    r"^(The|This|These|However|Managers|Taken together|Across|Where|While|Alongside|"
    r"Four |Generating|Scale |Efficiency|AI |Brand |Hallucination|Agentic|Running|"
    r"Punishing|Looking |Content |Customer|Internal |Personalizing|On the |He notes|"
    r"Interviewee \d+,?\s+(an|who)\b)",
    re.I,
)


def heading_level(paragraph: Paragraph) -> int | None:
    style = paragraph.style.name if paragraph.style else ""
    match = re.match(r"Heading\s+(\d+)", style)
    return int(match.group(1)) if match else None


def chapter_four_bounds(paragraphs: list[Paragraph]) -> tuple[int, int] | None:
    start = end = None
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        level = heading_level(paragraph)
        if level == 1 and CH4_START.search(text) and "Findings" in text:
            start = index
        if start is not None and level == 1 and CH5_START.search(text):
            end = index
            break
    if start is None:
        return None
    return start, end if end is not None else len(paragraphs)


def load_participant_names() -> dict[int, str]:
    if not INTERVIEWEES_XLSX.exists():
        return {}

    try:
        from openpyxl import load_workbook
    except ImportError:
        return {}

    workbook = load_workbook(INTERVIEWEES_XLSX, read_only=True, data_only=True)
    sheet = workbook["Sheet1"]
    names: dict[int, str] = {}
    for raw_number, raw_name in sheet.iter_rows(min_row=2, max_col=2, values_only=True):
        if raw_number is None:
            continue
        try:
            number = int(raw_number)
        except (TypeError, ValueError):
            continue
        name = (raw_name or "").strip()
        if "anonym" in name.lower() or "anonim" in name.lower() or not name:
            name = "Anonymous"
        names[number] = name
    return names


def display_name(participant_id: int, names: dict[int, str]) -> str:
    return names.get(participant_id, "")


def load_participant_ids(document: Document) -> list[int]:
    for table in document.tables:
        if not table.rows:
            continue
        header = table.rows[0].cells[0].text.strip().lower()
        if "interviewee" not in header:
            continue
        ids: list[int] = []
        for row in table.rows[2:]:
            value = row.cells[0].text.strip()
            if value.isdigit():
                ids.append(int(value))
        if ids:
            return sorted(ids)
    return list(range(1, 18))


def find_inline_quote_ranges(text: str) -> list[tuple[int, int]]:
    ranges: list[tuple[int, int]] = []
    inside = False
    start = 0
    for index, char in enumerate(text):
        if char in OPEN_DOUBLE and not inside:
            inside = True
            start = index + 1
        elif char in CLOSE_DOUBLE and inside:
            if index > start:
                ranges.append((start, index))
            inside = False
    return ranges


def interviewee_before(text: str, position: int | None = None) -> int | None:
    prefix = text if position is None else text[:position]
    matches = list(INTERVIEWEE_RE.finditer(prefix))
    if not matches:
        return None
    return int(matches[-1].group(1))


def resolve_interviewee(paragraphs: list[Paragraph], index: int, position: int | None = None) -> int | None:
    number = interviewee_before(paragraphs[index].text, position)
    if number is not None:
        return number
    for back in range(1, 4):
        lookback = index - back
        if lookback < 0:
            break
        number = interviewee_before(paragraphs[lookback].text)
        if number is not None:
            return number
    return None


def attribution_interviewee(text: str) -> int | None:
    match = ATTR_RE.match(text.strip())
    return int(match.group(1)) if match else None


def continues_quote_intro(text: str) -> int | None:
    stripped = text.strip()
    if not stripped.endswith(":"):
        return None
    return interviewee_before(stripped)


def looks_like_analysis(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return True
    if heading_level(type("P", (), {"style": None, "text": stripped})()) is not None:
        return True
    return ANALYSIS_START.match(stripped) is not None


def count_participant_quotes(
    paragraphs: list[Paragraph],
    *,
    start: int = 0,
    end: int | None = None,
) -> tuple[Counter[int], list[str]]:
    counts: Counter[int] = Counter()
    notes: list[str] = []
    end = len(paragraphs) if end is None else end

    index = start
    while index < end:
        paragraph = paragraphs[index]
        text = paragraph.text.strip()
        style = paragraph.style.name if paragraph.style else ""

        if attribution_interviewee(text) is not None:
            index += 1
            continue

        if style in QUOTE_STYLES and text:
            number = None
            if index + 1 < end:
                number = attribution_interviewee(paragraphs[index + 1].text)
            if number is None:
                number = resolve_interviewee(paragraphs, index)
            if number is None:
                notes.append(f"Unassigned block quote at paragraph {index + 1}")
            else:
                counts[number] += 1
            index += 1
            continue

        if text and index + 1 < end:
            next_number = attribution_interviewee(paragraphs[index + 1].text)
            if next_number is not None and style not in QUOTE_STYLES:
                counts[next_number] += 1
                index += 2
                continue

        if text and index > start:
            intro_number = continues_quote_intro(paragraphs[index - 1].text)
            if intro_number is not None and not looks_like_analysis(text):
                counts[intro_number] += 1
                index += 1
                continue

        inline_ranges = find_inline_quote_ranges(text)
        if inline_ranges:
            for range_start, _ in inline_ranges:
                number = resolve_interviewee(paragraphs, index, range_start)
                if number is None:
                    notes.append(f"Unassigned inline quote at paragraph {index + 1}")
                else:
                    counts[number] += 1
            index += 1
            continue

        index += 1

    return counts, notes


def format_report(
    participant_ids: list[int],
    counts: Counter[int],
    notes: list[str],
    *,
    scope_label: str,
    names: dict[int, str],
) -> str:
    lines = [
        f"Participant quote counts ({scope_label})",
        "",
        f"{'#':<4} {'Name':<24} {'Quotes':>8}",
        "-" * 40,
    ]

    for participant_id in participant_ids:
        name = display_name(participant_id, names) or "—"
        lines.append(
            f"{participant_id:<4} {name:<24} {counts.get(participant_id, 0):>8}"
        )

    lines.extend(
        [
            "-" * 40,
            f"{'Total':<4} {'':<24} {sum(counts.values()):>8}",
        ]
    )

    if notes:
        lines.extend(["", f"Unassigned quotes: {len(notes)}"])
        for note in notes[:10]:
            lines.append(f"  - {note}")
        if len(notes) > 10:
            lines.append(f"  - ... and {len(notes) - 10} more")

    quoted_participants = sum(1 for participant_id in participant_ids if counts.get(participant_id, 0) > 0)
    lines.extend(
        [
            "",
            f"Participants quoted: {quoted_participants}/{len(participant_ids)}",
        ]
    )
    return "\n".join(lines)


@app.command()
def list_counts(
    input_path: Path = typer.Argument(
        None,
        help="Thesis docx to analyse. Defaults to the main thesis draft.",
    ),
    chapter: str = typer.Option(
        "4",
        "--chapter",
        help="Chapter scope: '4' for Findings only, or 'all' for the full document.",
    ),
    json_output: bool = typer.Option(False, "--json", help="Print JSON instead of a table."),
    output: Path | None = typer.Option(None, "--output", "-o", help="Optional output file."),
) -> None:
    """List how often each interview participant is quoted."""
    path = input_path or DEFAULT_DOCX
    if not path.exists():
        typer.echo(f"File not found: {path}", err=True)
        raise typer.Exit(code=1)

    document = Document(str(path))
    paragraphs = document.paragraphs
    participant_ids = load_participant_ids(document)
    names = load_participant_names()

    if chapter.lower() == "all":
        start, end = 0, len(paragraphs)
        scope_label = "full document"
    else:
        bounds = chapter_four_bounds(paragraphs)
        if bounds is None:
            typer.echo("Chapter 4 heading not found.", err=True)
            raise typer.Exit(code=1)
        start, end = bounds
        scope_label = "Chapter 4 Findings"

    counts, notes = count_participant_quotes(paragraphs, start=start, end=end)

    if json_output:
        payload = {
            "scope": scope_label,
            "participants": [
                {
                    "participant_id": participant_id,
                    "name": display_name(participant_id, names),
                    "quote_count": counts.get(participant_id, 0),
                }
                for participant_id in participant_ids
            ],
            "total_quotes": sum(counts.values()),
            "unassigned": notes,
        }
        rendered = json.dumps(payload, indent=2)
    else:
        rendered = format_report(
            participant_ids,
            counts,
            notes,
            scope_label=scope_label,
            names=names,
        )

    typer.echo(rendered)

    if output:
        output.write_text(rendered + "\n", encoding="utf-8")
        typer.echo(f"Wrote {output}")


def main() -> None:
    app()


if __name__ == "__main__":
    main()
