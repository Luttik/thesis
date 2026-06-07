# -*- coding: utf-8 -*-
"""Add the foundational contribution to §5.1 — that the study answers explicit calls
for empirical research on agentic AI in marketing — plus a matching roadmap clause in
the §5 intro. Tracked, author 'Claude'. All-or-nothing save."""
from __future__ import annotations
import copy, shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
BACKUP = ROOT / "Thesis Draft - Daan Luttik - MBA.ch5calls-backup.docx"
AUTHOR = "Claude"; DATE = "2026-06-07T15:00:00Z"
XS = "{http://www.w3.org/XML/1998/namespace}space"
APO = "’"

NEW_PARA = (
    "The study" + APO + "s most fundamental contribution precedes the specific theoretical claims "
    "developed below. Agentic AI has moved rapidly from a technical novelty to a consequential "
    "development in marketing practice, yet empirical research into its practical implications within "
    "marketing organizations remains scarce; the existing literature concentrates on generative AI and, "
    "more narrowly, on content generation. Scholars have responded by issuing explicit calls for research "
    "into agentic AI in the marketing field (Kim, 2025; Mogaji & Jain, 2024; Jain & Eastman, 2024). This "
    "study answers those calls directly: it provides one of the first grounded, empirical accounts of how "
    "marketing managers create value with agentic AI in practice, developing the empirical understanding "
    "the literature has so far lacked and moving the field beyond its prevailing focus on content "
    "generation toward the broader question of organizational value creation. The theoretical "
    "contributions set out below build on this foundation."
)

ROADMAP_OLD = "Section 5.1 develops the theoretical contributions, situating the empirical model within"
ROADMAP_NEW = ("Section 5.1 develops the theoretical contributions — beginning with the study" + APO +
               "s response to recent calls for empirical research on agentic AI in marketing — and "
               "situates the empirical model within")


def used_ids(doc):
    return {int(el.get(qn("w:id"), 0)) for el in doc.element.body.iter() if el.get(qn("w:id")) is not None}
def nid(u): n = max(u, default=0) + 1; u.add(n); return n
def acc(p): return "".join(t.text or "" for t in p.iter(qn("w:t")))

def ins_para(text, used):
    p = etree.Element(qn("w:p")); etree.SubElement(p, qn("w:pPr"))
    ins = etree.SubElement(p, qn("w:ins"))
    ins.set(qn("w:id"), str(nid(used))); ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), DATE)
    r = etree.SubElement(ins, qn("w:r")); t = etree.SubElement(r, qn("w:t"))
    t.set(XS, "preserve"); t.text = text
    return p

def ins_run(text, used):
    ins = etree.Element(qn("w:ins")); ins.set(qn("w:id"), str(nid(used)))
    ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), DATE)
    r = etree.SubElement(ins, qn("w:r")); t = etree.SubElement(r, qn("w:t"))
    t.set(XS, "preserve"); t.text = text
    return ins

def replace_fragment_tracked(p_elem, old, new, used):
    runs = p_elem.findall(qn("w:r")); segs, total = [], ""
    for r in runs:
        t = r.find(qn("w:t")); txt = t.text if (t is not None and t.text) else ""
        segs.append((r, txt)); total += txt
    start = total.find(old)
    if start < 0: return False
    end = start + len(old)
    def mk(rPr, s):
        nr = etree.Element(qn("w:r"))
        if rPr is not None: nr.append(copy.deepcopy(rPr))
        tt = etree.SubElement(nr, qn("w:t")); tt.set(XS, "preserve"); tt.text = s; return nr
    def mkdel(rPr, s):
        d = etree.Element(qn("w:del")); d.set(qn("w:id"), str(nid(used)))
        d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), DATE)
        dr = etree.SubElement(d, qn("w:r"))
        if rPr is not None: dr.insert(0, copy.deepcopy(rPr))
        dt = etree.SubElement(dr, qn("w:delText")); dt.set(XS, "preserve"); dt.text = s; return d
    pos, inserted = 0, False
    for r, txt in segs:
        rs, re_ = pos, pos + len(txt); pos = re_
        if not txt or re_ <= start or rs >= end: continue
        rPr = r.find(qn("w:rPr")); ls = max(start, rs) - rs; le = min(end, re_) - rs
        before, deleted, after = txt[:ls], txt[ls:le], txt[le:]
        parent = r.getparent(); idx = list(parent).index(r); parts = []
        if before: parts.append(mk(rPr, before))
        if not inserted: parts.append(ins_run(new, used)); inserted = True
        if deleted: parts.append(mkdel(rPr, deleted))
        if after: parts.append(mk(rPr, after))
        for j, prt in enumerate(parts): parent.insert(idx + j, prt)
        parent.remove(r)
    return inserted


def main():
    doc = Document(str(DOCX)); used = used_ids(doc)

    # 1. insert new opening paragraph right after the §5.1 heading
    heading = None
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and "Contributions to Theory" in acc(p._element):
            heading = p._element; break
    assert heading is not None, "§5.1 heading not found"
    heading.addnext(ins_para(NEW_PARA, used))
    print("1. §5.1 opening paragraph inserted")

    # 2. roadmap clause in §5 intro
    road = None
    for p in doc.paragraphs:
        if ROADMAP_OLD in acc(p._element):
            road = p._element; break
    assert road is not None, "roadmap paragraph not found"
    ok = replace_fragment_tracked(road, ROADMAP_OLD, ROADMAP_NEW, used)
    assert ok, "roadmap fragment not matched"
    print("2. §5 roadmap clause added")

    shutil.copy(DOCX, BACKUP); doc.save(str(DOCX))
    print(f"\nSaved: {DOCX}\nBackup: {BACKUP}")


if __name__ == "__main__":
    main()
