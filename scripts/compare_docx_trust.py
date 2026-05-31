"""Compare trustworthiness sections across thesis docx variants."""

from __future__ import annotations

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
FILES = [
    ROOT / "Thesis Draft - Daan Luttik - MBA.docx",
    ROOT / "Thesis Draft - Daan Luttik - MBA - trustworthiness-shenton.docx",
    ROOT / "Thesis Draft - Daan Luttik - MBA.docx.bak",
]
OUT = ROOT / ".cache" / "compare-trust.txt"


def snap(path: Path) -> str:
    if not path.exists():
        return f"{path.name}: MISSING"
    try:
        d = Document(str(path))
    except Exception as exc:
        return f"{path.name}: CANNOT OPEN — {exc}"
    trust = next(
        (
            t
            for t in d.tables
            if t.rows and "trustworthiness" in t.rows[0].cells[0].text.lower()
        ),
        None,
    )
    if trust is None:
        has_34 = any(
            p.text.strip().startswith("3.4") and "rustworthiness" in p.text.lower()
            for p in d.paragraphs
        )
        return f"{path.name}: NO trust table; §3.4 heading={has_34}; tables={len(d.tables)}"
    intro = next(
        (
            p.text[:100]
            for p in d.paragraphs
            if "Shenton" in p.text and ("Table 2" in p.text or "trustworthiness" in p.text.lower())
        ),
        "no shenton intro",
    )
    row1 = trust.rows[1].cells[1].text[:60] if len(trust.rows) > 1 else ""
    return (
        f"{path.name}: OK rows={len(trust.rows)} size={path.stat().st_size}\n"
        f"  intro: {intro}\n"
        f"  cred-how: {row1}..."
    )


def main() -> None:
    lines = [snap(p) for p in FILES]
    text = "\n\n".join(lines)
    OUT.write_text(text, encoding="utf-8")
    print(text)


if __name__ == "__main__":
    main()
