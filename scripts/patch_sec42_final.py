"""
Section 4.2 final pass:
1. Fix para 203: "richly coded themes" → plain language
2. Fix para 209: dangling "which address" grammar
3. Merge paras 205-206 (4.2.4 overlap)
4. Insert culture-as-mediator paragraph after navigate (4.2.1)
5. Insert 4.2.3 conditions paragraph before "A third cluster..."

All direct edits — no tracked changes.
"""
from __future__ import annotations
import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT      = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"


def ft(p) -> str:
    elem = p._p if hasattr(p, "_p") else p
    return "".join(t.text or "" for t in elem.iter(qn("w:t")))


def find_para(doc: Document, marker: str) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if marker in ft(p):
            return i
    return None


def set_text(p_elem, new_text: str) -> None:
    """Replace all text in a paragraph's first w:t; clear the rest."""
    t_elems = list(p_elem.iter(qn("w:t")))
    if not t_elems:
        return
    t_elems[0].text = new_text
    for t in t_elems[1:]:
        t.text = ""


def replace_in_para(p_elem, old: str, new: str) -> bool:
    for t in p_elem.iter(qn("w:t")):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)
            return True
    return False


def new_normal_para(text: str) -> etree._Element:
    p = etree.Element(qn("w:p"))
    r = etree.SubElement(p, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return p


def insert_before(anchor: etree._Element, paras: list) -> None:
    parent = anchor.getparent()
    idx = list(parent).index(anchor)
    for i, p in enumerate(paras):
        parent.insert(idx + i, p)


def insert_after(anchor: etree._Element, paras: list) -> None:
    parent = anchor.getparent()
    idx = list(parent).index(anchor) + 1
    for i, p in enumerate(paras):
        parent.insert(idx + i, p)


# ---------------------------------------------------------------------------
# New paragraph content
# ---------------------------------------------------------------------------

CULTURE_PARA = (
    "The effectiveness of both education and experimentation is mediated by a condition "
    "that is harder to reshape: organizational culture. Where a culture of innovation "
    "exists, tolerating failure and treating experimentation as intrinsically valuable, "
    "both behaviors compound rapidly. Where it is absent, even well-designed programs "
    "stall. Interviewee 7 captured the ideal condition: “A culture of innovation: "
    "forgiving. When you pioneer, you break things fast.” The absence of this "
    "permissiveness amplifies the barriers described above, making the default response "
    "navigation rather than reshaping."
)

CONDITIONS_423 = (
    "The conditions that make strategic direction necessary are broadly consistent across "
    "the data. Organizations frequently approach agentic AI without a clear rationale: "
    "they want AI because others are using it, not because they have identified a problem "
    "it can solve. Interviewee 1 described this pattern: “There will be a lot of "
    "clients and companies that want AI, but they won’t know why, outside of the "
    "fact that someone tells them they should.” Interviewee 17 noted a related "
    "pattern: organizations think in tools rather than business objectives, seeking "
    "applications before establishing what they need to achieve. The consequence is that "
    "even when AI is deployed, it generates activity without direction. Interviewee 10 "
    "observed that when organizations give employees time for AI experimentation, they "
    "attach no goals, so potential and initiative are both lost."
)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    shutil.copy(DOCX_PATH, DOCX_PATH.with_suffix(".bak.docx"))
    doc = Document(str(DOCX_PATH))

    # ------------------------------------------------------------------
    # 1. Fix "richly coded themes" in para 203
    # ------------------------------------------------------------------
    idx = find_para(doc, "most richly coded themes")
    assert idx is not None, "richly coded para not found"
    ok = replace_in_para(doc.paragraphs[idx]._p,
                         "most richly coded themes in this block",
                         "most consistently documented behaviors in this block")
    print(f"  Para 203 language fix: {'done' if ok else 'FAILED'}")

    # ------------------------------------------------------------------
    # 2. Fix dangling "which address" in para 209
    # ------------------------------------------------------------------
    idx = find_para(doc, "navigating friction, which address")
    assert idx is not None, "closing para not found"
    ok = replace_in_para(
        doc.paragraphs[idx]._p,
        "navigating friction, which address multiple conditions simultaneously "
        "and in rough sequence.",
        "navigating friction. These behaviors address multiple conditions "
        "simultaneously and in rough sequence."
    )
    print(f"  Para 209 grammar fix: {'done' if ok else 'FAILED'}")

    # ------------------------------------------------------------------
    # 3. Merge paras 205-206 (4.2.4 intro + bureaucracy overlap)
    #    Keep 205's framing + fold 206's empirical detail into it, delete 206
    # ------------------------------------------------------------------
    idx_205 = find_para(doc,
        "A fourth cluster of conditions concerns governance and organizational friction")
    idx_206 = find_para(doc,
        "Alongside these enabling behaviors, bureaucratic governance emerged")
    assert idx_205 is not None and idx_206 is not None, "4.2.4 paras not found"

    # Build merged text: 205 framing + 206 empirical detail
    text_205 = ft(doc.paragraphs[idx_205])
    text_206 = ft(doc.paragraphs[idx_206])

    # Strip the awkward opening of 206 ("Alongside these enabling behaviors,")
    # and replace with a natural join
    text_206_clean = text_206.replace(
        "Alongside these enabling behaviors, bureaucratic governance emerged as a "
        "significant opposing force. Governance friction appeared in several forms: ",
        "Governance friction appears in several forms: "
    )

    merged = text_205.rstrip(".") + ". " + text_206_clean
    set_text(doc.paragraphs[idx_205]._p, merged)

    # Delete para 206
    p206_elem = doc.paragraphs[idx_206]._p
    p206_elem.getparent().remove(p206_elem)
    print("  Paras 205-206 merged.")

    # ------------------------------------------------------------------
    # 4. Insert culture-as-mediator paragraph after navigate (para 184)
    #    After merging above, indices may have shifted — re-find
    # ------------------------------------------------------------------
    idx = find_para(doc,
        "some managers navigate around organizational capacity constraints")
    assert idx is not None, "navigate para not found"
    insert_after(doc.paragraphs[idx]._p, [new_normal_para(CULTURE_PARA)])
    print("  Culture mediator paragraph inserted in 4.2.1.")

    # ------------------------------------------------------------------
    # 5. Insert 4.2.3 conditions paragraph before "A third cluster..."
    # ------------------------------------------------------------------
    idx = find_para(doc,
        "A third cluster of conditions concerns strategic direction and leadership")
    assert idx is not None, "4.2.3 intro para not found"
    insert_before(doc.paragraphs[idx]._p, [new_normal_para(CONDITIONS_423)])
    print("  4.2.3 conditions paragraph inserted.")

    doc.save(str(DOCX_PATH))
    print(f"\nSaved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
