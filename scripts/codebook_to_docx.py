"""Render `code-book.xlsx` as a formatted DOCX table.

Reads the `Codes` sheet from ``code-book.xlsx`` at the repo root and writes a
single-page DOCX containing the descriptive / interpretive / aggregated code
columns.

Run with:
    poetry run python scripts/codebook_to_docx.py
"""

from __future__ import annotations

from pathlib import Path

import typer
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_XLSX = ROOT / "code-book.xlsx"
DEFAULT_DOCX = ROOT / "output" / "Code Book - Daan Luttik - MBA.docx"

FONT_NAME = "Georgia"
FONT_SIZE_PT = 10
# 1.5em line height -> 1.5 line spacing (multiple).
LINE_SPACING = 1.5

ARROW = "\u2192  "  # → followed by two spaces

COLUMNS = ("Descriptive coding", "Interpretive coding", "Aggregated codes")


def _set_cell_border(cell, *, top: bool = False, bottom: bool = False) -> None:
    """Apply explicit top/bottom borders to a single cell; clear the rest."""
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)

    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        if (edge == "top" and top) or (edge == "bottom" and bottom):
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "8")  # 1pt
            element.set(qn("w:color"), "000000")
        else:
            element.set(qn("w:val"), "nil")
        tc_borders.append(element)
    tc_pr.append(tc_borders)


def _disable_table_borders(table) -> None:
    """Set all table-level borders to ``nil`` so only per-cell borders show."""
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        tbl_borders.append(element)
    tbl_pr.append(tbl_borders)


def _style_paragraph(paragraph, *, bold: bool = False) -> None:
    paragraph.paragraph_format.line_spacing = LINE_SPACING
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_PT)
        run.bold = bold
        # Ensure East Asian font fallback also picks up Georgia.
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), FONT_NAME)


def _write_cell(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if text:
        paragraph.add_run(text)
    _style_paragraph(paragraph, bold=bold)


def _read_rows(xlsx_path: Path) -> list[tuple[str, str, str]]:
    wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    ws = wb["Codes"]
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise ValueError(f"No rows found in {xlsx_path}")

    header = [str(c).strip() if c is not None else "" for c in rows[0][:3]]
    if tuple(header) != COLUMNS:
        raise ValueError(
            f"Unexpected header in {xlsx_path}: {header!r}; expected {list(COLUMNS)!r}"
        )

    body: list[tuple[str, str, str]] = []
    for row in rows[1:]:
        if row is None:
            continue
        descriptive = (row[0] or "").strip() if isinstance(row[0], str) else ""
        interpretive = (row[1] or "").strip() if isinstance(row[1], str) else ""
        aggregated = (row[2] or "").strip() if isinstance(row[2], str) else ""
        if not any((descriptive, interpretive, aggregated)):
            continue
        body.append((descriptive, interpretive, aggregated))
    return body


def build_docx(xlsx_path: Path, docx_path: Path) -> Path:
    rows = _read_rows(xlsx_path)

    document = Document()

    # Default style: Georgia 11pt across the document.
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE_PT)

    table = document.add_table(rows=len(rows) + 1, cols=len(COLUMNS))
    table.autofit = True
    _disable_table_borders(table)

    header_cells = table.rows[0].cells
    for idx, label in enumerate(COLUMNS):
        _write_cell(header_cells[idx], label, bold=True)
        _set_cell_border(header_cells[idx], top=True, bottom=True)

    last_index = len(rows)
    for r_idx, (descriptive, interpretive, aggregated) in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        _write_cell(cells[0], descriptive)
        _write_cell(cells[1], f"{ARROW}{interpretive}" if interpretive else "")
        _write_cell(cells[2], f"{ARROW}{aggregated}" if aggregated else "")
        if r_idx == last_index:
            for cell in cells:
                _set_cell_border(cell, bottom=True)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)
    return docx_path


app = typer.Typer(add_completion=False, help=__doc__)


@app.command()
def main(
    xlsx: Path = typer.Option(
        DEFAULT_XLSX,
        "--xlsx",
        exists=True,
        readable=True,
        dir_okay=False,
        help="Path to the source code-book .xlsx file.",
    ),
    docx: Path = typer.Option(
        DEFAULT_DOCX,
        "--docx",
        dir_okay=False,
        help="Path to write the generated .docx file.",
    ),
) -> None:
    """Render the code book Excel sheet as a formatted DOCX table."""
    output = build_docx(xlsx, docx)
    size_kb = output.stat().st_size / 1024
    typer.echo(f"Wrote {output} ({size_kb:,.1f} KB)")


if __name__ == "__main__":
    app()
