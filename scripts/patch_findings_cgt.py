"""Apply CGT verb/subject realignment to Findings chapter in thesis docx."""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
UNPACKED_CUSTOM = ROOT / ".cache" / "thesis_unpacked" / "customXML"
LOG_PATH = ROOT / ".cache" / "patch_findings_cgt.log"

TABLE3_UPDATES: dict[tuple[int, int], str] = {
    (1, 2): "Enabling behavior",  # Experimenting aggregated
    (8, 1): "Mobilising infrastructure & access",
    (9, 1): "Mobilising infrastructure & access",
    (10, 1): "Mobilising infrastructure & access",
    (11, 1): "Building specification capability",
    (11, 2): "Mobilising internal conditions",
    (4, 2): "Enabling behavior",  # Bringing people along
    (13, 1): "Enabling organizational conditions",
    (8, 2): "Mobilising internal conditions",
    (14, 1): "Navigating blocked conditions",
    (15, 1): "Navigating blocked conditions",
}


def repair_docx_if_needed(path: Path) -> None:
    try:
        Document(str(path))
        return
    except KeyError as exc:
        if "customXml" not in str(exc):
            raise
    backup = path.with_suffix(".docx.bak")
    shutil.copy2(path, backup)
    with zipfile.ZipFile(path, "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist()}
    custom_files = [
        "customXml/item1.xml",
        "customXml/itemProps1.xml",
        "customXml/_rels/item1.xml.rels",
    ]
    for rel_path in custom_files:
        if rel_path not in entries:
            src = UNPACKED_CUSTOM / Path(rel_path).name
            if rel_path.endswith(".rels"):
                src = UNPACKED_CUSTOM / "_rels" / "item1.xml.rels"
            if src.exists():
                entries[rel_path] = src.read_bytes()
    tmp = path.with_suffix(".docx.repairing")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in entries.items():
            zout.writestr(name, data)
    tmp.replace(path)


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())


def set_para_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    if text:
        paragraph.add_run(text)


def insert_paragraph_after(
    paragraph: Paragraph, text: str, style: str = "normal"
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    style_map = {
        "h1": "Heading 1",
        "h2": "Heading 2",
        "h3": "Heading 3",
        "h4": "Heading 4",
        "normal": "Normal",
    }
    try:
        new_para.style = style_map.get(style, "Normal")
    except (KeyError, ValueError):
        new_para.style = "Normal"
    if text:
        new_para.add_run(text)
    return new_para


def remove_following_until(start: Paragraph, stop: Paragraph) -> None:
    nxt = start._element.getnext()
    while nxt is not None and nxt is not stop._element:
        to_remove = nxt
        nxt = nxt.getnext()
        to_remove.getparent().remove(to_remove)


def find_heading(doc: Document, section: str) -> Paragraph:
    """Find a chapter section heading (e.g. '4.2') but not subsections ('4.2.1')."""
    pattern = re.compile(rf"^{re.escape(section)}[\t ]")
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        if pattern.match(p.text.strip()):
            return p
    raise ValueError(f"Section heading not found: {section}")


def find_para(doc: Document, needle: str, *, heading: bool = False) -> Paragraph:
    for p in doc.paragraphs:
        if needle.lower() not in normalize(p.text):
            continue
        if heading and not p.style.name.startswith("Heading"):
            continue
        return p
    raise ValueError(f"Paragraph not found: {needle[:80]}")


def find_next_heading(doc: Document, after: Paragraph) -> Paragraph:
    found = False
    for p in doc.paragraphs:
        if p._element is after._element:
            found = True
            continue
        if found and p.style.name.startswith("Heading"):
            return p
    raise ValueError("No following heading found")


Block = tuple[str, str]  # kind, text — kind: h2|h3|body|quote|attr


def insert_blocks(after: Paragraph, blocks: list[Block]) -> Paragraph:
    current = after
    for kind, text in blocks:
        style = {"h2": "h2", "h3": "h3", "body": "normal", "quote": "normal", "attr": "normal"}[
            kind
        ]
        current = insert_paragraph_after(current, text, style)
    return current


def replace_section(after_heading: Paragraph, stop_heading: Paragraph, blocks: list[Block]) -> None:
    remove_following_until(after_heading, stop_heading)
    insert_blocks(after_heading, blocks)


def patch_intro(doc: Document) -> None:
    set_para_text(
        find_para(doc, "The findings indicate a wide range"),
        "The findings indicate a wide range of applications and use cases of agentic AI in the "
        "marketing field. Managerial behaviors that enable adoption and mediate value creation "
        "emerged through iterative coding (Charmaz, 2014) as interpretive categories—not a "
        "pre-designed competency model. The chapter follows how managers observe external "
        "conditions, experiment and guide to enable adoption, navigate internal conditions "
        "(mobilising what is available or working around what is blocked), apply agentic AI, "
        "and capture value outcomes—including paradoxes that recur across these stages.",
    )
    set_para_text(
        find_para(doc, "Below we first describe the impulse"),
        "Below we first describe external conditions managers observe (§4.1). We then examine "
        "enabling work—experimenting and guiding the organization (§4.2.1–2)—before treating "
        "internal conditions as one analytical unit in two states: mobilising infrastructure, "
        "specification capability, and organizational alignment when present (§4.2.3), and "
        "navigating the same subjects when absent or blocked (§4.2.4). From there we describe "
        "how agentic AI is applied (§4.3), the value portfolio that results (§4.4), and "
        "structural paradoxes (§4.5).",
    )
    set_para_text(
        find_para(doc, "Lastly we note a set of observed paradoxical"),
        "Lastly we note paradoxical patterns in which agentic AI simultaneously creates and "
        "constrains the value it is meant to deliver.",
    )
    fig = find_para(doc, "How managerial interventions drive", heading=True)
    set_para_text(
        fig,
        "How managerial work links observation to value creation with agentic AI. Enabling "
        "behaviors (experimenting, guiding) precede navigating internal conditions—mobilising "
        "or working around the same three subjects—before applying AI and capturing outcomes. "
        "Bold arrows indicate managerial behaviors; boxes represent conditions and outcomes.",
    )
    fig_next = find_para(doc, "4.1", heading=True)
    reading = insert_paragraph_after(
        fig,
        "Reading guide: verbs name what managers do; subjects name the conditions they work "
        "with. The state of an internal condition selects mobilising versus navigating—it is "
        "not a separate taxonomy of obstacles.",
        "normal",
    )
    _ = reading


def patch_table3(doc: Document) -> None:
    if len(doc.tables) < 4:
        raise RuntimeError(f"Expected at least 4 tables, found {len(doc.tables)}")
    table = doc.tables[3]
    for (row, col), value in TABLE3_UPDATES.items():
        cell = table.rows[row].cells[col]
        cell.text = ""
        if value:
            cell.paragraphs[0].add_run(value)


def patch_section_41(doc: Document) -> None:
    set_para_text(
        find_para(doc, "Managers observe this pattern and might leverage"),
        "Managers observe this pattern; how they respond is developed in §4.2.",
    )
    set_para_text(
        find_para(doc, "We will note how managers can leverage the progression"),
        "How managers respond to this progression is developed in §4.2; it is also relevant to "
        "note that continuous progress is not purely beneficial. As interviewee 4 noted:",
    )

    market_h = find_para(doc, "4.1.2", heading=True)
    supplier_h = find_para(doc, "4.1.3", heading=True)
    replace_section(
        market_h,
        supplier_h,
        [
            (
                "body",
                "A further dimension of market pressure comes from shifting consumer behaviour. "
                "Consumers are increasingly deploying AI agents of their own—tools that conduct "
                "research, compare products, and route purchase decisions on their behalf. "
                "Interviewee 12 described consumer-owned agents as the most structurally "
                "disruptive development for marketers: they act on the customer's behalf and "
                "bypass the firm's traditional points of engagement. Most organizations in the "
                "data have not yet encountered this shift at scale, but it is widely "
                "anticipated and is already prompting investment in new channel designs and "
                "agentic customer engagement models.",
            ),
            (
                "body",
                "A parallel shift is already observable in how consumers discover products and "
                "information online. AI-powered search is disrupting the traffic channels that "
                "marketing has historically relied on. Interviewee 14 observed:",
            ),
            (
                "quote",
                '"the world of search is now changing. And fortunately it\'s not going as fast as '
                "everyone thought it would change. But the fact that information of course "
                "through LLMs — ultimately in the bots that are contained within those LLMs — "
                'read those agents — yes, they of course need to do that web scraping."',
            ),
            ("attr", "— Interviewee 14"),
            (
                "body",
                "For marketing managers, this shift translates directly into a new operational "
                "concern: visibility within AI systems rather than only within traditional "
                "search engines. Interviewee 3 described this as an active focus area for her "
                'team this year, observing that organic search and paid search are "becoming '
                'less and less relevant" as AI-powered queries grow.',
            ),
            (
                "body",
                "Looking further ahead, some interviewees described a shift that is not yet "
                "operational at scale but is already shaping how they think and invest. "
                "Interviewee 15 articulated what she sees as the coming form of consumer "
                "engagement with agentic AI:",
            ),
            (
                "quote",
                '"there\'s no longer doing business with a human. But there\'s a bot, a project, '
                "an artifact, whatever we're going to call it. Yes, that goes to the platform. "
                "The platform therefore probably also needs to look different. And it's going "
                "to search for information, select a product, make comparisons. So that for me "
                'is the further form of agentic AI."',
            ),
            ("attr", "— Interviewee 15"),
            (
                "body",
                "This scenario is anticipated rather than observed: most organizations in the "
                "data have not yet encountered consumer-owned purchasing agents at scale. Its "
                "importance lies in how it is already motivating concrete choices—including "
                "Interviewee 15's own decision to build an MCP connector that enables agentic "
                "re-ordering (§4.1.3), developed precisely because this future appeared close "
                "enough to warrant immediate action.",
            ),
        ],
    )

    set_para_text(supplier_h, "4.1.3\tSupplier communication")
    sec42 = find_heading(doc, "4.2")
    replace_section(
        supplier_h,
        sec42,
        [
            (
                "body",
                "A third external force shaping managers' engagement with agentic AI is the "
                "strategy and communication of their software suppliers. Many organizations "
                "adopt AI not through internal initiative but through the road maps of the "
                "vendors they already depend on. Interviewee 1 noted that \"most of the thinking "
                "behind an AI strategy comes from their default supplier — it's Microsoft's AI "
                "strategy or IBM's AI, because that is their source of knowledge.\" Interviewee "
                "2, at a national sports federation, described a similar pattern: \"we don't have "
                "much... mainly relying on the products that we use and agents that they provide "
                "us with.\" This dependency means that supplier road maps function as a de-facto "
                "external trigger: when a supplier ships a new connector, agent template, or "
                "native integration, it makes a previously abstract use case immediately "
                "actionable.",
            ),
            (
                "body",
                "Concrete supplier releases act as direct triggers for new initiatives. New "
                "technical standards — particularly MCP-compatible connectors and integrations "
                "— are routinely turning abstract AI capability into immediately deployable use "
                "cases. Several participants described a specific supplier release as the "
                "proximate trigger for a new agentic initiative rather than an internally "
                "generated idea. Interviewee 15 described how a newly available MCP connector "
                "enabled Claude to read and action orders directly from the company's webshop "
                "catalogue, creating a genuinely new customer channel: \"Look, you can couple "
                "connectors to Claude... I see that as the future, so to speak. It's a new "
                "channel, we need to facilitate that.\" In this pattern, suppliers do not merely "
                "provide tools — they also define which use cases become organizationally "
                "legible as 'agentic.'",
            ),
            (
                "body",
                "Together, these external conditions motivate the enabling work—experimenting "
                "and guiding adoption—developed in §4.2.",
            ),
        ],
    )


def section_42_blocks() -> list[Block]:
    return [
        ("h2", "4.2\tNavigating internal conditions"),
        (
            "body",
            "Internal conditions determine whether agentic AI moves from observation to "
            "implementation. The same three subjects recur throughout the data—infrastructure "
            "and access, specification capability, and organizational alignment—but the "
            "managerial verb depends on their state. When a condition is present and "
            "mobilisable, managers leverage it; when it is absent or blocked, they navigate "
            "gaps and work-arounds. §4.2.3 and §4.2.4 tell one story in two states, not "
            "independent taxonomies. Upstream enabling work—experimenting (§4.2.1) and guiding "
            "the organization (§4.2.2)—creates the conditions under which mobilising becomes "
            "possible at scale.",
        ),
        ("h3", "4.2.1\tExperimenting"),
        (
            "body",
            "Experimenting is the primary mode through which marketing managers discover where "
            "agentic AI creates value. Three patterns recur across the data: bounded 'lighthouse' "
            "projects designed to prove a use case before scaling; continuous personal "
            "experimentation that builds personal capability; and structured training programmes "
            "that help teams use agentic AI in a way that drives value.",
        ),
        (
            "body",
            "The lighthouse project pattern is the most deliberately designed. By limiting scope "
            "to a single data set or a clearly bounded workflow, a lighthouse project avoids "
            "triggering multi-committee approval processes (discussed in §4.2.4) while generating "
            "the concrete evidence that leadership and compliance require before authorizing "
            "broader rollout. Interviewee 13, an external AI advisor working extensively with "
            "German enterprises, described the approach:",
        ),
        (
            "quote",
            "I typically try to do y— lighthouse projects, right? … So don't do anything, uh, "
            "with all the data at the company. Don't put it to the whole IT, security, legal "
            "stuff. Just go to legal, say, 'Hey, we have this data set, and with this kind of "
            "data, we want to do something.' … And usually before that, I need to get the buy-in "
            "from the stakeholder, from the people in the company, but they usually get me "
            "because they want to build something like this.",
        ),
        ("attr", "— Interviewee 13"),
        (
            "body",
            "Personal experimentation is the most prevalent form across all participant types. "
            "Participants who report the strongest results describe daily hands-on use as the "
            "primary source of practical capability. Interviewee 7, an AI advisor who works "
            "across multiple sectors, captured the decisive disposition:",
        ),
        (
            "quote",
            "those who are really flying with AI, and I would think you are part of it, is when "
            "you're dealing with AI, the, the answer of the AI might not be the last answer. You "
            "say, \"No, I don't like this answer, let's move on.\" And then it becomes a "
            "communication until you get what you want.",
        ),
        ("attr", "— Interviewee 7"),
        (
            "body",
            "This iterative back-and-forth is both a skill and a disposition, one that training "
            "can introduce but only practice can entrench.",
        ),
        (
            "body",
            "Training programs are the organizational complement to individual experimentation. "
            "Several participants described structured learning as essential to moving teams "
            "beyond one-off prompt use toward repeatable, embedded workflows. Interviewee 1 noted "
            "that the appetite for foundational AI education is high but that it needs to be "
            "grounded in specific organizational contexts rather than generic tool tutorials: "
            "\"a lot of it has to do with the baseline understanding … there's a huge appetite "
            "for what AI is or what AI does that needs to be clarified before any further "
            "engagement.\"",
        ),
        (
            "body",
            "A recurring obstacle in the experimenting phase is the difficulty of finding "
            "product-market fit for AI solutions. Interviewee 5 described this in consumer-facing "
            "AI: the hardest constraint is not what the technology can do but what customers "
            "actually want, and whether the right customer data is even collectible. Finding the "
            "specific use case that is both technically feasible and genuinely valued requires "
            "the same market-sensing work as any other product development process — and often "
            "takes longer than the technology itself.",
        ),
        ("h3", "4.2.2\tGuiding the organization"),
        (
            "body",
            "Guiding the organization involves four interconnected managerial behaviors: "
            "providing clarity about direction and expectations; actively bringing people along "
            "through communication and training; securing and sustaining leadership support; "
            "and mobilizing AI champions who demonstrate what is possible rather than lobbying "
            "informally for permission.",
        ),
        (
            "body",
            "Clarity is foundational. Without a clear direction from leadership, teams experiment "
            "idiosyncratically and the resulting knowledge cannot be consolidated or compared. "
            "Interviewee 1, an AI advisor with board-level experience across multiple "
            "organizations, framed clarity as a precondition that precedes execution through a "
            "CARE model—clarity, awareness, readiness, and execution:",
        ),
        (
            "quote",
            "I have a simple model … It's called the abbreviations are CARE: C-A-R-E. And that "
            "stands for clarity, awareness, readiness, and execution. And I think if you don't "
            "have clarity, jumping into execution will only surface a lot of problems there.",
        ),
        ("attr", "— Interviewee 1"),
        (
            "body",
            "Where that clarity is absent, ambition diffuses across incompatible experiments "
            "rather than accumulating into organizational capability.",
        ),
        (
            "body",
            "Bringing people along — active change management rather than passive communication "
            "— is the most labor-intensive guiding behavior in the data. Interviewee 9, who leads "
            "AI transformation across an agency network, described the operationalization burden "
            "plainly:",
        ),
        (
            "quote",
            "People don't know it, so they just kinda go back to doing their own thing until "
            "someone puts it right in front of them and says, 'This is what it is. This is how "
            "you use it in your day-to-day.' Um, so I think that's honestly my role.",
        ),
        ("attr", "— Interviewee 9"),
        (
            "body",
            "Leadership support functions both as a resource and as a signal. Where leaders "
            "actively champion AI initiatives, teams describe materially stronger adoption than "
            "the technology alone predicts. Interviewee 6 captured the champion posture: \"No, "
            "this is the vision. This is where we're going. Help me get there.\" Interviewee 4 "
            "described the human side of that signal when pace accelerates: \"people are drinking "
            "from a fire hose\" and need leadership that affirms \"their importance, they're part "
            "of the team, and their work matters.\" Interviewee 17 emphasized that alignment "
            "requires a safe environment for experimentation:",
        ),
        (
            "quote",
            "you have to give the right people the feeling that they're allowed to experiment and "
            "that they're allowed to do things in a safe environment.",
        ),
        ("attr", "— Interviewee 17"),
        (
            "body",
            "AI champions make abstract capability concrete. Rather than absorbing organizational "
            "friction through informal governance, champions in the data build proof first and "
            "then widen adoption. Interviewee 8 described testing new capabilities quickly and "
            "then securing buy-in: after a five-second test tells him whether a tool is ready, "
            "\"I then try to get buy-in.\" He evangelizes through training sessions that show "
            "colleagues agents they can build themselves. Interviewee 9 described a similar "
            "proof-first pattern with Affinity Hub—demonstrating value before asking teams to "
            "change daily practice.",
        ),
        (
            "body",
            "Scaling requires moving from proof to organizational rollout. Interviewee 6 "
            "described how commercial success in marketing enabled a company-wide programme—"
            "Infinigym—with twenty prioritized agentic use cases rolled out to facility "
            "management, supplier follow-up, and other departments after marketing had "
            "demonstrated results. The pattern across interviews is consistent: guiding work "
            "converts experiments into shared direction, psychological safety, and visible proof "
            "that others can adopt.",
        ),
        ("h3", "4.2.3\tMobilising internal conditions"),
        (
            "body",
            "When internal conditions are present, managers mobilise them to move agentic AI from "
            "experiment to embedded practice. Three subjects structure the data: infrastructure "
            "and access; specification capability; and organizational conditions. External expert "
            "networks can accelerate learning but are not treated here as internal resources.",
        ),
        (
            "body",
            "Infrastructure and access combine data, tooling, budgets, and integration capacity. "
            "Data availability and quality is the most frequently cited enabler. Without "
            "accessible, well-structured data, even nominally AI-capable tools remain unused. "
            "Interviewee 13 described a pattern he encounters repeatedly across marketing clients:",
        ),
        (
            "quote",
            "if I look at my customers, there's now tools that have AI use cases like Emarsys, "
            "Braze and so on, but people still don't use this because there's no time to put the "
            "right data in from the data warehouse. There's like a long list of backlog, and "
            "there's no one in the team itself that understands how to really bring their "
            "marketing knowledge and this IT data knowledge together.",
        ),
        ("attr", "— Interviewee 13"),
        (
            "body",
            "Where formal IT backlogs stall progress, some managers treat bounded shadow "
            "experimentation as a pragmatic bypass—lighthouse projects on limited data sets that "
            "demonstrate value before enterprise integration. Interviewee 13 noted that if "
            "organizations refuse this path, \"nothing happens,\" even when consumer-grade tools "
            "could prove a use case in days.",
        ),
        (
            "body",
            "Specification capability is the ability to decompose workflows and define what good "
            "agentic output looks like. Skill gaps appear at both ends: no one can wire "
            "integrations or evaluate model output, and no one on the marketing side can specify "
            "requirements. Interviewee 13 labelled the failure mode 'unoperationalized tool "
            "effectiveness'—buying tools that no one has the skills or accountability to embed. "
            "Process thinking is the deeper resource: organizations that can map data, decision, "
            "and action layers specify viable use cases; those that cannot default to vague "
            "ambitions ('we need AI'). Interviewee 6 made the cross-functional dependency "
            "explicit: \"The next CMO needs to be CTO's best friend … without the technology, I "
            "wouldn't be able to...\"",
        ),
        (
            "body",
            "Organizational conditions—alignment between marketing and technology leadership, "
            "culture, and focus—determine whether mobilised infrastructure and specification "
            "translate into rollout. Organizations with cultures that tolerate fast failure and "
            "reward learning iterate more quickly through the experimentation cycle in §4.2.1. "
            "Focus—choosing a small number of specific use cases and resourcing them "
            "adequately—distinguishes organizations that make progress from those that accumulate "
            "a long list of stalled experiments.",
        ),
        ("h3", "4.2.4\tWhen conditions block progress"),
        (
            "body",
            "The same three subjects appear when conditions are absent or blocked. Resistance, "
            "governance delay, and analysis paralysis are properties within these mirrors—not a "
            "parallel taxonomy. Where §4.2.3 describes assets managers can mobilise, §4.2.4 "
            "describes how their absence or obstruction stalls agentic AI.",
        ),
        (
            "body",
            "Infrastructure and access blocked. Data-warehouse backlogs, integration committees, "
            "and IT capacity constraints reproduce the dormant-tool pattern from §4.2.3 at "
            "organizational scale. Interviewee 13's committee example—three approvals for a "
            "Google Analytics integration—shows how governance designed for risk can block "
            "routine data access. The mirror image is shadow AI: the same formal caution is "
            "rarely applied to unsanctioned consumer-grade tools employees run on their own "
            "desktops, creating a parallel layer outside governance.",
        ),
        (
            "body",
            "Specification capability blocked. Without process thinking, teams stall in analysis "
            "paralysis or generic 'we need AI' mandates that never operationalize. Interviewee "
            "13's 'unoperationalized tool effectiveness' label applies equally here: tools are "
            "procured but no one owns the workflow specification that would make them valuable. "
            "Interviewee 12 named the asymmetry of delay in a fast-moving field—standing still "
            "is itself a form of risk-taking, not a conservative default (developed further in "
            "§4.5.3). Interviewee 6 illustrated organizational delay when an internal build "
            "preference imposed a two-year wait for IT backlog capacity on a chatbot initiative.",
        ),
        (
            "body",
            "Organizational conditions blocked. Resistance rarely appears as explicit anti-AI "
            "opposition; it manifests as loss aversion—fear that automation will reduce role "
            "relevance or team size. Interviewee 13 offered a characterization echoed across "
            "the corpus:",
        ),
        (
            "quote",
            '"Germans are not risk averse. Germans are loss averse." And I think there\'s also '
            'this like, "Oh, if it\'s not automated, what, what is my team doing?" So that\'s '
            "like, it's always stupid discussions about this. And at this customer, even when we "
            "made it, it took, uh, three different committees to go through. I mean, this was a "
            "stupid tool just to, hey, this is Google Analytics, get me the data. So it needed to "
            "go through the IT committee, to the cloud committee, to the AI committee,",
        ),
        ("attr", "— Interviewee 13"),
        (
            "body",
            "Punishing failure amplifies loss aversion. Interviewee 7 described the opposite of "
            "psychological safety: when a pilot goes wrong, \"they say, 'Hey, Don, you stupid. "
            "Duh, it cannot be. Next time, you're fired.'\" Interviewee 8 noted that corporate "
            "politics—not model quality—often determines whether a promising test receives "
            "broader buy-in.",
        ),
        (
            "body",
            "A lack of AI understanding at key decision levels is distinct from resistance: "
            "leadership supports adoption without the conceptual grounding to govern it. "
            "Interviewee 1, who regularly engages with boards on AI strategy, observed:",
        ),
        (
            "quote",
            "comfort that boards have, and executive in general have with the topic of AI in my "
            "books, does not necessarily reflect their understanding.",
        ),
        ("attr", "— Interviewee 1"),
        (
            "body",
            "The consequence is an organization formally committed to AI while structurally "
            "unable to calibrate conversations about risk, pace, or investment—a configuration "
            "that amplifies every other blocker in this section.",
        ),
    ]


def patch_section_42(doc: Document) -> None:
    sec42_old = find_para(doc, "Affecting change", heading=True)
    sec43 = find_heading(doc, "4.3")
    replace_section(sec42_old, sec43, section_42_blocks())


def patch_cross_refs(doc: Document) -> None:
    applying = find_heading(doc, "4.3")
    nxt = find_next_heading(doc, applying)
    intro = None
    for p in doc.paragraphs:
        if p._element is applying._element:
            continue
        if p._element is nxt._element:
            break
        if p.style.name == "Normal" and p.text.strip():
            intro = p
            break
    if intro:
        set_para_text(
            intro,
            intro.text
            + " Use cases materialize when the internal conditions in §4.2.3 are met; they stall "
            "or stay cosmetic when blockers in §4.2.4 dominate.",
        )

    optics = find_para(doc, "Across the interviews, a consistent gap emerges")
    set_para_text(
        optics,
        optics.text
        + " Organizational blockers in §4.2.4—especially weak specification and loss aversion—"
        "help explain why positioning outruns production; the prototype-to-scale gap is developed "
        "in §4.5.3.",
    )

    strategies = find_para(doc, "On the strategy side, four patterns recur")
    set_para_text(
        strategies,
        strategies.text
        + " These strategies assume specification capability (§4.2.3): determining input, "
        "evaluating output, and engineering context are only valuable when teams can define what "
        "good agentic work looks like.",
    )

    value_intro = find_para(doc, "Value outcomes through agentic AI is not a single outcome")
    set_para_text(
        value_intro,
        "Value outcomes through agentic AI is not a single outcome but a portfolio of effects. "
        "The data reveal benefits, sacrifices, and risks that co-occur rather than trade off "
        "cleanly. Which part of the portfolio is captured depends on whether internal conditions "
        "were mobilised (§4.2.3) or blockers dominated (§4.2.4), mediated by experimenting and "
        "guiding (§4.2.1–2).",
    )

    mediating = find_para(doc, "The data make the mediating role of managerial behavior visible")
    if mediating:
        set_para_text(mediating, "")

    paradox = find_para(doc, "Running through the data, a recurring structural pattern emerges")
    set_para_text(
        paradox,
        paradox.text
        + " The same conditional logic structures §4.2.4: managers navigate design tensions when "
        "internal conditions block progress.",
    )


def patch_coding_placeholder(doc: Document) -> None:
    placeholder = find_para(doc, "Placeholder for coding table", heading=True)
    nxt = find_next_heading(doc, placeholder)
    replace_section(
        placeholder,
        nxt,
        [
            ("h2", "4.6\tCoding development"),
            (
                "body",
                "Table 2 shows how descriptive codes consolidated into interpretive and aggregated "
                "categories through iterative coding (Charmaz, 2014). Saturation was reached at the "
                "category level—observing external conditions, enabling adoption through "
                "experimenting and guiding, navigating internal conditions (mobilised or blocked), "
                "applying agentic AI, and value outcomes—rather than at every descriptive code. "
                "Aggregated labels such as Mobilising internal conditions and Experiencing obstacles "
                "name two states of the same analytical unit; interpretive labels identify the "
                "subjects within each state. This emergent structure reflects the analytical process "
                "documented in Methodology §3.4.",
            ),
        ],
    )


def main() -> None:
    repair_docx_if_needed(DOCX_PATH)
    doc = Document(str(DOCX_PATH))
    log: list[str] = []

    patch_intro(doc)
    log.append("intro + figure caption + reading guide")
    patch_table3(doc)
    log.append("table 3 coding labels")
    patch_section_41(doc)
    log.append("§4.1 restore + light edits")
    patch_section_42(doc)
    log.append("§4.2 full rewrite")
    patch_cross_refs(doc)
    log.append("§4.3–§4.5 cross-refs")
    patch_coding_placeholder(doc)
    log.append("§4.6 coding development")

    doc.save(str(DOCX_PATH))
    LOG_PATH.write_text("\n".join(log), encoding="utf-8")
    print(f"Patched {DOCX_PATH}")
    for line in log:
        print(f"  - {line}")


if __name__ == "__main__":
    main()
