"""Fix regressions from initial CGT patch."""

from __future__ import annotations

from docx import Document
from docx.text.paragraph import Paragraph

from patch_findings_cgt import (
    DOCX_PATH,
    find_heading,
    find_para,
    insert_blocks,
    insert_paragraph_after,
    normalize,
    set_para_text,
)


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def find_next_body(doc: Document, after: Paragraph) -> Paragraph:
    found = False
    for p in doc.paragraphs:
        if p._element is after._element:
            found = True
            continue
        if found and not p.style.name.startswith("Heading") and p.text.strip():
            return p
    raise ValueError("No body paragraph found after heading")


def main() -> None:
    doc = Document(str(DOCX_PATH))

    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and "Affecting change" in p.text:
            remove_paragraph(p)
            break

    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and "Placeholder for coding" in p.text:
            remove_paragraph(p)
            break

    short_intro = normalize(
        "Managers observe this pattern; how they respond is developed in §4.2."
    )
    full_intro = (
        "The first external condition that the marketing manager might observe is AI "
        "progression. This concept describes how AI is evolving; improving in quality, "
        "gaining new capabilities, and becoming more widely adopted. Managers observe "
        "this pattern; how they respond is developed in §4.2."
    )
    for p in doc.paragraphs:
        if normalize(p.text) == short_intro:
            set_para_text(p, full_intro)
            break

    market_h = find_para(doc, "4.1.2", heading=True)
    if "A second external condition" not in find_next_body(doc, market_h).text:
        insert_blocks(
            market_h,
            [
                (
                    "body",
                    "A second external condition that is frequently observed by marketing managers "
                    "and affects the adoption of agentic AI is market pressure. We define market "
                    "pressure broadly as a combination of two categories: 1. pressure from "
                    "competitors 2. customer needs or demands. Interviewees note shifts in both "
                    "categories.",
                ),
                (
                    "body",
                    "Interviewee 9, noted how different marketing agencies are looking at each "
                    "other's offerings and trying to mimic what others are doing. She noted that "
                    "organizations are primarily trying \"to get to parity in the present with "
                    "what they're saying they can do.\" Interviewee 13 noted the value of "
                    "leveraging insights from competitors to gain insights",
                ),
                (
                    "quote",
                    "\"Talk to your peers at conferences. I think this needs to happen more on the "
                    "C level, that they exchange ideas with their peers, because that's actually "
                    "where they will understand the quickest what the benefit is, and this makes "
                    "it a lot better, right?\"",
                ),
                ("attr", "— Interviewee 13"),
                (
                    "body",
                    "Participants described organizations feeling compelled to demonstrate AI "
                    "capability, primarily to avoid the perception of falling behind rather than "
                    "to generate novel value. The dominant managerial posture is defensive rather "
                    "than offensive. Interviewee 8, who maintains an active AI practice in "
                    "e-commerce and digital marketing, captured the organizational reaction to "
                    "new AI initiatives:",
                ),
                (
                    "quote",
                    "Almost everyone's been very positive about it because they want to not be "
                    "left behind, you know?",
                ),
                ("attr", "— Interviewee 8"),
            ],
        )

    applying = find_heading(doc, "4.3")
    first = find_next_body(doc, applying)
    if "materialize" not in first.text and "Four use-case archetypes" in first.text:
        insert_paragraph_after(
            applying,
            "Agentic use cases materialize when the internal conditions in §4.2.3 are met; they "
            "stall or stay cosmetic when blockers in §4.2.4 dominate.",
            "normal",
        )

    optics = find_para(doc, "Across the interviews, a consistent gap emerges")
    if "§4.2.4" not in optics.text:
        set_para_text(
            optics,
            optics.text.rstrip()
            + " Organizational blockers in §4.2.4—especially weak specification and loss "
            "aversion—help explain why positioning outruns production; the prototype-to-scale "
            "gap is developed in §4.5.3.",
        )

    strategies = find_para(doc, "On the strategy side, four patterns recur")
    if "§4.2.3" not in strategies.text:
        set_para_text(
            strategies,
            strategies.text.rstrip()
            + " These strategies assume specification capability (§4.2.3): determining input, "
            "evaluating output, and engineering context are only valuable when teams can define "
            "what good agentic work looks like.",
        )

    paradox = find_para(doc, "Running through the data, a recurring structural pattern emerges")
    if "§4.2.4" not in paradox.text:
        set_para_text(
            paradox,
            paradox.text.rstrip()
            + " The same conditional logic structures §4.2.4: managers navigate design tensions "
            "when internal conditions block progress.",
        )

    doc.save(str(DOCX_PATH))
    print(f"Fixed {DOCX_PATH}")


if __name__ == "__main__":
    main()
