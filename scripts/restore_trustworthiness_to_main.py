"""Restore §3.4 trustworthiness (Shenton table) into the main thesis docx."""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
MAIN = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
SOURCE = ROOT / "Thesis Draft - Daan Luttik - MBA - trustworthiness-shenton.docx"


def _run_patch() -> None:
    subprocess.run(
        [sys.executable, str(ROOT / "scripts" / "patch_trustworthiness_shenton.py")],
        check=True,
        cwd=str(ROOT),
    )


def _has_trust_table(path: Path) -> bool:
    from docx import Document

    doc = Document(str(path))
    for table in doc.tables:
        if table.rows and "Criteria of trustworthiness" in table.rows[0].cells[0].text:
            return True
    return False


def main() -> None:
    # Ensure source/main content is current
    _run_patch()
    if SOURCE.exists() and _has_trust_table(SOURCE) and not _has_trust_table(MAIN):
        shutil.copy2(SOURCE, MAIN)
        print(f"Copied trustworthiness section from {SOURCE.name} -> {MAIN.name}")
        return

    if _has_trust_table(MAIN):
        print(f"§3.4 trustworthiness table is already present in {MAIN.name}")
        return

    # Missing table: insert full section before Findings
    from convert_trustworthiness_table import _insert_new_section, _find_findings_paragraph

    from docx import Document

    doc = Document(str(MAIN))
    _insert_new_section(doc, _find_findings_paragraph(doc))
    _run_patch()
    doc = Document(str(MAIN))
    try:
        doc.save(str(MAIN))
    except PermissionError:
        shutil.copy2(SOURCE if SOURCE.exists() else MAIN, SOURCE)
        print(f"Main file locked. Close Word, then run: Copy-Item -Force '{SOURCE}' '{MAIN}'")
        raise SystemExit(1) from None
    print(f"Inserted and patched §3.4 in {MAIN.name}")


if __name__ == "__main__":
    main()
