# -*- coding: utf-8 -*-
"""Targeted Ch4/5/6 refinements (tracked, author 'Claude'): C1, C3, C4, C6, C7 + Ch4 lab/factory.
(§5.2 consolidation, C2, is a separate script.) All-or-nothing save."""
from __future__ import annotations
import copy, shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
BACKUP = ROOT / "Thesis Draft - Daan Luttik - MBA.ch456-backup.docx"
AUTHOR = "Claude"; DATE = "2026-06-07T19:00:00Z"
XS = "{http://www.w3.org/XML/1998/namespace}space"

def used_ids(doc):
    return {int(el.get(qn("w:id"), 0)) for el in doc.element.body.iter() if el.get(qn("w:id")) is not None}
def nid(u): n = max(u, default=0) + 1; u.add(n); return n
def atext(el): return "".join(t.text or "" for t in el.iter(qn("w:t")))

def para(doc, marker, style=None):
    for p in doc.paragraphs:
        if marker in atext(p._element) and (style is None or p.style.name == style):
            return p._element
    return None
def para_obj(doc, marker, style=None):
    for i, p in enumerate(doc.paragraphs):
        if marker in atext(p._element) and (style is None or p.style.name == style):
            return i, p
    return None, None

def ins_run(text, used):
    ins = etree.Element(qn("w:ins")); ins.set(qn("w:id"), str(nid(used)))
    ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), DATE)
    r = etree.SubElement(ins, qn("w:r")); t = etree.SubElement(r, qn("w:t"))
    t.set(XS, "preserve"); t.text = text
    return ins

def append_ins(p, text, used):
    p.append(ins_run(text, used))

def prepend_ins(p, text, used):
    pPr = p.find(qn("w:pPr"))
    ins = ins_run(text, used)
    if pPr is not None: pPr.addnext(ins)
    else: p.insert(0, ins)

def del_run(r, used):
    d = etree.Element(qn("w:del")); d.set(qn("w:id"), str(nid(used)))
    d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), DATE)
    rc = copy.deepcopy(r)
    for t in rc.findall(qn("w:t")):
        dt = etree.Element(qn("w:delText")); dt.set(XS, "preserve"); dt.text = t.text or ""
        t.getparent().replace(t, dt)
    d.append(rc); return d

def del_para(p, used):
    pPr = p.find(qn("w:pPr"))
    if pPr is None: pPr = etree.Element(qn("w:pPr")); p.insert(0, pPr)
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None: rPr = etree.SubElement(pPr, qn("w:rPr"))
    dm = etree.Element(qn("w:del")); dm.set(qn("w:id"), str(nid(used)))
    dm.set(qn("w:author"), AUTHOR); dm.set(qn("w:date"), DATE); rPr.insert(0, dm)
    for r in p.findall(qn("w:r")):
        p.replace(r, del_run(r, used))

def flatten_para(p):
    """Accept all tracked changes within a paragraph so the target text becomes bare runs
    (used only on paragraphs whose prior pending edits overlap a new replacement; ok per C8)."""
    for ins in p.findall(qn("w:ins")):
        idx = list(p).index(ins)
        for child in list(ins):
            p.insert(idx, child); idx += 1
        p.remove(ins)
    for d in p.findall(qn("w:del")):
        p.remove(d)
    pPr = p.find(qn("w:pPr"))
    if pPr is not None:
        rPr = pPr.find(qn("w:rPr"))
        if rPr is not None:
            for dm in rPr.findall(qn("w:del")):
                rPr.remove(dm)

def replace_frag(p, old, new, used):
    runs = p.findall(qn("w:r")); segs, total = [], ""
    for r in runs:
        t = r.find(qn("w:t")); txt = t.text if (t is not None and t.text) else ""
        segs.append((r, txt)); total += txt
    s = total.find(old)
    if s < 0: return False
    e = s + len(old)
    def mk(rPr, x):
        nr = etree.Element(qn("w:r"))
        if rPr is not None: nr.append(copy.deepcopy(rPr))
        tt = etree.SubElement(nr, qn("w:t")); tt.set(XS, "preserve"); tt.text = x; return nr
    def dl(rPr, x):
        d = etree.Element(qn("w:del")); d.set(qn("w:id"), str(nid(used)))
        d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), DATE)
        rr = etree.SubElement(d, qn("w:r"))
        if rPr is not None: rr.insert(0, copy.deepcopy(rPr))
        dt = etree.SubElement(rr, qn("w:delText")); dt.set(XS, "preserve"); dt.text = x; return d
    pos = 0; done = False
    for r, txt in segs:
        rs, re_ = pos, pos + len(txt); pos = re_
        if not txt or re_ <= s or rs >= e: continue
        rPr = r.find(qn("w:rPr")); ls = max(s, rs) - rs; le = min(e, re_) - rs
        before, mid, after = txt[:ls], txt[ls:le], txt[le:]
        parent = r.getparent(); idx = list(parent).index(r); parts = []
        if before: parts.append(mk(rPr, before))
        if not done: parts.append(ins_run(new, used)); done = True
        if mid: parts.append(dl(rPr, mid))
        if after: parts.append(mk(rPr, after))
        for j, prt in enumerate(parts): parent.insert(idx + j, prt)
        parent.remove(r)
    return done

DQ, DQR, APO, ELL = "“", "”", "’", "[…]"

def main():
    doc = Document(str(DOCX)); used = used_ids(doc)

    # ---- C1: §6 contribution paragraph ----
    p = para(doc, "The contribution of this study is a grounded account")
    assert p is not None, "C1 anchor missing"
    prepend_ins(p, f"Empirically, this study answers recent calls for research into agentic AI in "
                   f"marketing (Kim, 2025; Mogaji & Jain, 2024; Jain & Eastman, 2024): it provides one of "
                   f"the first grounded accounts of how marketing managers create value with agentic AI in "
                   f"practice, moving beyond the field{APO}s prevailing focus on content generation. ", used)
    print("C1 ok")

    # ---- C3: remove 5.1.5 (heading + body), fold one sentence into the model paragraph ----
    i_h, ph = para_obj(doc, "contested category", style="Heading 3")
    assert ph is not None, "C3 heading missing"
    body = doc.paragraphs[i_h + 1]._element  # body paragraph follows the heading
    assert "instability of the category" in atext(body), "C3 body mismatch"
    del_para(ph._element, used); del_para(body, used)
    model = para(doc, "The principal theoretical contribution of this study is a process model")
    assert model is not None, "C3 model para missing"
    append_ins(model, " The study also treats the construct agentic AI itself constructively — not as a "
                      "stable category but as a contested, boundary-object term that performs legitimating "
                      "work — an instability whose methodological consequences are discussed in Section 5.3.", used)
    print("C3 ok")

    # ---- C4: broaden §5.1.2 beyond the harness ----
    p = para(doc, "The harness was the component that participants who could articulate")
    assert p is not None, "C4 anchor missing"
    flatten_para(p)   # accept the prior §5.1.2 pending edit so the sentence is bare
    ok = replace_frag(p,
        "The harness was the component that participants who could articulate the mechanism named most directly.",
        f"The study{APO}s most technically experienced participants located this most precisely, pointing "
        f"beyond the agent{APO}s harness to the surrounding technical context: the state of the data "
        f"infrastructure and pipelines, the integration with existing systems and tooling, and the maturity "
        f"of the organization{APO}s wider stack.", used)
    assert ok, "C4 replace failed"
    print("C4 ok")

    # ---- C6: expand future research in §5.3 ----
    p = para(doc, "Future research that triangulates interview data")
    assert p is not None, "C6 anchor missing"
    flatten_para(p)   # accept the prior §5.3 pending reword so the sentence is bare
    ok = replace_frag(p,
        "Future research that triangulates interview data with observational methods, workflow documentation, "
        "or objectively measured performance indicators would substantially strengthen the empirical grounding "
        "of the theory developed here, and would be better positioned to assess where the dependence on internal "
        "conditions and configuration identified in section 4 actually produces the value creation it describes.",
        "Several directions of future research follow. The most important is to move below the high-level "
        "treatment typical of non-technical AI research and to study, for a given use case, which specific "
        "configuration, implementation, and cost choices most determine the value created — since the findings "
        "suggest that value is differentiated by context and implementation detail rather than by the use case "
        "itself. Beyond this, longitudinal designs could track initiatives from experiment to production; "
        "triangulating interview data with observational methods, workflow documentation, or objectively "
        "measured performance indicators would test where the dependence on internal conditions and configuration "
        "identified in Section 4 actually produces the value it describes; and the emerging agentic demand side "
        "— consumers acting through their own agents — warrants study as it matures.", used)
    assert ok, "C6 replace failed"
    print("C6 ok")

    # ---- C7: reflexivity sentence in §5.3 sample/scope paragraph ----
    p = para(doc, "Both factors limit how far the empirical details generalize")
    assert p is not None, "C7 anchor missing"
    append_ins(p, f" Finally, the analysis reflects a single researcher{APO}s interpretation, and the "
                  f"researcher{APO}s own background as a marketing-technology practitioner inevitably shaped "
                  f"the coding and the categories developed; consistent with constructivist grounded theory "
                  f"this positionality is treated as part of the analytic instrument rather than a flaw, and "
                  f"it was partially offset by regular debriefing with the thesis supervisor, who reviewed and "
                  f"discussed the evolving coding.", used)
    print("C7 ok")

    # ---- Ch4: laboratory/factory in the experimentation paragraph ----
    p = para(doc, "Experimenting with AI primarily addresses resistance")
    assert p is not None, "Ch4 anchor missing"
    append_ins(p, f" Interviewee 12 framed this separation explicitly as a {DQ}laboratory{DQR} and a "
                  f"{DQ}factory{DQR}: organizations need {DQ}the factory where {ELL} the trains are running "
                  f"on time{DQR} alongside {DQ}the laboratory, where you{APO}re in a position to experiment and "
                  f"try things and have things fail and break,{DQR} transferring an approach into the factory "
                  f"{DQ}only once you{APO}ve found things in the laboratory that {ELL} you{APO}ve proven.{DQR}", used)
    print("Ch4 ok")

    shutil.copy(DOCX, BACKUP); doc.save(str(DOCX))
    print(f"\nSaved: {DOCX}\nBackup: {BACKUP}")

if __name__ == "__main__":
    main()
