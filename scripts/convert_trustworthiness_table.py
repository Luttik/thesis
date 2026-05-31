"""Replace §3.4 prose subsections with a trustworthiness table (Beninger & Francis, 2021)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
DOCX_FALLBACK = ROOT / "Thesis Draft - Daan Luttik - MBA - trustworthiness-table.docx"

INTRO = (
    "Qualitative inquiry does not seek statistical generalization to a population; "
    "instead, readers assess whether findings are sufficiently grounded in data and "
    "whether the inquiry was conducted transparently. This study adopts Lincoln and "
    "Guba's (1985) trustworthiness criteria and implements them through strategies "
    "proposed by Shenton (2004) for interview-based projects of limited duration. "
    "Following Beninger and Francis (2021), the assessment is summarized in table form "
    "(Table 3). Tactics such as prolonged engagement and persistent observation, "
    "while valuable in long-term ethnography, were not feasible in this bounded study "
    "(Shenton, 2004)."
)

TABLE_CAPTION = "Table 3."
TABLE_TITLE = "Trustworthiness assessment."
TRUST_HDR = "Criteria of trustworthiness"

# (criterion definition, how addressed) — Beninger-style order
TABLE_ROWS: list[tuple[str, str]] = [
    (
        "Transferability concerns whether findings may have meaning in other contexts "
        "(Lincoln & Guba, 1985; Shenton, 2004), including sufficient detail about the "
        "fieldwork context when reporting results.",
        "Purposive and snowball sampling targeted marketing managers and AI implementation "
        "experts who could speak to agentic AI in practice (Section 3.3; Table 1). "
        "Thick description—process accounts and extensive verbatim quotation—is provided "
        "in Chapter 4. Scope conditions (geography, network sampling, and sample "
        "composition) are discussed in Section 5.3.",
    ),
    (
        "Dependability requires a consistent, auditable inquiry process (Lincoln & Guba, "
        "1985; Shenton, 2004), including transparent reporting of what was done in the field.",
        "An audit trail links interview recordings and transcripts, cleaned transcript "
        "files, versioned QDPX analysis projects (April–May 2026), analytic memos, the code "
        "structure in Table 2, and the Findings narrative. The semi-structured interview "
        "guide and protocol (Appendix A) provided a baseline; backup and validation "
        "questions (Appendix A.1.3) show how later interviews probed emerging themes. "
        "Coding can be traced from quotation-level codings through focused categories to "
        "aggregated concepts in Chapter 4.",
    ),
    (
        "Confirmability requires that interpretations are supported by the data rather "
        "than researcher preference alone (Lincoln & Guba, 1985; Shenton, 2004).",
        "Analytic memos recorded reasoning separately from transcript text. Findings "
        "claims are anchored in participant quotations; supporting excerpts are collated "
        "in Appendix D. Reflexive documentation of the researcher's marketing and AI "
        "background appears in Section 3.2.1.",
    ),
    (
        "Credibility concerns whether findings are a plausible, acceptable representation "
        "of participants' accounts (Merriam, 1998; Shenton, 2004).",
        "The study applied constructivist grounded theory (Charmaz, 2014) with iterative "
        "interviewing, memo writing, and coding (Appendix C). Credibility was strengthened "
        "through triangulation across marketing practitioners and AI experts; iterative "
        "initial and focused coding (Table 2); theoretical sampling; retention of negative "
        "and paradoxical accounts; prior field familiarity with reflexivity (Section 3.2.1); "
        "and thesis-supervisor peer scrutiny of emerging categories. Member checking was not "
        "conducted (Charmaz, 2014). Theoretical saturation was assessed during data "
        "collection: later interviews chiefly densified existing categories rather than "
        "introducing new top-level concepts, with seventeen interviews completed (Table 1).",
    ),
    (
        "Integrity concerns whether interpretations are impaired by misinformation from "
        "informants while recognizing that data never perfectly reflect phenomena studied "
        "(Wallendorf & Belk, 1989).",
        "Interviews were semi-structured, confidential, and conducted to build rapport in "
        "a non-threatening manner. Table 1 reports participants in anonymized form; "
        "Findings use interviewee labels tied to verbatim transcript evidence.",
    ),
]


def _table_style_name(doc: Document) -> str | None:
    if doc.tables:
        return doc.tables[0].style.name
    for name in ("Table Grid", "Table Normal", "Normal Table", "Tabelraster"):
        try:
            doc.styles[name]
            return name
        except KeyError:
            continue
    return None


def insert_table_after(paragraph: Paragraph, rows: list[tuple[str, str]]) -> Table:
    doc = paragraph.part.document
    table = doc.add_table(rows=1 + len(rows), cols=2)
    style_name = _table_style_name(doc)
    if style_name:
        table.style = style_name
    table.autofit = False
    table.columns[0].width = Inches(2.4)
    table.columns[1].width = Inches(4.1)

    hdr = table.rows[0].cells
    hdr[0].text = "Criteria of trustworthiness"
    hdr[1].text = "How addressed"

    for i, (crit, how) in enumerate(rows, start=1):
        row = table.rows[i].cells
        row[0].text = crit
        row[1].text = how

    tbl_el = table._tbl
    body = doc.element.body
    body.remove(tbl_el)
    paragraph._p.addnext(tbl_el)
    return table


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para


def has_trustworthiness_table(doc: Document) -> bool:
    for table in doc.tables:
        if table.rows and TRUST_HDR in table.rows[0].cells[0].text:
            return True
    return False


def _find_findings_paragraph(doc: Document) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip().startswith("4.") and "Findings" in p.text:
            return p
    msg = "Could not find '4. Findings' heading."
    raise ValueError(msg)


def _find_section_34_paragraph(doc: Document) -> Paragraph | None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("3.4") or ("3.4" in t and "trustworthiness" in t.lower()):
            return p
    return None


def _style_like_section(doc: Document, section_prefix: str, fallback: str) -> object:
    for p in doc.paragraphs:
        if p.text.strip().startswith(section_prefix):
            return p.style
    return doc.styles[fallback]


def _append_table_block(anchor: Paragraph) -> None:
    cap = insert_paragraph_after(anchor, TABLE_CAPTION)
    title_p = insert_paragraph_after(cap, TABLE_TITLE)
    insert_table_after(title_p, TABLE_ROWS)


def _convert_existing_section(doc: Document, section_34: Paragraph, findings: Paragraph) -> None:
    start = next(i for i, p in enumerate(doc.paragraphs) if p._p is section_34._p)
    end = next(i for i, p in enumerate(doc.paragraphs) if p._p is findings._p)

    intro_p = doc.paragraphs[start + 1]
    intro_p.text = INTRO
    intro_p.style = _style_like_section(doc, "3.3.", "Normal")

    to_remove = [doc.paragraphs[i]._element for i in range(start + 2, end)]
    for el in to_remove:
        el.getparent().remove(el)

    intro_p = next(p for p in doc.paragraphs if p.text == INTRO)
    _append_table_block(intro_p)


def _insert_new_section(doc: Document, findings: Paragraph) -> None:
    heading = findings.insert_paragraph_before("3.4.\tEnsuring trustworthiness")
    heading.style = _style_like_section(doc, "3.3.", "Heading 2")

    intro = insert_paragraph_after(heading, INTRO)
    intro.style = _style_like_section(doc, "3.3.", "Normal")
    _append_table_block(intro)


def main() -> None:
    doc = Document(str(DOCX_PATH))

    for p in doc.paragraphs:
        if "see Section 3.4.5" in p.text:
            p.text = p.text.replace("see Section 3.4.5", "see Table 3")
        if "see section 4)" in p.text.lower():
            p.text = p.text.replace("see section 4)", "see Table 3)")

    if has_trustworthiness_table(doc):
        print(f"Trustworthiness table already present in {DOCX_PATH}")
        _save(doc)
        return

    findings = _find_findings_paragraph(doc)
    section_34 = _find_section_34_paragraph(doc)

    if section_34 is not None:
        _convert_existing_section(doc, section_34, findings)
        action = "Converted existing §3.4 to table"
    else:
        _insert_new_section(doc, findings)
        action = "Inserted new §3.4 with table before Findings"

    _save(doc, action)


def _save(doc: Document, action: str = "Saved") -> None:
    try:
        doc.save(str(DOCX_PATH))
        print(f"{action} in {DOCX_PATH}")
    except PermissionError:
        doc.save(str(DOCX_FALLBACK))
        print(
            f"Could not overwrite {DOCX_PATH} (file may be open in Word). "
            f"Saved to {DOCX_FALLBACK}"
        )


if __name__ == "__main__":
    main()
