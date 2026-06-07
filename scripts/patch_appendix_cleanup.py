# -*- coding: utf-8 -*-
"""Phase 0 + bullets (tracked, author 'Claude').
- Remove the 'process of constructive grounded theory' appendix (pre-existing -> tracked delete)
  and the 'Supporting Quotes' appendix (heading/placeholder tracked-deleted; this-session quote
  insertions dropped outright) plus the empty spacer headings between them.
- Renumber: checklist 'Appendix D.' -> 'Appendix B.' (in-place edit of pending insertion);
  'Appendix E. What to include?' -> 'Appendix C.' (tracked rename).
- Convert the checklist items from '☐ ' paragraphs to a real bulleted list (numId 1)."""
from __future__ import annotations
import copy, shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
BACKUP = ROOT / "Thesis Draft - Daan Luttik - MBA.cleanup-backup.docx"
AUTHOR = "Claude"; DATE = "2026-06-07T18:00:00Z"
XS = "{http://www.w3.org/XML/1998/namespace}space"
BULLET_NUMID = "1"   # abstractNum 0, lvl0 '●'

def used_ids(doc):
    return {int(el.get(qn("w:id"), 0)) for el in doc.element.body.iter() if el.get(qn("w:id")) is not None}
def nid(u): n = max(u, default=0) + 1; u.add(n); return n
def atext(el): return "".join(t.text or "" for t in el.iter(qn("w:t")))

def find_h1(doc, marker):
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and marker in atext(p._element):
            return p._element
    return None

def del_run(r, used):
    d = etree.Element(qn("w:del")); d.set(qn("w:id"), str(nid(used)))
    d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), DATE)
    rc = copy.deepcopy(r)
    for t in rc.findall(qn("w:t")):
        dt = etree.Element(qn("w:delText")); dt.set(XS, "preserve"); dt.text = t.text or ""
        t.getparent().replace(t, dt)
    d.append(rc); return d

def del_para(p, used):
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.Element(qn("w:pPr")); p.insert(0, pPr)
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None: rPr = etree.SubElement(pPr, qn("w:rPr"))
    dm = etree.Element(qn("w:del")); dm.set(qn("w:id"), str(nid(used)))
    dm.set(qn("w:author"), AUTHOR); dm.set(qn("w:date"), DATE); rPr.insert(0, dm)
    for r in p.findall(qn("w:r")):
        p.replace(r, del_run(r, used))

def remove_para_tracked(p, used):
    for ins in list(p.findall(qn("w:ins"))):
        if ins.get(qn("w:author")) == AUTHOR:
            p.remove(ins)
    if not (p.findall(qn("w:r")) or p.findall(qn("w:del")) or p.findall(qn("w:hyperlink"))):
        p.getparent().remove(p)              # empty / all-insertion -> drop outright
    else:
        del_para(p, used)                    # pre-existing content -> tracked delete

def replace_keepfmt(p, old, new, used):
    runs = p.findall(qn("w:r")); segs, total = [], ""
    for r in runs:
        t = r.find(qn("w:t")); txt = t.text if (t is not None and t.text) else ""
        segs.append((r, txt)); total += txt
    s = total.find(old)
    if s < 0: return False
    e = s + len(old)
    def mk(rPr, x):
        nr = etree.Element(qn("w:r"))
        if rPr is not None: nr.append(copy.deepcopy(rPr))
        tt = etree.SubElement(nr, qn("w:t")); tt.set(XS, "preserve"); tt.text = x; return nr
    def insr(rPr, x):
        ins = etree.Element(qn("w:ins")); ins.set(qn("w:id"), str(nid(used)))
        ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), DATE)
        ins.append(mk(rPr, x)); return ins
    def delr(rPr, x):
        d = etree.Element(qn("w:del")); d.set(qn("w:id"), str(nid(used)))
        d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), DATE)
        rr = etree.SubElement(d, qn("w:r"))
        if rPr is not None: rr.insert(0, copy.deepcopy(rPr))
        dt = etree.SubElement(rr, qn("w:delText")); dt.set(XS, "preserve"); dt.text = x; return d
    pos = 0; done = False
    for r, txt in segs:
        rs, re_ = pos, pos + len(txt); pos = re_
        if not txt or re_ <= s or rs >= e: continue
        rPr = r.find(qn("w:rPr")); ls = max(s, rs) - rs; le = min(e, re_) - rs
        before, mid, after = txt[:ls], txt[ls:le], txt[le:]
        parent = r.getparent(); idx = list(parent).index(r); parts = []
        if before: parts.append(mk(rPr, before))
        if not done: parts.append(insr(rPr, new)); done = True
        if mid: parts.append(delr(rPr, mid))
        if after: parts.append(mk(rPr, after))
        for j, prt in enumerate(parts): parent.insert(idx + j, prt)
        parent.remove(r)
    return done

def add_bullet(p):
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.Element(qn("w:pPr")); p.insert(0, pPr)
    numPr = etree.Element(qn("w:numPr"))
    ilvl = etree.SubElement(numPr, qn("w:ilvl")); ilvl.set(qn("w:val"), "0")
    numId = etree.SubElement(numPr, qn("w:numId")); numId.set(qn("w:val"), BULLET_NUMID)
    ps = pPr.find(qn("w:pStyle"))
    pPr.insert(list(pPr).index(ps) + 1 if ps is not None else 0, numPr)


def main():
    doc = Document(str(DOCX)); used = used_ids(doc)
    cgt = find_h1(doc, "constructive grounded theory")
    quotes = find_h1(doc, "Supporting Quotes")
    checklist = find_h1(doc, "Putting the model to work")
    whatinc = find_h1(doc, "What to include")
    assert all([cgt, quotes, checklist, whatinc]), "an appendix heading was not found"

    # checklist item elements (Claude-inserted Normal paras starting with the box glyph)
    items = []
    seen = False
    for p in doc.paragraphs:
        if p._element is checklist: seen = True; continue
        if p._element is whatinc: break
        if seen and atext(p._element).strip().startswith("☐"):
            items.append(p._element)

    # --- retitle checklist heading D -> B (edit pending insertion text in place) ---
    for t in checklist.iter(qn("w:t")):
        if t.text and "Appendix D." in t.text:
            t.text = t.text.replace("Appendix D.", "Appendix B."); break
    print("retitled checklist heading -> Appendix B")

    # --- rename What to include E -> C (tracked) ---
    assert replace_keepfmt(whatinc, "Appendix E", "Appendix C", used), "What-to-include rename failed"
    print("renamed What to include -> Appendix C")

    # --- bullets: strip glyph + apply numbering ---
    for el in items:
        for t in el.iter(qn("w:t")):
            if t.text and "☐" in t.text:
                t.text = t.text.replace("☐ ", "").replace("☐", ""); break
        add_bullet(el)
    print(f"bulleted {len(items)} checklist items")

    # --- removal range: empty-spacer-before-CGT .. up to (excluding) checklist heading ---
    body = doc.element.body; kids = list(body)
    i_cgt = kids.index(cgt); i_check = kids.index(checklist)
    start = i_cgt
    prev = kids[i_cgt - 1]
    if prev.tag == qn("w:p") and not atext(prev).strip():
        start = i_cgt - 1
    to_remove = kids[start:i_check]
    assert all(el.tag == qn("w:p") for el in to_remove), "unexpected table in removal range"
    for el in to_remove:
        remove_para_tracked(el, used)
    print(f"removed {len(to_remove)} paragraphs (CGT + Supporting Quotes + spacers)")

    shutil.copy(DOCX, BACKUP); doc.save(str(DOCX))
    print(f"\nSaved: {DOCX}\nBackup: {BACKUP}")


if __name__ == "__main__":
    main()
