# -*- coding: utf-8 -*-
"""Follow-up: reconcile Ellström/Leone years so refs match in-text citations.
 - §2.3 Ellström in-text (2021->2022): tracked del+ins (plain run).
 - Ellström ref entry (2021->2022) and Leone ref entry (2020->2021): the entries are
   the user's own pending tracked insertions, so the year typo is corrected in place
   (keeps each as one clean insertion). Reported explicitly for review.
"""
import copy, shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
BACKUP = ROOT / "Thesis Draft - Daan Luttik - MBA.yearfix-backup.docx"
AUTHOR, DATE = "Claude", "2026-06-07T00:00:00Z"
XS = "{http://www.w3.org/XML/1998/namespace}space"

def ftext(p): return "".join(t.text or "" for t in p.iter(qn("w:t")))
def used_ids(doc): return {int(e.get(qn("w:id"),0)) for e in doc.element.body.iter() if e.get(qn("w:id")) is not None}

def tracked_replace_plain(p, old, new, used):
    """tracked del+ins on a direct-child plain run containing `old`."""
    for r in p.findall(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is None or not t.text or old not in t.text:
            continue
        before, _, after = t.text.partition(old)
        parent = r.getparent(); i = list(parent).index(r); n=lambda: (max(used)+1)
        parts = []
        rpr = r.find(qn("w:rPr"))
        def mk(text):
            rr=etree.Element(qn("w:r"))
            if rpr is not None: rr.append(copy.deepcopy(rpr))
            tt=etree.SubElement(rr,qn("w:t")); tt.set(XS,"preserve"); tt.text=text; return rr
        if before: parts.append(mk(before))
        nid=max(used)+1; used.add(nid)
        d=etree.Element(qn("w:del")); d.set(qn("w:id"),str(nid)); d.set(qn("w:author"),AUTHOR); d.set(qn("w:date"),DATE)
        rd=etree.SubElement(d,qn("w:r"))
        if rpr is not None: rd.append(copy.deepcopy(rpr))
        dt=etree.SubElement(rd,qn("w:delText")); dt.set(XS,"preserve"); dt.text=old; parts.append(d)
        nid=max(used)+1; used.add(nid)
        ins=etree.Element(qn("w:ins")); ins.set(qn("w:id"),str(nid)); ins.set(qn("w:author"),AUTHOR); ins.set(qn("w:date"),DATE)
        ins.append(mk(new)); parts.append(ins)
        if after: parts.append(mk(after))
        for j,part in enumerate(parts): parent.insert(i+j,part)
        parent.remove(r); return True
    return False

def inplace_year(doc, para_marker, old, new):
    """Correct a year inside an existing (user-inserted) run, in place."""
    for p in doc.paragraphs:
        if para_marker in ftext(p._element):
            for t in p._element.iter(qn("w:t")):
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new, 1)
                    return True
    return False

def main():
    shutil.copy(DOCX, BACKUP)
    doc = Document(str(DOCX)); used = used_ids(doc)

    # 1. §2.3 Ellström in-text year (tracked)
    done = False
    for p in doc.paragraphs:
        if "identify sensing, seizing" in ftext(p._element):
            done = tracked_replace_plain(p._element, "(2021) identify", "(2022) identify", used); break
    print(f"1. §2.3 Ellström in-text 2021->2022 (tracked): {'ok' if done else 'FAIL'}")

    # 2. Ellström ref entry year (in place)
    ok2 = inplace_year(doc, "Ellström, D., Holtström", "Josefsson, C. (2021)", "Josefsson, C. (2022)")
    print(f"2. Ellström ref 2021->2022 (in place): {'ok' if ok2 else 'FAIL'}")

    # 3. Leone ref entry year (in place; avoid the DOI 2020)
    ok3 = inplace_year(doc, "Leone, D., Schiavone", "(2020). How does", "(2021). How does")
    print(f"3. Leone ref 2020->2021 (in place): {'ok' if ok3 else 'FAIL'}")

    doc.save(str(DOCX))
    print(f"Saved. Backup: {BACKUP}")

if __name__ == "__main__":
    main()
