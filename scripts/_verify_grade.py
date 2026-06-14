# -*- coding: utf-8 -*-
"""Simulate accept-all / reject-all on the patched doc and print key items."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
XMLSPACE="{http://www.w3.org/XML/1998/namespace}space"
DOCX=Path('Thesis Draft - Daan Luttik - MBA - copy from 2026-06-14.docx')

def ftext(pe): return "".join(t.text or "" for t in pe.iter(qn("w:t")))

def accept_all(doc):
    body=doc.element.body
    for d in list(body.iter(qn("w:del"))): d.getparent().remove(d)
    for ins in list(body.iter(qn("w:ins"))):
        par=ins.getparent(); i=list(par).index(ins)
        for c in list(ins): par.insert(i,c); i+=1
        par.remove(ins)

def reject_all(doc):
    body=doc.element.body
    for ins in list(body.iter(qn("w:ins"))): ins.getparent().remove(ins)
    for d in list(body.iter(qn("w:del"))):
        par=d.getparent(); i=list(par).index(d)
        for c in list(d):
            for dt in c.findall(qn("w:delText")):
                t=etree.Element(qn("w:t")); t.set(XMLSPACE,"preserve"); t.text=dt.text; c.replace(dt,t)
            par.insert(i,c); i+=1
        par.remove(d)

def ref_order(doc):
    out=[]
    for p in doc.paragraphs:
        pe=p._element; tx=ftext(pe).strip()
        pPr=pe.find(qn("w:pPr")); ind=pPr.find(qn("w:ind")) if pPr is not None else None
        if ind is not None and ind.get(qn("w:hanging"))=="720" and tx[:1].isalpha():
            out.append(tx[:34])
    return out

def show(doc,label):
    print(f"\n===================== {label} =====================")
    order=ref_order(doc)
    # regions
    def region(names):
        return [e for e in order if any(e.startswith(n) for n in names)]
    print("A-head :", order[:3])
    print("C-zone :", region(["Charmaz","ChatGPT","Claude","Cottier","Burkhardt"]))
    print("K-zone :", region(["Kim","Kohavi","Krizhevsky"]))
    print("P-zone :", region(["Parker","Patton","Polsa","Radford"]))
    for p in doc.paragraphs:
        tx=ftext(p._element).strip()
        if tx.startswith("Blumer"): print("Blumer :", tx)
        if tx.startswith("Abou Elgheit"): print("AbouElg:", tx)
        if tx.startswith("Charmaz, K. (2014)"): print("Charmaz:", tx)
    for p in doc.paragraphs:
        tx=ftext(p._element)
        if "data gathering" in tx and "theoretical sampling" in tx and "saturation" in tx:
            print("Meth3.3:", tx.strip()[:230])
        if "constructivist grounded theory, a form of qualitative" in tx:
            print("Meth3.1:", tx.strip()[:120])
    ny=sum(1 for h in doc.element.body.iter(qn("w:highlight")) if h.get(qn("w:val"))=="yellow")
    print("yellow highlights:", ny)

a=Document(str(DOCX)); accept_all(a); show(a,"ACCEPT-ALL (final result)")
r=Document(str(DOCX)); reject_all(r); show(r,"REJECT-ALL (should = original)")
