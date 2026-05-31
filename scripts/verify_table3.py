from docx import Document
from pathlib import Path

d = Document(Path(__file__).resolve().parents[1] / "Thesis Draft - Daan Luttik - MBA.docx")
lines = [f"tables={len(d.tables)}"]
for i, p in enumerate(d.paragraphs):
    if 110 <= i <= 140 or "3.4" in p.text or "Table 3" in p.text:
        lines.append(f"{i}|{p.style.name}|{p.text[:90]}")
if d.tables:
    t = d.tables[-1]
    lines.append(f"last_table_rows={len(t.rows)}")
    lines.append(f"hdr={t.rows[0].cells[0].text}|{t.rows[0].cells[1].text}")
    lines.append(f"r1crit={t.rows[1].cells[0].text[:60]}...")
Path(__file__).resolve().parents[1] / "verify-table3.txt"
Path(__file__).resolve().parents[1].joinpath("verify-table3.txt").write_text(
    "\n".join(lines), encoding="utf-8"
)
