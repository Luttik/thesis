"""
Accept all tracked changes in section 4.2, fix Heading 4 → bold inline,
and replace em-dashes with appropriate punctuation.

Steps (in order):
1. Accept tracked changes in 4.2 range
2. Fix Heading 4 paragraphs → bold inline text at start of following paragraph
3. Replace em-dashes with context-appropriate punctuation
"""
from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT      = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
OUT_PATH  = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"

EM_DASH = "—"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def full_text(p_elem: etree._Element) -> str:
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def para_is_fully_deleted(p_elem: etree._Element) -> bool:
    """True if the paragraph mark is flagged as deleted (whole para was deleted)."""
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return False
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        return False
    return rPr.find(qn("w:del")) is not None


def accept_ins(elem: etree._Element) -> None:
    """Unwrap a <w:ins>: move its children to its parent, remove the wrapper."""
    parent = elem.getparent()
    if parent is None:
        return
    idx = list(parent).index(elem)
    for i, child in enumerate(list(elem)):
        parent.insert(idx + i, child)
    parent.remove(elem)


def accept_tracked_changes(body: etree._Element,
                            start_elem: etree._Element,
                            end_elem: etree._Element) -> int:
    """
    Accept all tracked changes between start_elem and end_elem (exclusive).
    Returns count of changes accepted.
    """
    # Collect direct body children in the range
    children = list(body)
    try:
        s_idx = children.index(start_elem)
        e_idx = children.index(end_elem)
    except ValueError:
        return 0

    in_range = children[s_idx + 1 : e_idx]  # paragraphs between start and end

    count = 0

    # First pass: remove fully-deleted paragraphs
    for p_elem in list(in_range):
        if p_elem.tag == qn("w:p") and para_is_fully_deleted(p_elem):
            body.remove(p_elem)
            in_range.remove(p_elem)
            count += 1

    # Second pass: process ins/del within remaining paragraphs
    for p_elem in in_range:
        if p_elem.tag != qn("w:p"):
            continue

        # Remove <w:del> in pPr/rPr (paragraph mark deletion marker, if any remain)
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is not None:
            rPr = pPr.find(qn("w:rPr"))
            if rPr is not None:
                for d in rPr.findall(qn("w:del")):
                    rPr.remove(d)
                    count += 1

        # Process all <w:ins> and <w:del> in document order (iterative, deepest first)
        changed = True
        while changed:
            changed = False
            # Handle <w:del> — remove entirely
            for d in p_elem.findall(f".//{qn('w:del')}"):
                parent = d.getparent()
                if parent is not None:
                    parent.remove(d)
                    count += 1
                    changed = True
                    break
            if changed:
                continue
            # Handle <w:ins> — unwrap
            for ins in p_elem.findall(f".//{qn('w:ins')}"):
                accept_ins(ins)
                count += 1
                changed = True
                break

        # Convert any remaining <w:delText> to <w:t> (safety net)
        for dt in p_elem.findall(f".//{qn('w:delText')}"):
            dt.tag = qn("w:t")
            count += 1

    return count


# ---------------------------------------------------------------------------
# Step 2: Fix Heading 4 → bold inline
# ---------------------------------------------------------------------------

def fix_heading4(doc: Document, body: etree._Element,
                 start_elem: etree._Element, end_elem: etree._Element) -> int:
    """
    For each Heading 4 paragraph in range, make the first words of the
    following paragraph bold (matching the heading text), then delete the heading.
    Returns count of headings fixed.
    """
    children = list(body)
    s_idx = children.index(start_elem)
    e_idx = children.index(end_elem)
    in_range = children[s_idx + 1 : e_idx]

    count = 0
    to_remove = []

    for i, p_elem in enumerate(in_range):
        if p_elem.tag != qn("w:p"):
            continue
        # Check style
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None or pStyle.get(qn("w:val"), "") not in ("Heading4", "4"):
            continue

        heading_text = "".join(t.text or "" for t in p_elem.iter(qn("w:t")))
        if not heading_text.strip():
            to_remove.append(p_elem)
            continue

        # Find the next Normal paragraph in the range
        next_p = None
        for j in range(i + 1, len(in_range)):
            cand = in_range[j]
            if cand.tag == qn("w:p"):
                next_p = cand
                break

        if next_p is None:
            to_remove.append(p_elem)
            continue

        # Find the first run(s) in next_p whose text starts with heading_text
        runs = next_p.findall(qn("w:r"))
        if not runs:
            to_remove.append(p_elem)
            continue

        # Collect leading runs that together form the heading text
        accumulated = ""
        bold_runs = []
        for r in runs:
            r_text = "".join(t.text or "" for t in r.findall(qn("w:t")))
            accumulated += r_text
            bold_runs.append(r)
            if accumulated.strip().startswith(heading_text.strip()):
                break

        # Make those runs bold
        for r in bold_runs:
            rPr = r.find(qn("w:rPr"))
            if rPr is None:
                rPr = etree.Element(qn("w:rPr"))
                r.insert(0, rPr)
            # Add <w:b/> if not already present
            if rPr.find(qn("w:b")) is None:
                b = etree.SubElement(rPr, qn("w:b"))

        to_remove.append(p_elem)
        count += 1
        print(f"  Made '{heading_text}' bold in following paragraph.")

    for p_elem in to_remove:
        if p_elem.getparent() is not None:
            p_elem.getparent().remove(p_elem)

    return count


# ---------------------------------------------------------------------------
# Step 3: Replace em-dashes
# ---------------------------------------------------------------------------

EM_DASH_REPLACEMENTS: list[tuple[str, str]] = [
    ("reinforcing — there is no",
     "reinforcing; there is no"),
    ("in systems — to see how data, tools, and processes interact — was",
     "in systems (seeing how data, tools, and processes interact) was"),
    ("across the organization — through training sessions",
     "across the organization, through training sessions"),
    ("without fixing it — a pragmatic",
     "without fixing it, a pragmatic"),
    ("systems thinking — the ability to understand",
     "systems thinking, understood as the ability to understand"),
    ("deep process knowledge — a feel for system logic — is",
     "deep process knowledge (a feel for system logic) is"),
    ("and leadership — the organizational structures",
     "and leadership, covering the organizational structures"),
    ("Providing clarity — naming",
     "Providing clarity, that is, naming"),
    ("in her own organization — consistently pushing",
     "in her own organization, consistently pushing"),
    (
        "people along — change management at the team and organizational level "
        "— was",
        "people along, understood as change management at the team and "
        "organizational level, was",
    ),
    ("organizational friction — the IT/legal/compliance",
     "organizational friction, covering the IT/legal/compliance"),
    # Quote from Interviewee 6 (para 214): keep interviewee voice but fix punctuation
    ("the beginning — “You cannot",
     "the beginning: “You cannot"),
    ("these sections — educating",
     "these sections: educating"),
]


def replace_em_dashes(body: etree._Element,
                      start_elem: etree._Element,
                      end_elem: etree._Element) -> int:
    """
    Apply all em-dash replacements to <w:t> text in the 4.2 range.
    Operates on paragraph-level concatenated text to handle multi-run spans,
    then writes back by splitting runs as needed.
    """
    children = list(body)
    s_idx = children.index(start_elem)
    e_idx = children.index(end_elem)

    count = 0
    for p_elem in children[s_idx + 1 : e_idx]:
        if p_elem.tag != qn("w:p"):
            continue

        # Collect all w:t elements
        t_elems = list(p_elem.iter(qn("w:t")))
        if not t_elems:
            continue

        # Build concatenated text
        full = "".join(t.text or "" for t in t_elems)
        if EM_DASH not in full:
            continue

        # Apply replacements
        new_full = full
        for old, new in EM_DASH_REPLACEMENTS:
            if old in new_full:
                new_full = new_full.replace(old, new)
                count += 1

        if new_full == full:
            # No replacement matched — do a generic fallback for any remaining em-dash
            if EM_DASH in new_full:
                new_full = new_full.replace(f" {EM_DASH} ", ", ")
                if EM_DASH in new_full:
                    new_full = new_full.replace(EM_DASH, ",")
                count += 1

        if new_full == full:
            continue

        # Write back: put all text into the first w:t, clear the rest
        t_elems[0].text = new_full
        for t in t_elems[1:]:
            t.text = ""

    return count


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    backup = DOCX_PATH.with_suffix(".bak.docx")
    shutil.copy(DOCX_PATH, backup)
    print(f"Backup: {backup}")

    doc = Document(str(DOCX_PATH))
    body = doc.element.body

    # Find the 4.2 and 4.3 heading elements
    start_elem = end_elem = None
    for p in doc.paragraphs:
        txt = "".join(t.text or "" for t in p._element.iter(qn("w:t")))
        if start_elem is None and "4.2" in txt and "Creating" in txt:
            start_elem = p._element
        elif start_elem is not None and "4.3" in txt and "Applying" in txt:
            end_elem = p._element
            break

    assert start_elem is not None, "4.2 heading not found"
    assert end_elem is not None, "4.3 heading not found"

    # Step 1: Accept tracked changes
    n = accept_tracked_changes(body, start_elem, end_elem)
    print(f"  Accepted {n} tracked change elements.")

    # Step 2: Fix Heading 4 → bold inline
    n = fix_heading4(doc, body, start_elem, end_elem)
    print(f"  Fixed {n} Heading 4 paragraph(s).")

    # Step 3: Replace em-dashes
    n = replace_em_dashes(body, start_elem, end_elem)
    print(f"  Replaced {n} em-dash occurrence(s).")

    # Verification
    remaining_ins  = sum(1 for p in doc.paragraphs
                         for _ in p._element.findall(f".//{qn('w:ins')}"))
    remaining_del  = sum(1 for p in doc.paragraphs
                         for _ in p._element.findall(f".//{qn('w:del')}"))
    remaining_h4   = sum(1 for p in doc.paragraphs
                         if p.style and "Heading4" in p.style.name.replace(" ", ""))
    remaining_dash = sum(1 for p in doc.paragraphs
                         for t in p._element.iter(qn("w:t"))
                         if EM_DASH in (t.text or ""))

    print(f"\nVerification (whole doc):")
    print(f"  Remaining <w:ins>:  {remaining_ins}")
    print(f"  Remaining <w:del>:  {remaining_del}")
    print(f"  Remaining Heading4: {remaining_h4}")
    print(f"  Remaining em-dashes in 4.2 range: {remaining_dash} "
          f"(note: may include other sections)")

    doc.save(str(OUT_PATH))
    print(f"\nSaved: {OUT_PATH}")


if __name__ == "__main__":
    main()
