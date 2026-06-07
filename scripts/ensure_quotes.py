"""Check and fix quotation mark formatting in Word documents.

Two fixes are applied:

1. Smart quotes: convert straight ASCII " and ' to typographic “ ” and ‘ ’.
2. Block quotes: remove redundant outer quote marks from Quote-style paragraphs.

Run ``check`` to report issues and ``fix`` to apply both fixes in one pass
(smart quotes first, then block-quote cleanup).
"""

from __future__ import annotations

import html
import re
import zipfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import typer
from docx import Document
from docx.document import Document as DocumentType
from docx.table import _Cell
from docx.text.paragraph import Paragraph

app = typer.Typer(add_completion=False, no_args_is_help=True)

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOCX = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"

LEFT_DOUBLE = "\u201c"
RIGHT_DOUBLE = "\u201d"
LEFT_SINGLE = "\u2018"
RIGHT_SINGLE = "\u2019"

STRAIGHT_DOUBLE = '"'
STRAIGHT_SINGLE = "'"
OPEN_DOUBLE = {STRAIGHT_DOUBLE, LEFT_DOUBLE}
CLOSE_DOUBLE = {STRAIGHT_DOUBLE, RIGHT_DOUBLE}
OPEN_SINGLE = {STRAIGHT_SINGLE, LEFT_SINGLE}
CLOSE_SINGLE = {STRAIGHT_SINGLE, RIGHT_SINGLE}

QUOTE_STYLES = {"Quote", "Block Text"}

APOSTROPHE_RE = re.compile(r"(?<=\w)'(?=\w)")
CONTRACTION_RE = re.compile(
    r"(?<=\w)'(?=(?:s|t|re|ve|ll|d|m|M|em|til|round)\b)",
    re.IGNORECASE,
)

PARAGRAPH_RE = re.compile(r"(<w:p\b.*?</w:p>)", re.DOTALL)
TEXT_NODE_RE = re.compile(
    r"(<w:(?:t|delText)\b(?:\s[^>]*)?>)(.*?)(</w:(?:t|delText)>)",
    re.DOTALL,
)
QUOTE_STYLE_RE = re.compile(r"""<w:pStyle w:val=["'](?P<style>Quote|Block Text)["']""")

QUOTE_NORMALIZE = str.maketrans(
    {
        "\u201c": '"',
        "\u201d": '"',
        "\u2018": "'",
        "\u2019": "'",
    }
)

XML_PARTS = (
    "document.xml",
    "footnotes.xml",
    "endnotes.xml",
    "comments.xml",
)


@dataclass(frozen=True)
class StraightQuoteIssue:
    location: str
    paragraph_index: int
    snippet: str
    count: int


@dataclass(frozen=True)
class WrappedBlockQuoteIssue:
    location: str
    paragraph_index: int
    style: str
    snippet: str


def smarten_quotes(text: str) -> str:
    """Convert straight ASCII quotes to typographic opening/closing quotes."""
    if STRAIGHT_DOUBLE not in text and STRAIGHT_SINGLE not in text:
        return text

    converted = CONTRACTION_RE.sub(RIGHT_SINGLE, text)
    converted = APOSTROPHE_RE.sub(RIGHT_SINGLE, converted)

    result: list[str] = []
    open_double = True
    open_single = True

    for char in converted:
        if char == STRAIGHT_DOUBLE:
            result.append(LEFT_DOUBLE if open_double else RIGHT_DOUBLE)
            open_double = not open_double
            continue
        if char == STRAIGHT_SINGLE:
            result.append(LEFT_SINGLE if open_single else RIGHT_SINGLE)
            open_single = not open_single
            continue
        result.append(char)

    return "".join(result)


def find_straight_quotes(text: str) -> list[str]:
    return [char for char in text if char in {STRAIGHT_DOUBLE, STRAIGHT_SINGLE}]


def is_quote_paragraph(paragraph: Paragraph) -> bool:
    style = paragraph.style.name if paragraph.style else ""
    return style in QUOTE_STYLES


def is_quote_paragraph_xml(paragraph_xml: str) -> bool:
    return QUOTE_STYLE_RE.search(paragraph_xml) is not None


def has_wrapping_quotes(text: str) -> bool:
    stripped = text.strip()
    if len(stripped) < 2:
        return False
    if stripped[0] in OPEN_DOUBLE and stripped[-1] in CLOSE_DOUBLE:
        return True
    return stripped[0] in OPEN_SINGLE and stripped[-1] in CLOSE_SINGLE


def strip_wrapping_quotes(text: str) -> tuple[str, bool]:
    if not has_wrapping_quotes(text):
        return text, False

    leading = text[: len(text) - len(text.lstrip())]
    trailing = text[len(text.rstrip()) :]
    stripped = text.strip()
    return leading + stripped[1:-1] + trailing, True


def decode_xml_text(text: str) -> str:
    return html.unescape(text)


def encode_xml_text(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def rewrite_paragraph_text(paragraph_xml: str, converted: str, matches: list[re.Match[str]]) -> str:
    rebuilt: list[str] = []
    last_end = 0
    position = 0
    for match in matches:
        decoded_part = decode_xml_text(match.group(2))
        rebuilt.append(paragraph_xml[last_end : match.start(2)])
        part_len = len(decoded_part)
        rebuilt.append(encode_xml_text(converted[position : position + part_len]))
        position += part_len
        last_end = match.end(2)
    rebuilt.append(paragraph_xml[last_end:])
    return "".join(rebuilt)


def fix_paragraph_block(paragraph_xml: str) -> tuple[str, bool, bool]:
    matches = list(TEXT_NODE_RE.finditer(paragraph_xml))
    if not matches:
        return paragraph_xml, False, False

    original = "".join(decode_xml_text(match.group(2)) for match in matches)
    converted = smarten_quotes(original)
    smart_changed = converted != original

    block_changed = False
    if is_quote_paragraph_xml(paragraph_xml):
        converted, block_changed = strip_wrapping_quotes(converted)

    if not smart_changed and not block_changed:
        return paragraph_xml, False, False

    return rewrite_paragraph_text(paragraph_xml, converted, matches), smart_changed, block_changed


def fix_xml_text(xml: str) -> tuple[str, int, int]:
    smart_paragraphs = 0
    block_paragraphs = 0

    def replace_paragraph(match: re.Match[str]) -> str:
        nonlocal smart_paragraphs, block_paragraphs
        updated, smart_changed, block_changed = fix_paragraph_block(match.group(1))
        if smart_changed:
            smart_paragraphs += 1
        if block_changed:
            block_paragraphs += 1
        return updated

    return PARAGRAPH_RE.sub(replace_paragraph, xml), smart_paragraphs, block_paragraphs


def should_fix_xml_part(filename: str) -> bool:
    if not filename.startswith("word/") or not filename.endswith(".xml"):
        return False
    name = Path(filename).name
    if name in XML_PARTS:
        return True
    return name.startswith(("header", "footer"))


def fix_docx_file(source: Path, destination: Path) -> tuple[int, int]:
    smart_paragraphs = 0
    block_paragraphs = 0
    buffer = BytesIO()

    with zipfile.ZipFile(source, "r") as source_zip:
        with zipfile.ZipFile(buffer, "w") as destination_zip:
            for item in source_zip.infolist():
                data = source_zip.read(item.filename)
                if should_fix_xml_part(item.filename):
                    xml = data.decode("utf-8")
                    xml, smart_count, block_count = fix_xml_text(xml)
                    smart_paragraphs += smart_count
                    block_paragraphs += block_count
                    data = xml.encode("utf-8")
                destination_zip.writestr(item, data)

    destination.write_bytes(buffer.getvalue())
    return smart_paragraphs, block_paragraphs


def _iter_paragraphs(parent: DocumentType | _Cell) -> list[tuple[str, int, Paragraph]]:
    paragraphs: list[tuple[str, int, Paragraph]] = []

    for index, paragraph in enumerate(parent.paragraphs):
        paragraphs.append(("body", index, paragraph))

    for table_index, table in enumerate(parent.tables):
        seen_cells: set[int] = set()
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                location = f"table {table_index + 1}, row {row_index + 1}, cell {cell_index + 1}"
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    paragraphs.append((location, paragraph_index, paragraph))

    if isinstance(parent, DocumentType):
        for section_index, section in enumerate(parent.sections):
            for kind, header in (
                ("header", section.header),
                ("first-page header", section.first_page_header),
                ("even-page header", section.even_page_header),
                ("footer", section.footer),
                ("first-page footer", section.first_page_footer),
                ("even-page footer", section.even_page_footer),
            ):
                if header is None:
                    continue
                location = f"section {section_index + 1} {kind}"
                for paragraph_index, paragraph in enumerate(header.paragraphs):
                    paragraphs.append((location, paragraph_index, paragraph))

    return paragraphs


def _snippet(text: str, position: int, radius: int = 45) -> str:
    start = max(0, position - radius)
    end = min(len(text), position + radius)
    snippet = text[start:end].replace("\n", " ")
    if start > 0:
        snippet = f"...{snippet}"
    if end < len(text):
        snippet = f"{snippet}..."
    return snippet


def scan_straight_quotes(document: DocumentType) -> list[StraightQuoteIssue]:
    issues: list[StraightQuoteIssue] = []
    for location, paragraph_index, paragraph in _iter_paragraphs(document):
        text = paragraph.text
        straight = find_straight_quotes(text)
        if not straight:
            continue
        first = text.find(straight[0])
        issues.append(
            StraightQuoteIssue(
                location=location,
                paragraph_index=paragraph_index,
                snippet=_snippet(text, first),
                count=len(straight),
            )
        )
    return issues


def scan_wrapped_block_quotes(document: DocumentType) -> list[WrappedBlockQuoteIssue]:
    issues: list[WrappedBlockQuoteIssue] = []
    for location, paragraph_index, paragraph in _iter_paragraphs(document):
        if not is_quote_paragraph(paragraph):
            continue
        text = paragraph.text
        if not has_wrapping_quotes(text):
            continue
        issues.append(
            WrappedBlockQuoteIssue(
                location=location,
                paragraph_index=paragraph_index,
                style=paragraph.style.name,
                snippet=text.strip()[:90],
            )
        )
    return issues


def count_fixable_paragraphs(document: DocumentType) -> tuple[int, int]:
    smart_fixable = 0
    block_fixable = 0
    for _, _, paragraph in _iter_paragraphs(document):
        text = paragraph.text
        if smarten_quotes(text) != text:
            smart_fixable += 1
        if is_quote_paragraph(paragraph) and strip_wrapping_quotes(text)[1]:
            block_fixable += 1
    return smart_fixable, block_fixable


def normalized_document_text(path: Path) -> str:
    document = Document(str(path))
    chunks: list[str] = []
    for _, _, paragraph in _iter_paragraphs(document):
        text = paragraph.text.translate(QUOTE_NORMALIZE)
        if is_quote_paragraph(paragraph):
            text, _ = strip_wrapping_quotes(text)
        chunks.append(text)
    return "\n".join(chunks)


def verify_quote_only_change(before_path: Path, after_path: Path) -> tuple[bool, str]:
    if normalized_document_text(before_path) == normalized_document_text(after_path):
        return True, "PASS — only quotation mark formatting changed"
    return False, "FAIL — non-quote content changed"


@app.command()
def check(
    input_path: Path = typer.Argument(
        None,
        help="Word document to scan. Defaults to the main thesis draft.",
    ),
) -> None:
    """Report straight ASCII quotes and wrapped block-quote paragraphs."""
    path = input_path or DEFAULT_DOCX
    if not path.exists():
        typer.echo(f"File not found: {path}", err=True)
        raise typer.Exit(code=1)

    document = Document(str(path))
    straight_issues = scan_straight_quotes(document)
    block_issues = scan_wrapped_block_quotes(document)

    if not straight_issues and not block_issues:
        typer.echo(f"OK: quotation formatting looks good in {path.name}")
        raise typer.Exit(code=0)

    if straight_issues:
        total = sum(issue.count for issue in straight_issues)
        typer.echo(
            f"Straight quotes: {total} character(s) in {len(straight_issues)} paragraph(s)"
        )
        for issue in straight_issues:
            typer.echo(
                f"- {issue.location}, paragraph {issue.paragraph_index + 1} "
                f"({issue.count}): {issue.snippet}"
            )

    if block_issues:
        if straight_issues:
            typer.echo("")
        typer.echo(f"Wrapped block quotes: {len(block_issues)} paragraph(s)")
        for issue in block_issues:
            typer.echo(
                f"- {issue.location}, paragraph {issue.paragraph_index + 1} "
                f"({issue.style}): {issue.snippet}..."
            )

    raise typer.Exit(code=1)


@app.command()
def fix(
    input_path: Path = typer.Argument(
        None,
        help="Word document to fix. Defaults to the main thesis draft.",
    ),
    output_path: Path | None = typer.Option(
        None,
        "--output",
        "-o",
        help="Write the fixed document here. Defaults to overwriting the input file.",
    ),
    dry_run: bool = typer.Option(
        False,
        "--dry-run",
        help="Show what would change without writing a file.",
    ),
) -> None:
    """Apply smart quotes and block-quote cleanup in one pass."""
    path = input_path or DEFAULT_DOCX
    if not path.exists():
        typer.echo(f"File not found: {path}", err=True)
        raise typer.Exit(code=1)

    document = Document(str(path))
    straight_issues = scan_straight_quotes(document)
    block_issues = scan_wrapped_block_quotes(document)

    if not straight_issues and not block_issues:
        typer.echo(f"OK: quotation formatting looks good in {path.name}")
        raise typer.Exit(code=0)

    smart_fixable, block_fixable = count_fixable_paragraphs(document)
    straight_total = sum(issue.count for issue in straight_issues)

    if dry_run:
        typer.echo(
            f"Would smarten {smart_fixable} paragraph(s) "
            f"({straight_total} straight quote characters)"
        )
        typer.echo(f"Would unwrap {block_fixable} block-quote paragraph(s)")
        raise typer.Exit(code=0)

    destination = output_path or path
    backup_path = path.with_suffix(path.suffix + ".bak")
    if destination == path and not backup_path.exists():
        backup_path.write_bytes(path.read_bytes())

    smart_fixed, block_fixed = fix_docx_file(path, destination)
    fixed_document = Document(str(destination))
    straight_after = scan_straight_quotes(fixed_document)
    block_after = scan_wrapped_block_quotes(fixed_document)

    typer.echo(
        f"Smart quotes: updated {smart_fixed} paragraph(s); "
        f"{straight_total} straight quote characters converted."
    )
    typer.echo(f"Block quotes: unwrapped {block_fixed} paragraph(s).")

    if straight_after:
        typer.echo(
            f"Warning: {sum(issue.count for issue in straight_after)} straight quotes remain:",
            err=True,
        )
        for issue in straight_after[:5]:
            typer.echo(
                f"- {issue.location}, paragraph {issue.paragraph_index + 1}: {issue.snippet}",
                err=True,
            )

    if block_after:
        typer.echo(
            f"Warning: {len(block_after)} wrapped block quote(s) remain:",
            err=True,
        )
        for issue in block_after[:5]:
            typer.echo(
                f"- {issue.location}, paragraph {issue.paragraph_index + 1}: {issue.snippet}...",
                err=True,
            )

    compare_source = backup_path if backup_path.exists() and destination == path else path
    if compare_source != destination:
        ok, message = verify_quote_only_change(compare_source, destination)
    elif backup_path.exists():
        ok, message = verify_quote_only_change(backup_path, destination)
    else:
        ok, message = True, "Skipped content verification (no backup available)"

    typer.echo(message)
    typer.echo(f"Saved {destination}")
    raise typer.Exit(code=0 if ok and not straight_after and not block_after else 1)


def main() -> None:
    app()


if __name__ == "__main__":
    main()
