# -*- coding: utf-8 -*-
"""
Grade-fix pass (all tracked changes, author "Claude", 2026-06-15).

Content fixes (in-place tracked del+ins):
  - Blumer entry 1986 -> 1969, "Univ of California Press" -> "Prentice-Hall" (match in-text 1969)
  - Charmaz (2014) entry: remove doubled title -> "Constructing grounded theory (2nd ed.). Sage."
  - Vidal entry DOI: trailing 07.0200 -> 07.020 (conditional)
  - Methodology tense: will be applied/analyzed/categorized/steered -> past; "yields" -> "yielded"
Reordering (tracked move = del_para_deep old + ins_ref_para at correct slot):
  - Abou Elgheit: fix name/journal/pages/DOI AND move E-slot -> A-slot (before Acharya)
  - Claude: move (before Charmaz) -> after ChatGPT
  - Kim: move -> before Kohavi
  - Parker: move -> before Patton
Formatting:
  - remove the single stray yellow highlight (on the 4.4.3 Risks heading)
"""
from __future__ import annotations
import copy, shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT   = Path(__file__).resolve().parents[1]
DOCX   = ROOT / "Thesis Draft - Daan Luttik - MBA - copy from 2026-06-14.docx"
BACKUP = ROOT / "Thesis Draft - Daan Luttik - MBA - copy from 2026-06-14.gradefix-backup.docx"
AUTHOR = "Claude"
DATE   = "2026-06-15T00:00:00Z"
XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"

# ---------- id helpers ----------
def _used_ids(doc):
    return {int(el.get(qn("w:id"), 0)) for el in doc.element.body.iter()
            if el.get(qn("w:id")) is not None}
def _nid(used):
    n = max(used, default=0) + 1; used.add(n); return n
def _stamp(el, used):
    el.set(qn("w:id"), str(_nid(used))); el.set(qn("w:author"), AUTHOR); el.set(qn("w:date"), DATE)

def ftext(p_elem):
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))

# ---------- run builders (for replace_tracked) ----------
def _mk_run(text, rpr_template=None, italic=False):
    r = etree.Element(qn("w:r"))
    if rpr_template is not None:
        rpr = copy.deepcopy(rpr_template)
        for bad in rpr.findall(qn("w:ins")) + rpr.findall(qn("w:del")):
            rpr.remove(bad)
    else:
        rpr = etree.Element(qn("w:rPr"))
    if italic and rpr.find(qn("w:i")) is None:
        rpr.insert(0, etree.Element(qn("w:iCs"))); rpr.insert(0, etree.Element(qn("w:i")))
    if len(rpr): r.append(rpr)
    t = etree.SubElement(r, qn("w:t")); t.set(XMLSPACE, "preserve"); t.text = text
    return r
def _wrap_ins(children, used):
    ins = etree.Element(qn("w:ins")); _stamp(ins, used)
    for c in children: ins.append(c)
    return ins
def _wrap_del(deleted_text, rpr_template, used):
    d = etree.Element(qn("w:del")); _stamp(d, used)
    r = etree.SubElement(d, qn("w:r"))
    if rpr_template is not None:
        rpr = copy.deepcopy(rpr_template)
        for bad in rpr.findall(qn("w:ins")) + rpr.findall(qn("w:del")):
            rpr.remove(bad)
        if len(rpr): r.append(rpr)
    dt = etree.SubElement(r, qn("w:delText")); dt.set(XMLSPACE, "preserve"); dt.text = deleted_text
    return d

def replace_tracked(p, old, new, used):
    runs = p.findall(qn("w:r")); spans = []; full = ""
    for r in runs:
        t = r.find(qn("w:t")); txt = (t.text if t is not None and t.text else "")
        spans.append((r, len(full), len(full)+len(txt), txt)); full += txt
    idx = full.find(old)
    if idx == -1: return False
    end = idx + len(old); inserted = False
    for r, rs, re_, txt in spans:
        if re_ <= idx or rs >= end: continue
        rpr = r.find(qn("w:rPr"))
        before = txt[:idx-rs] if rs < idx else ""
        after  = txt[end-rs:] if re_ > end else ""
        os_, oe_ = max(idx, rs)-rs, min(end, re_)-rs
        deleted = txt[os_:oe_]
        parent = r.getparent(); i = list(parent).index(r); parts = []
        if before: parts.append(_mk_run(before, rpr))
        if deleted: parts.append(_wrap_del(deleted, rpr, used))
        if not inserted and new:
            parts.append(_wrap_ins([_mk_run(new, rpr)], used)); inserted = True
        if after: parts.append(_mk_run(after, rpr))
        for j, part in enumerate(parts): parent.insert(i+j, part)
        parent.remove(r)
    return True

# ---------- deep tracked delete (handles hyperlinks + nested ins) ----------
def _wrap_run_in_del(r_elem, used):
    d = etree.Element(qn("w:del")); _stamp(d, used)
    rc = copy.deepcopy(r_elem)
    for t in rc.findall(qn("w:t")):
        dt = etree.Element(qn("w:delText")); dt.set(XMLSPACE, "preserve"); dt.text = t.text or ""
        t.getparent().replace(t, dt)
    d.append(rc); return d
def del_para_deep(p_elem, used):
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None: pPr = etree.Element(qn("w:pPr")); p_elem.insert(0, pPr)
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None: rPr = etree.SubElement(pPr, qn("w:rPr"))
    if rPr.find(qn("w:del")) is None:
        dmark = etree.Element(qn("w:del")); _stamp(dmark, used); rPr.insert(0, dmark)
    for r in p_elem.findall(qn("w:r")): p_elem.replace(r, _wrap_run_in_del(r, used))
    for ins in p_elem.findall(qn("w:ins")):
        for r in ins.findall(qn("w:r")): ins.replace(r, _wrap_run_in_del(r, used))
    for hl in p_elem.findall(qn("w:hyperlink")):
        for r in hl.findall(qn("w:r")): hl.replace(r, _wrap_run_in_del(r, used))

def ins_ref_para(segments, used):
    p = etree.Element(qn("w:p"))
    pPr = etree.SubElement(p, qn("w:pPr"))
    ind = etree.SubElement(pPr, qn("w:ind")); ind.set(qn("w:left"), "720"); ind.set(qn("w:hanging"), "720")
    rPr = etree.SubElement(pPr, qn("w:rPr"))
    _stamp(etree.SubElement(rPr, qn("w:ins")), used)
    etree.SubElement(rPr, qn("w:highlight")).set(qn("w:val"), "white")
    ins = etree.SubElement(p, qn("w:ins")); _stamp(ins, used)
    for text, italic in segments:
        r = etree.SubElement(ins, qn("w:r")); rpr = etree.SubElement(r, qn("w:rPr"))
        if italic: etree.SubElement(rpr, qn("w:i")); etree.SubElement(rpr, qn("w:iCs"))
        etree.SubElement(rpr, qn("w:highlight")).set(qn("w:val"), "white")
        t = etree.SubElement(r, qn("w:t")); t.set(XMLSPACE, "preserve"); t.text = text
    return p

# ---------- locators ----------
def find_start(doc, prefix):
    for p in doc.paragraphs:
        if ftext(p._element).strip().startswith(prefix): return p._element
    return None
def find_contains(doc, sub):
    for p in doc.paragraphs:
        if sub in ftext(p._element): return p._element
    return None

def rep(doc, anchor_sub, old, new, used, label):
    p = find_contains(doc, anchor_sub)
    if p is None: print(f"  FAIL (no para) {label}"); return False
    ok = replace_tracked(p, old, new, used)
    print(f"  {'ok  ' if ok else 'MISS'} {label}")
    return ok

def move(doc, old_prefix, anchor_prefix, where, segs, used, label):
    old = find_start(doc, old_prefix)
    anc = find_start(doc, anchor_prefix)
    if old is None or anc is None:
        print(f"  FAIL (old={old is not None}, anchor={anc is not None}) {label}"); return False
    del_para_deep(old, used)
    newp = ins_ref_para(segs, used)
    if where == "before": anc.addprevious(newp)
    else: anc.addnext(newp)
    print(f"  ok   {label}")
    return True

def main():
    shutil.copy(DOCX, BACKUP)
    doc = Document(str(DOCX)); used = _used_ids(doc)

    print("PART A: in-place content fixes")
    rep(doc, "Blumer, H. (1986)", "(1986)", "(1969)", used, "Blumer year 1986->1969")
    rep(doc, "Blumer, H. (", "Univ of California Press", "Prentice-Hall", used, "Blumer publisher")
    rep(doc, "Charmaz, K. (2014)",
        "(introducing qualitative methods series). Constructing grounded theory (2nd ed.)",
        "(2nd ed.)", used, "Charmaz 2014 dedup title")
    rep(doc, "Managing digital transformation", "2022.07.0200", "2022.07.020", used, "Vidal DOI")

    print("PART A2: methodology tense -> past")
    rep(doc, "constructivist grounded theory, a form of qualitative research",
        "constructivist grounded theory, a form of qualitative research, will be applied",
        "constructivist grounded theory, a form of qualitative research, was applied",
        used, "3.1 will be applied")
    rep(doc, "Afterward, the data will be analyzed",
        "the data will be analyzed through the coding process",
        "the data was analyzed through the coding process", used, "3.2 will be analyzed")
    rep(doc, "statements from the interviews will be categorized",
        "statements from the interviews will be categorized",
        "statements from the interviews were categorized", used, "3.2 will be categorized")
    rep(doc, "data gathering will be steered",
        "data gathering will be steered by theoretical sampling",
        "data gathering was steered by theoretical sampling", used, "3.3 will be steered")
    rep(doc, "no longer yields new theoretical insights",
        "no longer yields new theoretical insights",
        "no longer yielded new theoretical insights", used, "3.3 yields->yielded")

    print("PART B: tracked reference moves")
    move(doc, "Elgheit, E. A. (2025)", "Acharya, D. B.", "before", [
        ("Abou Elgheit, E. (2025). Generative AI as a disruptive innovation: Implications for marketing strategic transformations. ", False),
        ("Foresight and STI Governance, 19", True),
        ("(1), 6-15. https://doi.org/10.17323/fstig.2025.24831", False),
    ], used, "Abou Elgheit fix+move E->A")
    move(doc, "Claude. (2025", "ChatGPT. (2026", "after", [
        ("Claude. (2025, October 16). ", False),
        ("Introducing Agent Skills", True),
        (". Claude. https://claude.com/blog/skills", False),
    ], used, "Claude move -> after ChatGPT")
    move(doc, "Kim, J. (2025", "Kohavi, R.", "before", [
        ("Kim, J. (2025). Advertising in the age of agentic AI: Call for research. ", False),
        ("Journal of Interactive Advertising, 25", True),
        ("(3), 215-221.", False),
    ], used, "Kim move -> before Kohavi")
    move(doc, "Parker, C., Scott", "Patton, M. Q.", "before", [
        ("Parker, C., Scott, S., & Geddes, A. (2019). Snowball sampling. ", False),
        ("SAGE research methods foundations", True),
        (".", False),
    ], used, "Parker move -> before Patton")

    print("PART C: remove stray yellow highlight")
    n = 0
    for h in list(doc.element.body.iter(qn("w:highlight"))):
        if h.get(qn("w:val")) == "yellow":
            h.getparent().remove(h); n += 1
    print(f"  ok   removed {n} yellow highlight(s)")

    doc.save(str(DOCX))
    print(f"\nSaved: {DOCX.name}\nBackup: {BACKUP.name}")

if __name__ == "__main__":
    main()
