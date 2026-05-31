"""Add Shenton (2004) and Lincoln & Guba (1985) to thesis reference list in docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"

LINCOLN = (
    "Lincoln, Y. S., & Guba, E. G. (1985). Naturalistic inquiry. Sage."
)
SHENTON = (
    "Shenton, A. K. (2004). Strategies for ensuring trustworthiness in "
    "qualitative research projects. Education for Information, 22(2), 63-75. "
    "https://doi.org/10.3233/EFI-2004-22201"
)


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = "normal"
    new_para.add_run(text)
    return new_para


def main() -> None:
    doc = Document(str(DOCX_PATH))
    text = "\n".join(p.text for p in doc.paragraphs)
    if "Shenton, A. K. (2004)" not in text:
        anchor = next(p for p in doc.paragraphs if p.text.startswith("Schafer,"))
        insert_after(anchor, SHENTON)
    text = "\n".join(p.text for p in doc.paragraphs)
    if "Lincoln, Y. S." not in text:
        idx = next(
            i
            for i, p in enumerate(doc.paragraphs)
            if p.text.startswith("Lewis, P.") or "9459" in p.text
        )
        while idx + 1 < len(doc.paragraphs):
            nxt = doc.paragraphs[idx + 1].text
            if nxt.startswith("Li,") or nxt.startswith("[Li,") or nxt.startswith("Little"):
                break
            idx += 1
        insert_after(doc.paragraphs[idx], LINCOLN)
    doc.save(str(DOCX_PATH))
    print(f"Updated references in {DOCX_PATH}")


if __name__ == "__main__":
    main()
