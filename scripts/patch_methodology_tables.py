# -*- coding: utf-8 -*-
"""Methodology table updates (tracked, author 'Claude').

1. Realign the coding table (Table 3, "Overview of code development") to chapter 4:
   focused codes map 1:1 to the ch4 subsections; condition rows balanced into
   enablers/obstacles; supplier communication folded into AI progression. Done by
   deleting the old table (tracked) and inserting a style-cloned realigned table.
2. Insert a new definitions table (Table 4) for the four theoretical categories,
   worded from ch4 §4.1-4.4, right after the coding table.
3. Renumber the divergent-outcomes caption "Table 4" -> "Table 5".

All-or-nothing save with backup.  Usage:
    python patch_methodology_tables.py --check   # validate, no write
    python patch_methodology_tables.py           # apply + save
"""
from __future__ import annotations
import copy, shutil, sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Thesis Draft - Daan Luttik - MBA - copy from 2026-06-13.docx"
BACKUP = ROOT / "Thesis Draft - Daan Luttik - MBA - copy from 2026-06-13.tables-backup.docx"
AUTHOR = "Claude"; DATE = "2026-06-13T16:00:00Z"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _used_ids(doc):
    return {int(el.get(qn("w:id"), 0)) for el in doc.element.body.iter()
            if el.get(qn("w:id")) is not None}
def _nid(used): n = max(used, default=0) + 1; used.add(n); return n
def full_text(p): return "".join(t.text or "" for t in p.iter(qn("w:t")))
def _esc(s): return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
def _ins_attr(used): return f'w:id="{_nid(used)}" w:author="{AUTHOR}" w:date="{DATE}"'


# ---- builders ---- #
def make_run(text, italic, bold):
    r = etree.Element(qn("w:r")); rPr = etree.SubElement(r, qn("w:rPr"))
    rf = etree.SubElement(rPr, qn("w:rFonts")); rf.set(qn("w:cs"), "Segoe UI")
    if bold: etree.SubElement(rPr, qn("w:b")); etree.SubElement(rPr, qn("w:bCs"))
    if italic: etree.SubElement(rPr, qn("w:i")); etree.SubElement(rPr, qn("w:iCs"))
    col = etree.SubElement(rPr, qn("w:color")); col.set(qn("w:val"), "0A0A0A")
    sz = etree.SubElement(rPr, qn("w:sz")); sz.set(qn("w:val"), "20")
    szcs = etree.SubElement(rPr, qn("w:szCs")); szcs.set(qn("w:val"), "20")
    t = etree.SubElement(r, qn("w:t")); t.set(XMLSPACE, "preserve"); t.text = text
    return r


def set_cell(tc, runs, used):
    """runs: list of (text, italic, bold). Rebuilds the cell paragraph as a tracked insertion."""
    p = tc.find(qn("w:p"))
    for el in list(p):
        if etree.QName(el).localname in ("r", "commentRangeStart", "commentRangeEnd", "commentReference"):
            p.remove(el)
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.Element(qn("w:pPr")); p.insert(0, pPr)
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None: rPr = etree.SubElement(pPr, qn("w:rPr"))
    im = etree.Element(qn("w:ins"))
    im.set(qn("w:id"), str(_nid(used))); im.set(qn("w:author"), AUTHOR); im.set(qn("w:date"), DATE)
    rPr.insert(0, im)
    if runs:
        ins = etree.SubElement(p, qn("w:ins"))
        ins.set(qn("w:id"), str(_nid(used))); ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), DATE)
        for (text, italic, bold) in runs:
            ins.append(make_run(text, italic, bold))


def mark_row_inserted(tr, used):
    trPr = tr.find(qn("w:trPr"))
    if trPr is None: trPr = etree.Element(qn("w:trPr")); tr.insert(0, trPr)
    ins = etree.SubElement(trPr, qn("w:ins"))
    ins.set(qn("w:id"), str(_nid(used))); ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), DATE)


def delete_table_tracked(tbl, used):
    for tr in tbl.findall(qn("w:tr")):
        trPr = tr.find(qn("w:trPr"))
        if trPr is None: trPr = etree.Element(qn("w:trPr")); tr.insert(0, trPr)
        d = etree.SubElement(trPr, qn("w:del"))
        d.set(qn("w:id"), str(_nid(used))); d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), DATE)
        for tc in tr.findall(qn("w:tc")):
            for p in tc.findall(qn("w:p")):
                pPr = p.find(qn("w:pPr"))
                if pPr is None: pPr = etree.Element(qn("w:pPr")); p.insert(0, pPr)
                rPr = pPr.find(qn("w:rPr"))
                if rPr is None: rPr = etree.SubElement(pPr, qn("w:rPr"))
                dm = etree.Element(qn("w:del"))
                dm.set(qn("w:id"), str(_nid(used))); dm.set(qn("w:author"), AUTHOR); dm.set(qn("w:date"), DATE)
                rPr.insert(0, dm)
                for r in p.findall(qn("w:r")):
                    t = r.find(qn("w:t"))
                    if t is not None:
                        dt = etree.Element(qn("w:delText")); dt.set(XMLSPACE, "preserve"); dt.text = t.text or ""
                        t.getparent().replace(t, dt)
                    dd = etree.Element(qn("w:del"))
                    dd.set(qn("w:id"), str(_nid(used))); dd.set(qn("w:author"), AUTHOR); dd.set(qn("w:date"), DATE)
                    r.addprevious(dd); dd.append(r)


def caption_para(runs_spec, used):
    p = etree.Element(qn("w:p")); pPr = etree.SubElement(p, qn("w:pPr"))
    ind = etree.SubElement(pPr, qn("w:ind")); ind.set(qn("w:firstLine"), "0")
    rPr = etree.SubElement(pPr, qn("w:rPr"))
    im = etree.SubElement(rPr, qn("w:ins"))
    im.set(qn("w:id"), str(_nid(used))); im.set(qn("w:author"), AUTHOR); im.set(qn("w:date"), DATE)
    ins = etree.SubElement(p, qn("w:ins"))
    ins.set(qn("w:id"), str(_nid(used))); ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), DATE)
    for (text, bold, italic) in runs_spec:
        r = etree.SubElement(ins, qn("w:r")); rp = etree.SubElement(r, qn("w:rPr"))
        if bold: etree.SubElement(rp, qn("w:b")); etree.SubElement(rp, qn("w:bCs"))
        if italic: etree.SubElement(rp, qn("w:i")); etree.SubElement(rp, qn("w:iCs"))
        t = etree.SubElement(r, qn("w:t")); t.set(XMLSPACE, "preserve"); t.text = text
    return p


def build_defs_table(used):
    widths = [2800, 6814]
    header = ["Theoretical category", "Definition"]
    rows = [
        ["Observing external conditions",
         "The managerial process of monitoring and interpreting the organization’s changing context. "
         "Two categories of external conditions sparked change: the progression of AI itself, and market pressure."],
        ["Navigating the organizational context",
         "How managers steer or move around conditions within the organization, across three scopes: the "
         "marketing department (which the manager can steer to create alignment), technical resources (which "
         "can be leveraged to implement initiatives), and the compliance framework (which bounds what is possible)."],
        ["Applying agentic AI",
         "Using or implementing agentic AI systems within the organization, across key marketing use cases: "
         "generating insights, creating and validating content and campaigns, utilizing generic agents for "
         "personal work, and employing customer-facing agents."],
        ["Obtaining value outcomes",
         "The outcomes that applying agentic AI creates, analyzed as value in three buckets: benefits (features "
         "beneficial to the team, organization, or customer), sacrifices (negative aspects highly likely to "
         "materialize), and risks (potential negative aspects whose likelihood is unlikely or unknown)."],
    ]
    def cell(text, width, is_head):
        bold = "<w:b/><w:bCs/>" if is_head else ""
        borders = ('<w:tcBorders><w:top w:val="single" w:sz="8" w:space="0" w:color="auto"/>'
                   '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="auto"/></w:tcBorders>') if is_head else ""
        return (f'<w:tc><w:tcPr><w:tcW w:w="{width}" w:type="dxa"/>{borders}</w:tcPr>'
                f'<w:p><w:pPr><w:pStyle w:val="NoSpacing"/><w:ind w:firstLine="0"/>'
                f'<w:rPr><w:ins {_ins_attr(used)}/><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr></w:pPr>'
                f'<w:ins {_ins_attr(used)}><w:r><w:rPr><w:rFonts w:cs="Segoe UI"/>{bold}'
                f'<w:color w:val="0A0A0A"/><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr>'
                f'<w:t xml:space="preserve">{_esc(text)}</w:t></w:r></w:ins></w:p></w:tc>')
    def row(cells, is_head):
        body = "".join(cell(c, widths[i], is_head) for i, c in enumerate(cells))
        return f'<w:tr><w:trPr><w:trHeight w:val="255"/><w:ins {_ins_attr(used)}/></w:trPr>{body}</w:tr>'
    grid = "<w:tblGrid>" + "".join(f'<w:gridCol w:w="{w}"/>' for w in widths) + "</w:tblGrid>"
    tblPr = ('<w:tblPr><w:tblW w:w="9614" w:type="dxa"/><w:tblLook w:val="04A0" w:firstRow="1" '
             'w:lastRow="0" w:firstColumn="1" w:lastColumn="0" w:noHBand="0" w:noVBand="1"/></w:tblPr>')
    body = row(header, True) + "".join(row(r, False) for r in rows)
    return etree.fromstring(f'<w:tbl xmlns:w="{W}">{tblPr}{grid}{body}</w:tbl>'.encode("utf-8"))


# ---- realigned coding-table content ---- #
I, B = True, False
HEADER = ["Open coding", "Focused coding", "Theoretical coding"]
DATA = [
    ([("AI’s improvement & potential; Increasing AI adoption; Vendor roadmaps & offerings", B, B)],
     "AI progression", "Observing external conditions"),
    ([("Competitor pressure; Changing consumer behavior; The rise of agents of consumers", B, B)],
     "Market pressure", ""),
    ([("Enablers:", I, B), (" bringing people along; educating & training; experimenting with AI; providing "
      "clarity & direction; leadership backing; AI champion; being opportunistic; innovation culture. ", B, B),
      ("Obstacles:", I, B), (" resistance to change; limited AI literacy; analysis paralysis & delay; absent "
      "strategy or focus; punishing failure", B, B)],
     "Steering the marketing department", "Navigating the organizational context"),
    ([("Enablers:", I, B), (" having data & infrastructure; having the right tooling; having technical talent "
      "& skills; leveraging systems & process thinking; leveraging external experts & vendors. ", B, B),
      ("Obstacles:", I, B), (" lacking data & infrastructure; poor data quality; lacking skills to execute; "
      "gap between marketing & IT", B, B)],
     "Leveraging technical resources", ""),
    ([("Enablers:", I, B), (" working around constraints (shadow IT); accepting bounded risk. ", B, B),
      ("Obstacles:", I, B), (" restrictive governance & bureaucracy; politics & silos; legal & data-sovereignty "
      "concerns", B, B)],
     "Dealing with compliance", ""),
    ([("Analyzing data; Generating insights & monitoring; Ingesting data & context", B, B)],
     "Generating insights", "Applying agentic AI"),
    ([("Creating & validating content; Building campaign pipelines; Segmenting & targeting; Controlling brand "
      "& messaging", B, B)],
     "Creating & validating content & campaigns", ""),
    ([("Using general-purpose assistants; Automating manual processes; Leveraging tool calling & integrations; "
      "Building agent harnesses", B, B)],
     "Utilizing generic agents for personal work", ""),
    ([("Deploying customer-service & sales agents; Supporting customers’ own agents", B, B)],
     "Employing customer-facing agents", ""),
    ([("Gaining efficiency & speed; Gaining scale; Extending the personal skillset; Improving quality of output", B, B)],
     "Benefits", "Obtaining value outcomes"),
    ([("Incurring financial costs; Displacing jobs", B, B)], "Sacrifices", ""),
    ([("Risking hallucination; Risking security & privacy violations; Risking brand degradation", B, B)],
     "Risks", ""),
]


def find_coding_table(doc):
    for tbl in doc.element.body.findall(qn("w:tbl")):
        tr0 = tbl.find(qn("w:tr"))
        if tr0 is not None:
            txt = "".join(t.text or "" for t in tr0.iter(qn("w:t")))
            if "Open coding" in txt and "Theoretical coding" in txt:
                return tbl
    return None


def find_caption_table4(doc):
    for p in doc.paragraphs:
        if full_text(p._element).strip() == "Table 4":
            return p._element
    return None


def build_new_coding_table(old_tbl, used):
    new = copy.deepcopy(old_tbl)
    trs = new.findall(qn("w:tr"))
    header_tr = trs[0]
    template_tr = copy.deepcopy(trs[3])
    for tr in trs[1:]:
        new.remove(tr)
    for i, tc in enumerate(header_tr.findall(qn("w:tc"))):
        set_cell(tc, [(HEADER[i], False, True)], used)
    mark_row_inserted(header_tr, used)
    for open_runs, focused, theo in DATA:
        tr = copy.deepcopy(template_tr)
        tcs = tr.findall(qn("w:tc"))
        set_cell(tcs[0], open_runs, used)
        set_cell(tcs[1], [(focused, False, False)], used)
        set_cell(tcs[2], [(theo, False, False)] if theo else [], used)
        mark_row_inserted(tr, used)
        new.append(tr)
    return new


def renumber_caption(p_elem, used):
    for r in p_elem.findall(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is not None and t.text and "4" in t.text:
            rPr = r.find(qn("w:rPr"))
            old, new = t.text, t.text.replace("4", "5")
            d = etree.Element(qn("w:del"))
            d.set(qn("w:id"), str(_nid(used))); d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), DATE)
            dr = etree.SubElement(d, qn("w:r"))
            if rPr is not None: dr.append(copy.deepcopy(rPr))
            dt = etree.SubElement(dr, qn("w:delText")); dt.set(XMLSPACE, "preserve"); dt.text = old
            ins = etree.Element(qn("w:ins"))
            ins.set(qn("w:id"), str(_nid(used))); ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), DATE)
            ir = etree.SubElement(ins, qn("w:r"))
            if rPr is not None: ir.append(copy.deepcopy(rPr))
            it = etree.SubElement(ir, qn("w:t")); it.set(XMLSPACE, "preserve"); it.text = new
            par = r.getparent(); idx = list(par).index(r)
            par.insert(idx, ins); par.insert(idx, d); par.remove(r)
            return True
    return False


def main():
    check = "--check" in sys.argv
    doc = Document(str(DOCX)); used = _used_ids(doc)
    coding = find_coding_table(doc)
    cap4 = find_caption_table4(doc)
    problems = []
    if coding is None: problems.append("coding table (Open coding header) NOT FOUND")
    elif len(coding.findall(qn("w:tr"))) != 16:
        problems.append(f"coding table has {len(coding.findall(qn('w:tr')))} rows, expected 16")
    if cap4 is None: problems.append("divergent caption 'Table 4' NOT FOUND")

    print(f"Validating against: {DOCX.name}")
    print(f"  coding table: {'found' if coding is not None else 'MISSING'}; "
          f"caption Table 4: {'found' if cap4 is not None else 'MISSING'}")
    print(f"  realigned coding rows: {len(DATA)} + header; definitions rows: 4 + header")
    if problems:
        for p in problems: print("    PROBLEM:", p)
    if check:
        print("\n--check only: no file written.")
        return
    if problems:
        print("\nNOT SAVED - resolve problems first."); return

    # build new structures
    new_coding = build_new_coding_table(coding, used)
    defcap1 = caption_para([("Table 4", True, False)], used)
    defcap2 = caption_para([("Definitions of the theoretical categories", False, True)], used)
    defs_tbl = build_defs_table(used)

    # delete old coding table (tracked)
    delete_table_tracked(coding, used)

    # insert after the empty paragraph that follows the old coding table
    parent = coding.getparent()
    anchor = coding.getnext()
    if anchor is None or etree.QName(anchor).localname != "p":
        anchor = coding
    idx = list(parent).index(anchor)
    for off, el in enumerate([new_coding, defcap1, defcap2, defs_tbl], start=1):
        parent.insert(idx + off, el)

    # renumber divergent caption Table 4 -> Table 5
    if not renumber_caption(cap4, used):
        print("WARNING: caption renumber did not match a run")

    shutil.copy(DOCX, BACKUP); doc.save(str(DOCX))
    print(f"\nApplied. Saved: {DOCX.name}\nBackup: {BACKUP.name}")


if __name__ == "__main__":
    main()
