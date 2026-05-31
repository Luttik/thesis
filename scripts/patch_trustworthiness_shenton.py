"""Align §3.4 trustworthiness table with Shenton (2004); preserve user credibility/transferability rows."""

from __future__ import annotations

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
TRUST_HDR = "Criteria of trustworthiness"

# User-edited rows (how addressed) — preserved verbatim
CREDIBILITY_HOW = (
    "The study applied constructivist grounded theory (Charmaz, 2014) with iterative "
    "interviewing, memo writing, and coding (Appendix C). Credibility was strengthened "
    "through triangulation across marketing practitioners and AI experts; iterative "
    "initial and focused coding (Table 3); theoretical sampling; retention of negative "
    "and paradoxical accounts; the researcher had prior familiarity with the field and "
    "applied reflexivity (Section 3.2.1); and regular thesis-supervisor debriefing on "
    "transcripts, memos, and emerging categories. Member checking was not conducted "
    "(Charmaz, 2014). Theoretical saturation was assessed during data collection: later "
    "interviews chiefly densified existing categories rather than introducing new "
    "top-level concepts, with seventeen interviews completed (Table 1)."
)

TRANSFERABILITY_HOW = (
    "In line with Shenton (2004), information regarding the sampling method, participant "
    "types, data collection methods, and the February–May 2026 fieldwork period are provided "
    "in this chapter (Sections 3.2–3.3; Table 1; Appendix A). Thick description in Chapter 4 "
    "and scope conditions in Section 5.3 allow readers to judge whether findings may apply in "
    "comparable settings."
)

INTRO = (
    "Qualitative inquiry does not seek statistical generalization to a population; readers "
    "instead judge whether findings are sufficiently grounded in data and whether the "
    "inquiry was conducted transparently (Lincoln & Guba, 1985). This study addresses "
    "trustworthiness through the four criteria Lincoln and Guba propose—credibility, "
    "transferability, dependability, and confirmability—and implements strategies "
    "adapted from Shenton (2004) for interview-based projects of limited duration. "
    "Shenton (2004) notes that prolonged engagement and persistent observation, while "
    "valuable in ethnography, are not realistic in shorter studies; the strategies below "
    "were selected accordingly. Table 2 summarises how each criterion was addressed."
)

ROWS: list[tuple[str, str]] = [
    (
        "Credibility concerns the congruency between the findings and the phenomenon "
        "studied (Merriam, 1998; Shenton, 2004).",
        CREDIBILITY_HOW,
    ),
    (
        "Transferability requires sufficient contextual detail for readers to judge "
        "whether findings may apply elsewhere—not statistical generalisation to a "
        "wider population (Lincoln & Guba, 1985; Shenton, 2004).",
        TRANSFERABILITY_HOW,
    ),
    (
        "Dependability concerns whether the inquiry process is consistent, auditable, "
        "and reported in enough depth that a future researcher could follow the same "
        "steps (Shenton, 2004).",
        "Interviews were transcribed; transcriptions, memos, and emerging models were "
        "reviewed in regular meetings with the thesis supervisor (debriefing; Shenton, 2004). "
        "Anonymised transcripts and supporting materials were submitted to the research "
        "institution for examination and secure archiving. The research design, "
        "field procedures, and reflective appraisal are documented in Chapter 3 and "
        "Appendix A (Shenton, 2004). Triangulation across marketing managers and AI "
        "experts provided overlapping methods. The code structure is reported in Table 3; "
        "quotation-level codings, QDPX projects, and analytic memos form an audit trail "
        "linking data to Findings.",
    ),
    (
        "Confirmability requires that findings emerge from participant accounts rather "
        "than researcher predisposition (Lincoln & Guba, 1985; Shenton, 2004).",
        "In line with Shenton (2004), analytic memos recorded reasoning separately from "
        "transcript text; preliminary interpretations that were not supported by later "
        "data were revised or discarded. Findings claims are anchored in participant "
        "quotations (Chapter 4; Appendix D). Reflexive documentation of the researcher's "
        "marketing and AI background and sampling implications appears in Section 3.2.1.",
    ),
    (
        "Informant honesty and rapport affect credibility of what participants report "
        "(Shenton, 2004).",
        "Participation was voluntary; interviews were semi-structured, confidential, and "
        "conducted in a non-threatening manner, with participants encouraged to speak "
        "frankly. Table 1 reports participants in anonymised form; Findings use "
        "interviewee labels tied to verbatim transcript evidence.",
    ),
]


def _find_trust_table(doc: Document):
    for table in doc.tables:
        if table.rows and TRUST_HDR in table.rows[0].cells[0].text:
            return table
    return None


def _set_section_34_intro(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("3.4") and "rustworthiness" in t.lower():
            continue
        if "Shenton" in t and ("trustworthiness" in t.lower() or "Table 2" in t):
            p.text = INTRO
            return
    msg = "Could not find §3.4 intro paragraph to update."
    raise ValueError(msg)


def main() -> None:
    doc = Document(str(DOCX_PATH))
    table = _find_trust_table(doc)
    if table is None:
        from convert_trustworthiness_table import (
            _find_findings_paragraph,
            _insert_new_section,
        )

        findings = _find_findings_paragraph(doc)
        if not any(
            p.text.strip().startswith("3.4") and "rustworthiness" in p.text.lower()
            for p in doc.paragraphs
        ):
            h = findings.insert_paragraph_before("3.4.\tTrustworthiness")
            try:
                h.style = doc.styles["Heading 2"]
            except KeyError:
                pass
        _insert_new_section(doc, findings)
        doc.save(str(DOCX_PATH))
        doc = Document(str(DOCX_PATH))
        table = _find_trust_table(doc)
        if table is None:
            raise SystemExit("Failed to insert trustworthiness table.")

    _set_section_34_intro(doc)

    # Drop empty rows; keep header + data rows
    while len(table.rows) > len(ROWS) + 1:
        last = table.rows[-1]
        if not any(c.text.strip() for c in last.cells):
            table._tbl.remove(last._tr)
        else:
            break

    for i, (crit, how) in enumerate(ROWS, start=1):
        if i >= len(table.rows):
            table.add_row()
        row = table.rows[i].cells
        row[0].text = crit
        row[1].text = how

    # Remove extra rows beyond our five criteria
    while len(table.rows) > len(ROWS) + 1:
        table._tbl.remove(table.rows[-1]._tr)

    try:
        doc.save(str(DOCX_PATH))
        print(f"Updated trustworthiness section in {DOCX_PATH}")
    except PermissionError:
        fallback = DOCX_PATH.with_name(DOCX_PATH.stem + " - trustworthiness-shenton.docx")
        doc.save(str(fallback))
        print(f"File locked; saved to {fallback}")


if __name__ == "__main__":
    main()
