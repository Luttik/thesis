"""
Restructure section 4.2 per the approved plan:

1. Rename 4.2.1/4.2.2/4.2.3 headings; add 4.2.4
2. Update 4.2 intro: four conditions, causal interaction, portfolio+dependency framing
3. Split 4.2.1 response para into #### Educating + #### Experimenting sub-blocks
4. Add portfolio framing sentence to 4.2.2
5. Update 4.2.3 intro; add "providing clarity" paragraph
6. Insert 4.2.4 heading before bureaucracy content
7. Update closing paragraph

All changes are tracked (author "Claude").
"""
from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
OUT_PATH  = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"

AUTHOR = "Claude"
DATE   = "2026-06-05T00:00:00Z"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _all_rev_ids(doc: Document) -> set[int]:
    return {
        int(el.get(qn("w:id"), 0))
        for el in doc.element.body.iter()
        if el.get(qn("w:id")) is not None
    }

def _next_id(used: set[int]) -> int:
    n = max(used, default=0) + 1
    used.add(n)
    return n


def ins_para(text: str, style: str, rev_ids: set[int],
             block_quote: bool = False) -> etree._Element:
    """New paragraph with a single tracked-inserted run."""
    p = etree.Element(qn("w:p"))
    pPr = etree.SubElement(p, qn("w:pPr"))
    if style and style.lower() not in ("normal", ""):
        ps = etree.SubElement(pPr, qn("w:pStyle"))
        ps.set(qn("w:val"), style)
    if block_quote:
        ind = etree.SubElement(pPr, qn("w:ind"))
        ind.set(qn("w:left"), "720")

    ins = etree.SubElement(p, qn("w:ins"))
    ins.set(qn("w:id"), str(_next_id(rev_ids)))
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), DATE)
    r = etree.SubElement(ins, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return p


def mark_para_deleted(elem: etree._Element, rev_ids: set[int]) -> None:
    """In-place: wrap all direct w:r children in <w:del> and convert w:t → w:delText."""
    for r in list(elem.findall(qn("w:r"))):
        parent = r.getparent()
        idx = list(parent).index(r)
        # Convert w:t → w:delText inside the run
        for t in r.findall(qn("w:t")):
            dt = etree.Element(qn("w:delText"))
            dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            dt.text = t.text or ""
            t.getparent().replace(t, dt)
        # Wrap run in <w:del>
        d = etree.Element(qn("w:del"))
        d.set(qn("w:id"), str(_next_id(rev_ids)))
        d.set(qn("w:author"), AUTHOR)
        d.set(qn("w:date"), DATE)
        parent.insert(idx, d)
        parent.remove(r)
        d.append(r)
    # Mark paragraph mark as deleted (pPr > rPr > del)
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(elem, qn("w:pPr"))
        elem.insert(0, pPr)
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = etree.SubElement(pPr, qn("w:rPr"))
    para_del = etree.SubElement(rPr, qn("w:del"))
    para_del.set(qn("w:id"), str(_next_id(rev_ids)))
    para_del.set(qn("w:author"), AUTHOR)
    para_del.set(qn("w:date"), DATE)


def rename_heading(elem: etree._Element, new_text: str, rev_ids: set[int]) -> None:
    """In-place: wrap existing runs in <w:del>, append <w:ins> with new_text."""
    # Delete existing runs
    for r in list(elem.findall(f"{qn('w:r')}")):
        idx = list(elem).index(r)
        d = etree.Element(qn("w:del"))
        d.set(qn("w:id"), str(_next_id(rev_ids)))
        d.set(qn("w:author"), AUTHOR)
        d.set(qn("w:date"), DATE)
        for t in r.findall(qn("w:t")):
            dt = etree.SubElement(r, qn("w:delText"))
            dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            dt.text = t.text or ""
            r.remove(t)
        elem.insert(idx, d)
        d.append(r)
    # Insert new text
    ins = etree.SubElement(elem, qn("w:ins"))
    ins.set(qn("w:id"), str(_next_id(rev_ids)))
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), DATE)
    r = etree.SubElement(ins, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = new_text


def insert_after(anchor: etree._Element, new_paras: list[etree._Element]) -> None:
    parent = anchor.getparent()
    idx = list(parent).index(anchor) + 1
    for i, p in enumerate(new_paras):
        parent.insert(idx + i, p)


def insert_before(anchor: etree._Element, new_paras: list[etree._Element]) -> None:
    parent = anchor.getparent()
    idx = list(parent).index(anchor)
    for i, p in enumerate(new_paras):
        parent.insert(idx + i, p)


def find_para(doc: Document, marker: str) -> int | None:
    """Find paragraph by marker, searching all w:t elements (incl. inside w:ins)."""
    for i, p in enumerate(doc.paragraphs):
        # Collect all text from w:t elements anywhere in the paragraph XML
        full_text = "".join(
            t.text or ""
            for t in p._element.iter(qn("w:t"))
        )
        if marker in full_text:
            return i
    return None


def replace_para(doc: Document, para_idx: int,
                 new_paras: list[etree._Element],
                 rev_ids: set[int]) -> None:
    """Delete para at para_idx (tracked in-place), insert new_paras after it."""
    elem = doc.paragraphs[para_idx]._element
    mark_para_deleted(elem, rev_ids)
    insert_after(elem, new_paras)


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

# --- 4.2 intro update ---

INTRO_P1_NEW = (
    "Internal conditions describe a set of organizational variables that shape whether "
    "and how agentic AI can be implemented and create value. Four interlocking themes "
    "recur throughout the data: (1) organizational capacity, encompassing AI literacy, "
    "resistance to change, and the ability to act quickly under uncertainty; (2) technical "
    "resources, including data infrastructure, tooling, and the systems-thinking capability "
    "to connect them; (3) strategic direction and leadership, covering the presence or "
    "absence of clear vision, leadership backing, and active change management; and "
    "(4) governance and organizational friction, including IT/legal/compliance gatekeeping, "
    "politics, and silos."
)

INTRO_CAUSAL = (
    "These four conditions are not independent: poor strategic direction amplifies "
    "capacity problems; resource gaps deepen paralysis; governance friction compounds "
    "all of the above. Strong resources do not substitute for clarity, and clarity "
    "does not substitute for literacy."
)

INTRO_PORTFOLIO = (
    "The behaviors managers apply in response to these conditions are overlapping and "
    "mutually reinforcing — there is no clean one-to-one mapping between conditions "
    "and responses. Experimentation builds literacy as well as countering paralysis; "
    "education reduces resistance as well as closing the literacy gap; providing "
    "clarity reduces resistance as well as setting direction. The data suggests a loose "
    "dependency: clarity and basic literacy tend to precede productive experimentation; "
    "infrastructure tends to precede scalable automation. The sections below organize "
    "around conditions — what managers find — but the behaviors should be read as a "
    "portfolio with rough sequencing, not a diagnostic chart."
)

# --- 4.2.1 replacement for paras 181+182 ---

EDU_HEADING_TEXT = "Educating and training"

EDU_PARA = (
    "Educating and training (28 code applications, 14 documents) is the primary response "
    "to the literacy and understanding barrier. Managers invest deliberately in raising "
    "AI capability across the organization — through training sessions, workshops, "
    "awareness programs, and structured learning — to create the cognitive preconditions "
    "for effective adoption. While education primarily targets the understanding deficit, "
    "it also reduces resistance by demystifying AI: employees who understand what AI can "
    "and cannot do are less fearful of it. However, education alone is insufficient to "
    "sustain change. Interviewee 17 observed a common pattern in her training work: "
    "“After two or three trainings companies basically stop, because then they go "
    "back to what they started with.” Sustained capacity-building requires not just "
    "initial education but ongoing structural embedding of learning."
)

EXP_HEADING_TEXT = "Experimenting with AI"

EXP_PARA = (
    "Experimenting with AI (37 code applications, 17 documents) primarily addresses "
    "resistance and analysis paralysis by making the value of AI tangible through doing "
    "rather than planning. Beyond this primary role, experimentation also surfaces "
    "use cases that were invisible in the abstract, and builds tacit literacy that "
    "structured education cannot easily provide. Rather than waiting for perfect "
    "conditions, effective managers create structured opportunities to learn through "
    "doing. Interviewee 6 summarized this orientation: “I’d rather have "
    "something live and test and improve than have two years down the road and then "
    "finally have something live.” Interviewee 13 described enabling this at the "
    "team level: “Provide freedom to experiment and don’t set this atmosphere "
    "of fear that AI is going to replace you.”"
)

NAVIGATE_PARA = (
    "A third pattern is less visible but analytically important: some managers navigate "
    "around organizational capacity constraints rather than reshaping them. When internal "
    "literacy is too low or resistance too entrenched to address in the short term, "
    "managers route AI initiatives through external experts, agencies, or low-code "
    "tooling that reduces the literacy bar. This bypasses the bottleneck without fixing "
    "it — a pragmatic short-term strategy that can stall long-term capability development "
    "if it becomes the default."
)

# --- 4.2.2 portfolio framing ---

PORTFOLIO_422 = (
    "Managers respond to resource conditions in three patterns: leveraging available "
    "infrastructure to build and deploy fast; navigating gaps through external "
    "partnerships and workarounds; and reshaping the foundation through deliberate "
    "data-infrastructure investment. The same resource base can function as either "
    "an accelerator or a bottleneck depending on whether the organizational capacity "
    "described in section 4.2.1 is in place to use it."
)

# --- 4.2.3 intro update and providing clarity ---

SEC423_INTRO_NEW = (
    "A third cluster of conditions concerns strategic direction and leadership — "
    "the organizational structures and behaviors that orient AI adoption toward "
    "meaningful goals. Where section 4.2.1 described the people-level capacity "
    "to change, this section describes the organizational-level direction that "
    "makes change purposeful. The data surfaces four related dynamics: the role "
    "of providing clarity, the impact of individual AI champions, the importance "
    "of leadership backing, and the process of bringing people along."
)

CLARITY_PARA = (
    "Providing clarity — naming where the organization is going with AI, which "
    "use cases matter, and what success looks like — emerged as the most broadly "
    "distributed enabling behavior in this section (18 code applications, 11 documents). "
    "Unlike championing, which is driven by individuals, or leadership backing, which "
    "is top-down, providing clarity is a managerial act that can happen at any "
    "organizational level. Its cross-condition effect is significant: clarity directly "
    "reduces the resistance and paralysis described in section 4.2.1, because employees "
    "who understand the direction and purpose of AI adoption are less likely to resist "
    "or stall. Interviewee 6 captured the essential move: “This is the vision. "
    "This is where we’re going. Help me get there.” Interviewee 1 noted "
    "the downstream consequence of its absence: “If you don’t have clarity, "
    "jumping into execution will only surface a lot of problems.”"
)

# --- 4.2.4 heading ---

SEC424_HEADING_TEXT = "4.2.4\tGovernance & organizational friction"

SEC424_INTRO_NEW = (
    "A fourth cluster of conditions concerns governance and organizational friction — "
    "the IT/legal/compliance gatekeeping, procurement processes, politics, and silos "
    "that impede AI adoption. Unlike the conditions in sections 4.2.1 to 4.2.3, which "
    "managers actively reshape, governance friction is predominantly navigated: "
    "unfavorable governance structures are treated as fixed features of the environment "
    "to be worked around rather than changed. This tendency is itself a finding: "
    "reshaping governance — actually changing the structures rather than routing "
    "around them — is rare in the data."
)

# --- Closing paragraph update ---

CLOSING_NEW = (
    "Taken together, the four conditions described in sections 4.2.1 to 4.2.4 represent "
    "the organizational substrate on which AI value creation depends. They are not "
    "independent: poor direction amplifies capacity problems; resource gaps deepen "
    "paralysis; governance friction compounds both. The behaviors documented across "
    "these sections — educating, experimenting, providing clarity, bringing people "
    "along, leveraging resources, navigating friction — address multiple conditions "
    "simultaneously and in rough sequence. Organizations that neglect any one dimension "
    "tend to find that the others stall. This pattern surfaces most acutely in section "
    "4.5.3: organizations that invest only in experimentation without clarity or "
    "education tend to generate pilots that do not scale."
)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    backup = DOCX_PATH.with_suffix(".bak.docx")
    shutil.copy(DOCX_PATH, backup)
    print(f"Backup: {backup}")

    doc = Document(str(DOCX_PATH))
    body = doc.element.body
    rev_ids = _all_rev_ids(doc)

    # ------------------------------------------------------------------
    # 1. Rename 4.2.1 heading
    # ------------------------------------------------------------------
    idx_421 = find_para(doc, "4.2.1")
    assert idx_421 is not None, "4.2.1 heading not found"
    rename_heading(doc.paragraphs[idx_421]._element,
                   "4.2.1\tOrganizational capacity", rev_ids)
    print("  Renamed 4.2.1 heading.")

    # ------------------------------------------------------------------
    # 2. Rename 4.2.2 heading
    # ------------------------------------------------------------------
    idx_422 = find_para(doc, "4.2.2")
    assert idx_422 is not None, "4.2.2 heading not found"
    rename_heading(doc.paragraphs[idx_422]._element,
                   "4.2.2\tTechnical resources", rev_ids)
    print("  Renamed 4.2.2 heading.")

    # ------------------------------------------------------------------
    # 3. Rename 4.2.3 heading
    # ------------------------------------------------------------------
    idx_423 = find_para(doc, "4.2.3")
    assert idx_423 is not None, "4.2.3 heading not found"
    rename_heading(doc.paragraphs[idx_423]._element,
                   "4.2.3\tStrategic direction & leadership", rev_ids)
    print("  Renamed 4.2.3 heading.")

    # ------------------------------------------------------------------
    # 4. Update 4.2 intro para 1 (three → four conditions)
    # ------------------------------------------------------------------
    idx_intro1 = find_para(doc, "Internal conditions describe a set of conditions")
    assert idx_intro1 is not None, "4.2 intro para 1 not found"
    replace_para(doc, idx_intro1,
                 [ins_para(INTRO_P1_NEW, "Normal", rev_ids)], rev_ids)
    print("  Updated 4.2 intro para 1.")

    # ------------------------------------------------------------------
    # 5. Insert causal-interaction + portfolio paragraphs after intro para 2
    #    (para 2 = "Managers respond to internal conditions in three distinct ways...")
    # ------------------------------------------------------------------
    idx_intro2 = find_para(doc,
        "Managers respond to internal conditions in three distinct ways")
    assert idx_intro2 is not None, "4.2 intro para 2 not found"
    insert_after(doc.paragraphs[idx_intro2]._element, [
        ins_para(INTRO_CAUSAL,    "Normal", rev_ids),
        ins_para(INTRO_PORTFOLIO, "Normal", rev_ids),
    ])
    print("  Inserted causal + portfolio intro paragraphs.")

    # ------------------------------------------------------------------
    # 6. Replace combined 4.2.1 response para + education-alone para
    #    with: #### Educating heading + edu para + #### Experimenting
    #           heading + exp para + navigate para
    # ------------------------------------------------------------------
    idx_combined = find_para(doc,
        "In response to these barriers, the data shows that experimenting with AI")
    idx_edu_alone = find_para(doc,
        "However, education alone is insufficient to sustain change")
    assert idx_combined is not None, "Combined 4.2.1 response para not found"
    assert idx_edu_alone is not None, "Education-alone para not found"

    # Grab references BEFORE modifying the document
    edu_alone_elem = doc.paragraphs[idx_edu_alone]._element

    # Delete old joint-definition tracked insertion — grab reference now
    idx_joint_def = find_para(doc,
        "Experimenting with AI refers to running small-scale trials")
    joint_def_elem = doc.paragraphs[idx_joint_def]._element if idx_joint_def is not None else None

    new_421 = [
        ins_para(EDU_HEADING_TEXT, "Heading4", rev_ids),
        ins_para(EDU_PARA,         "Normal",   rev_ids),
        ins_para(EXP_HEADING_TEXT, "Heading4", rev_ids),
        ins_para(EXP_PARA,         "Normal",   rev_ids),
        ins_para(NAVIGATE_PARA,    "Normal",   rev_ids),
    ]
    replace_para(doc, idx_combined, new_421, rev_ids)
    print("  Replaced combined 4.2.1 response paragraph.")

    # Delete education-alone paragraph (now embedded in EDU_PARA) — use pre-captured ref
    mark_para_deleted(edu_alone_elem, rev_ids)
    print("  Deleted education-alone paragraph (merged into edu sub-block).")

    # Remove old joint-definition tracked insertion — use pre-captured ref
    if joint_def_elem is not None:
        joint_def_elem.getparent().remove(joint_def_elem)
        print("  Removed old joint definition paragraph.")

    # ------------------------------------------------------------------
    # 7. Add portfolio framing sentence after 4.2.2 systems-thinking para
    # ------------------------------------------------------------------
    idx_sys = find_para(doc,
        "The problem is not data access per se, but the analytical capacity")
    assert idx_sys is not None, "Systems thinking para not found"
    insert_after(doc.paragraphs[idx_sys]._element,
                 [ins_para(PORTFOLIO_422, "Normal", rev_ids)])
    print("  Inserted 4.2.2 portfolio framing sentence.")

    # ------------------------------------------------------------------
    # 8. Replace 4.2.3 intro paragraph
    # ------------------------------------------------------------------
    idx_423_intro = find_para(doc,
        "The third cluster of internal conditions concerns leadership and governance")
    assert idx_423_intro is not None, "4.2.3 intro para not found"
    replace_para(doc, idx_423_intro,
                 [ins_para(SEC423_INTRO_NEW, "Normal", rev_ids)], rev_ids)
    print("  Replaced 4.2.3 intro paragraph.")

    # ------------------------------------------------------------------
    # 9. Insert "providing clarity" paragraph after 4.2.3 new intro
    # ------------------------------------------------------------------
    idx_423_intro2 = find_para(doc,
        "A third cluster of conditions concerns strategic direction and leadership")
    assert idx_423_intro2 is not None, "New 4.2.3 intro para not found"
    insert_after(doc.paragraphs[idx_423_intro2]._element,
                 [ins_para(CLARITY_PARA, "Normal", rev_ids)])
    print("  Inserted 'providing clarity' paragraph in 4.2.3.")

    # ------------------------------------------------------------------
    # 10. Insert 4.2.4 heading + new intro before bureaucracy paragraph
    # ------------------------------------------------------------------
    idx_bureau = find_para(doc,
        "Alongside these enabling behaviors, bureaucratic governance emerged")
    assert idx_bureau is not None, "Bureaucracy paragraph not found"

    # Insert heading + new 4.2.4 intro BEFORE the existing bureaucracy para
    insert_before(doc.paragraphs[idx_bureau]._element, [
        ins_para(SEC424_HEADING_TEXT, "Heading3", rev_ids),
        ins_para(SEC424_INTRO_NEW,    "Normal",   rev_ids),
    ])
    print("  Inserted 4.2.4 heading and intro.")

    # ------------------------------------------------------------------
    # 11. Replace closing paragraph (4.2.1 to 4.2.3 → 4.2.1 to 4.2.4)
    # ------------------------------------------------------------------
    idx_closing = find_para(doc,
        "Taken together, the internal conditions described in sections 4.2.1 to 4.2.3")
    assert idx_closing is not None, "Closing paragraph not found"
    replace_para(doc, idx_closing,
                 [ins_para(CLOSING_NEW, "Normal", rev_ids)], rev_ids)
    print("  Replaced closing paragraph.")

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------
    doc.save(str(OUT_PATH))
    print(f"\nSaved: {OUT_PATH}")


if __name__ == "__main__":
    main()
