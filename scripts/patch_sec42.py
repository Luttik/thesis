"""
Patch section 4.2 of the thesis draft with substantive content.

Replaces stub placeholder paragraphs in sections 4.2.1, 4.2.2, and 4.2.3
with full prose and quotes from the QDPX data, using tracked-change XML so
the author can review and accept/reject each change.

Usage:
    cd C:\\workspace\\thesis
    python scripts/patch_sec42.py
"""

from __future__ import annotations

import copy
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"

AUTHOR = "Claude"
DATE = "2026-06-01T00:00:00Z"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _new_ins_id(doc: Document) -> int:
    """Return an unused revision ID for <w:ins>/<w:del> elements."""
    used = {
        int(el.get(qn("w:id"), 0))
        for el in doc.element.body.iter()
        if el.get(qn("w:id")) is not None
    }
    return max(used, default=0) + 1


def make_ins_para(doc: Document, text: str, style: str = "Normal",
                  rev_id: int = 1, block_quote: bool = False) -> etree._Element:
    """Return a new <w:p> whose run is wrapped in <w:ins>."""
    p = etree.SubElement(etree.Element("dummy"), qn("w:p"))

    # Only add pPr when we need non-default formatting
    needs_pPr = block_quote or (style and style.lower() not in ("normal", ""))
    if needs_pPr:
        pPr = etree.SubElement(p, qn("w:pPr"))
        if style and style.lower() not in ("normal", ""):
            pStyle = etree.SubElement(pPr, qn("w:pStyle"))
            pStyle.set(qn("w:val"), style)
        if block_quote:
            ind = etree.SubElement(pPr, qn("w:ind"))
            ind.set(qn("w:left"), "720")

    ins = etree.SubElement(p, qn("w:ins"))
    ins.set(qn("w:id"), str(rev_id))
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), DATE)

    r = etree.SubElement(ins, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text

    # detach from dummy parent and return bare element
    p.getparent().remove(p) if p.getparent() is not None else None
    return p


def make_del_para(doc: Document, para_elem: etree._Element,
                  rev_id: int = 1) -> etree._Element:
    """Return a copy of para_elem with its runs converted to <w:del> runs."""
    p = copy.deepcopy(para_elem)
    for r in p.findall(f".//{qn('w:r')}"):
        parent = r.getparent()
        idx = list(parent).index(r)
        d = etree.Element(qn("w:del"))
        d.set(qn("w:id"), str(rev_id))
        d.set(qn("w:author"), AUTHOR)
        d.set(qn("w:date"), DATE)
        # convert w:t → w:delText
        for t in r.findall(qn("w:t")):
            dt = etree.SubElement(r, qn("w:delText"))
            dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            dt.text = t.text or ""
            r.remove(t)
        d.append(r)
        parent.insert(idx, d)
        parent.remove(r) if r.getparent() is parent else None
    return p


def insert_paras_after(anchor: etree._Element, new_paras: list[etree._Element]) -> None:
    """Insert new_paras immediately after anchor in the body."""
    parent = anchor.getparent()
    idx = list(parent).index(anchor) + 1
    for i, p in enumerate(new_paras):
        parent.insert(idx + i, p)


def para_text(p) -> str:
    return "".join(r.text or "" for r in p.runs)


# ---------------------------------------------------------------------------
# Section content
# ---------------------------------------------------------------------------

# Each entry: (text, is_block_quote)
SEC42_INTRO: list[tuple[str, bool]] = [
    (
        "Where section 4.1 described the external conditions that prompt marketing managers "
        "to engage with agentic AI, this section turns to the internal conditions that "
        "determine whether such engagement can be translated into value. Three interlocking "
        "themes recur throughout the data: the capacity to change, encompassing knowledge, "
        "culture, and organizational agility; the availability of enabling resources such "
        "as data infrastructure and technical tools; and the quality of leadership and "
        "governance. Where these internal conditions are favorable, managers leverage them; "
        "where they constrain progress, managers must navigate or actively reshape them.",
        False,
    ),
]

SEC421_PARAS: list[tuple[str, bool]] = [
    (
        "The most consistently identified barrier to adopting agentic AI is not technical "
        "but organizational: the capacity to change. This encompasses the level of AI "
        "understanding within the organization, employees’ openness to new ways of "
        "working, and the organizational ability to act quickly under uncertainty.",
        False,
    ),
    (
        "Resistance emerged as the dominant theme across the data (28 code applications, "
        "14 documents). Interviewee 12 described organizational change as inherently "
        "difficult:",
        False,
    ),
    (
        "“Change is frigging hard. It’s hard for us as individuals. Generally speaking, "
        "individuals don’t like to change, particularly if they’ve been doing something "
        "a certain way, and it works well.”",
        True,
    ),
    (
        "This resistance intensifies at the organizational level. Interviewee 12 continued:",
        False,
    ),
    (
        "“That challenge gets multiplied when you start talking about an organization "
        "… you’ve got all multiple people, you’ve got politics, existing incentives, "
        "existing processes.”",
        True,
    ),
    (
        "Alongside individual resistance, a lack of understanding of AI compounds the "
        "challenge. Interviewee 1 noted not just an absence of factual knowledge but a "
        "deeper cognitive barrier: limited understanding of how processes and systems "
        "connect. Interviewee 13 echoed this in practical terms: “The culprit is really "
        "still they don’t understand what the issues are … they don’t understand what "
        "processes they are doing right now that take time and money.” This inability to "
        "think in systems — to see how data, tools, and processes interact — was "
        "identified as a recurring constraint across organizations.",
        False,
    ),
    (
        "A related challenge is the tendency toward analysis paralysis. Interviewee 12 "
        "was direct: “You have to act. You can’t just be in a mode of studying this or "
        "writing memos about it, or thinking about planning about it.” Interviewee 4 "
        "observed a similar dynamic: “We spend so much time debating which shot we have "
        "to take that that is already some cost.”",
        False,
    ),
    (
        "In response to these barriers, the data shows that experimenting with AI "
        "(37 code applications, 17 documents) and educating and training "
        "(28 code applications, 14 documents) are the two dominant mechanisms through "
        "which managers build organizational capacity to change. Rather than waiting for "
        "perfect conditions, effective managers create structured opportunities to learn "
        "through doing. Interviewee 6 summarized this orientation: “I’d rather have "
        "something live and test and improve than have two years down the road and then "
        "finally have something live.” Interviewee 13 described enabling this at the team "
        "level: “Provide freedom to experiment and don’t set this atmosphere of fear "
        "that AI is going to replace you.”",
        False,
    ),
    (
        "However, education alone is insufficient to sustain change. Interviewee 17 "
        "observed a common pattern encountered in her training work: “After two or three "
        "trainings companies basically stop, because then they go back to what they started "
        "with.” Sustained capacity-building requires not just initial education but ongoing "
        "structural embedding of learning.",
        False,
    ),
]

SEC422_PARAS: list[tuple[str, bool]] = [
    (
        "A second cluster of internal conditions concerns the availability of enabling "
        "resources: the data, tools, and technical infrastructure that make AI "
        "implementation feasible. Having data and tools available and accessible emerged "
        "as a critical enabler (21 code applications, 12 documents), while lacking solid "
        "data infrastructure was a significant barrier (8 code applications, 6 documents).",
        False,
    ),
    (
        "Data infrastructure emerged as the foundational prerequisite. Interviewee 2 "
        "described his organization’s situation plainly: “The infrastructure is not yet "
        "in place, and the work on it is really big.” He elaborated that this is not a "
        "failure of will but of competing priorities: every improvement to the underlying "
        "data platform competes for attention with immediate operational demands. Without a "
        "robust and accessible data layer, AI systems cannot be reliably built or deployed.",
        False,
    ),
    (
        "A deeper resource constraint is the absence of systems thinking — the ability "
        "to understand how data flows connect across tools, departments, and processes. "
        "Interviewee 13 described this as the root cause behind many apparent AI failures:",
        False,
    ),
    (
        "“The stupidest example is … people want to understand their marketing data "
        "better. And it’s really just like you have 20 different sources, you don’t "
        "know how to connect them, you don’t know how to put this all together. … "
        "The culprit is really still they don’t understand what the issues are.”",
        True,
    ),
    (
        "The problem is not data access per se, but the analytical capacity to know what "
        "data is needed, in what form, and for what purpose. Interviewee 1 noted that this "
        "type of deep process knowledge — a feel for system logic — is precisely what "
        "is hardest to augment with AI, because it requires a level of structured thinking "
        "that many teams have not yet developed.",
        False,
    ),
    (
        "Where these foundational resources are available, they become powerful enablers. "
        "Interviewee 4 described his own setup: “I connected an MCP to their Google "
        "Analytics and Google Search Console. So I have direct access to that data source "
        "in my cloud, and I can run cron jobs and have an agent run reports.” "
        "Interviewee 5 similarly observed that once the back-end infrastructure is in "
        "place, rapid experimentation becomes possible: “The back-end stuff needs to be "
        "extremely robust, and then … you can really build quickly.” Both cases "
        "illustrate how data access and tool availability translate directly into the "
        "ability to deploy agentic workflows.",
        False,
    ),
    (
        "Budget appeared as a secondary resource constraint (2 code applications). "
        "Interviewee 9 captured the tension: “The passion is there, the need is there, "
        "the want is there, the brains are there. It’s time, money, and resources.” "
        "While budget is not the primary barrier in the data, it becomes a bottleneck "
        "when resource availability determines whether AI initiatives can scale from "
        "pilot to production.",
        False,
    ),
]

SEC423_PARAS: list[tuple[str, bool]] = [
    (
        "The third cluster of internal conditions concerns leadership and governance — "
        "the organizational structures and behaviors that either enable or impede AI "
        "adoption. The data surfaces three related dynamics: the role of individual AI "
        "champions, the importance of leadership backing, and the friction created by "
        "bureaucratic governance.",
        False,
    ),
    (
        "Individual AI champions emerged as a disproportionately influential factor "
        "(5 code applications but high narrative weight across the corpus). An AI champion "
        "is a person who, often without formal mandate, drives AI initiatives forward "
        "through personal initiative, visible results, and internal advocacy. Interviewee 9 "
        "described a vivid example from her organization:",
        False,
    ),
    (
        "“Affinity Hub was created originally by … our global creative director that "
        "dabbles in dev, and he created it, and it’s like blown up, and so now there’s "
        "millions of dollars of funding.”",
        True,
    ),
    (
        "A single individual’s initiative, when it produces visible results, can unlock "
        "significant organizational investment. Interviewee 6 embodied this champion role "
        "in her own organization — consistently pushing AI initiatives forward, absorbing "
        "resistance, and creating space for others to experiment.",
        False,
    ),
    (
        "Leadership backing amplifies the champion effect (14 code applications, "
        "10 documents). Without visible endorsement from leadership, AI initiatives tend "
        "to stall at the pilot stage. Interviewee 4 noted that the current moment requires "
        "leaders to help employees make sense of change: “People are drinking from a fire "
        "hose. They need to have that leadership that tells them their importance, they’re "
        "part of the team, and their work matters.” Interviewee 17 identified three "
        "conditions that must align for AI adoption to succeed: “Leadership wants us to "
        "work with AI. But can employees do that too? Do they do it and do they want to? "
        "Those are basically three things. And when those three things align with each "
        "other, you get a fantastic AI implementation.”",
        False,
    ),
    (
        "The process of bringing people along — change management at the team and "
        "organizational level — was one of the most richly coded themes in this block "
        "(24 code applications, 13 documents). Interviewee 6 captured the essential "
        "managerial move: “This is the vision. This is where we’re going. Help me get "
        "there.” Interviewee 13 described a practical approach: using lighthouse projects "
        "as demonstration cases that make the value of AI tangible to skeptical colleagues. "
        "“Show them the first one that can be automated or made a lot quicker … then, "
        "well, then often they understand much better what the use cases would be.”",
        False,
    ),
    (
        "Alongside these enabling behaviors, bureaucratic governance emerged as a "
        "significant opposing force (21 code applications, 13 documents). Governance "
        "friction appeared in several forms: IT security and procurement processes slowing "
        "deployment, legal and compliance concerns blocking tool use, and organizational "
        "politics preventing experimentation. Interviewee 13 observed a pattern in large "
        "enterprises:",
        False,
    ),
    (
        "“They are not allowed to use agentic AI a lot … there’s this weird distinction "
        "people make. It’s totally fine to buy any kind of tool you don’t know anything "
        "about and put it into the system. But agentic AI? That’s completely out of "
        "range.”",
        True,
    ),
    (
        "Interviewee 6 noted that governance structures, while legitimate, can prevent "
        "organizations from learning: “If we would’ve been stopped at the beginning — "
        "‘You cannot use this system because you’re already on this system’ — I think "
        "that is something that’s really holding people back.”",
        False,
    ),
    (
        "Taken together, the internal conditions described in sections 4.2.1 to 4.2.3 "
        "represent the organizational substrate on which AI value creation depends. How "
        "managers navigate these conditions — leveraging favorable ones and working "
        "around or reshaping unfavorable ones — directly shapes the adoption processes "
        "described in section 4.3.",
        False,
    ),
]

# ---------------------------------------------------------------------------
# Identify stub paragraphs
# ---------------------------------------------------------------------------

# Text fragments that identify the stub placeholders to delete
STUB_MARKERS = [
    "Describe elements like knowledge, culture",
    "Describe resources like data availability",
    "Describe elements like restrictive governance",
]

# Text fragment identifying the existing 4.2 intro to keep (but we replace it)
INTRO_MARKER = "Internal conditions describe a set of conditions"

# Heading text for the three subsections
H421_MARKER = "4.2.1"
H422_MARKER = "4.2.2"
H423_MARKER = "4.2.3"


def find_para_index(doc: Document, marker: str) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if p.text and marker in p.text:
            return i
    return None


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print(f"Opening: {DOCX_PATH}")
    doc = Document(str(DOCX_PATH))
    body = doc.element.body
    paras = list(body)  # direct children (paragraphs + tables)

    rev_id = _new_ins_id(doc)

    def next_id() -> int:
        nonlocal rev_id
        val = rev_id
        rev_id += 1
        return val

    # ------------------------------------------------------------------
    # 1. Replace the existing 4.2 intro paragraph
    # ------------------------------------------------------------------
    intro_idx = find_para_index(doc, INTRO_MARKER)
    if intro_idx is not None:
        intro_para_elem = doc.paragraphs[intro_idx]._element
        # Mark old intro as deleted
        del_intro = make_del_para(doc, intro_para_elem, rev_id=next_id())
        body_paras = [c for c in body if c.tag == qn("w:p") or c.tag == qn("w:tbl")]
        body_idx = list(body).index(intro_para_elem)
        # Insert del before the original, then new intro paras, then remove original
        body.insert(body_idx, del_intro)
        offset = 1
        for text, bq in SEC42_INTRO:
            p = make_ins_para(doc, text, style="Normal", rev_id=next_id(),
                              block_quote=bq)
            body.insert(body_idx + offset, p)
            offset += 1
        body.remove(intro_para_elem)
        print("  Replaced 4.2 intro paragraph.")
    else:
        print("  WARNING: 4.2 intro paragraph not found; skipping intro replacement.")

    # ------------------------------------------------------------------
    # 2. Replace stub paragraphs under each subsection heading
    # ------------------------------------------------------------------
    section_data = [
        (H421_MARKER, "4.2.1 Capacity to change", SEC421_PARAS),
        (H422_MARKER, "4.2.2 Available resources", SEC422_PARAS),
        (H423_MARKER, "4.2.3 Leadership", SEC423_PARAS),
    ]

    for heading_marker, label, content_paras in section_data:
        heading_idx = find_para_index(doc, heading_marker)
        if heading_idx is None:
            print(f"  WARNING: heading for {label} not found.")
            continue

        heading_elem = doc.paragraphs[heading_idx]._element

        # Find and delete the stub paragraph immediately after the heading
        # (look ahead up to 5 paragraphs for a stub marker)
        stub_elem = None
        for j in range(1, 6):
            candidate_idx = heading_idx + j
            if candidate_idx >= len(doc.paragraphs):
                break
            cand = doc.paragraphs[candidate_idx]
            if any(m in cand.text for m in STUB_MARKERS):
                stub_elem = cand._element
                break

        # Insert new content after heading
        insert_after = heading_elem
        if stub_elem is not None:
            del_stub = make_del_para(doc, stub_elem, rev_id=next_id())
            stub_body_idx = list(body).index(stub_elem)
            body.insert(stub_body_idx, del_stub)
            body.remove(stub_elem)
            insert_after = del_stub
            print(f"  Marked stub as deleted for {label}.")
        else:
            print(f"  No stub found for {label}; inserting after heading.")

        # Insert new paragraphs after heading (or after del_stub)
        insert_idx = list(body).index(insert_after) + 1
        for i, (text, bq) in enumerate(content_paras):
            p = make_ins_para(doc, text, style="Normal", rev_id=next_id(),
                              block_quote=bq)
            body.insert(insert_idx + i, p)

        print(f"  Inserted {len(content_paras)} paragraphs for {label}.")

    # ------------------------------------------------------------------
    # 3. Save in-place
    # ------------------------------------------------------------------
    doc.save(str(DOCX_PATH))
    print(f"\nSaved in-place: {DOCX_PATH}")


if __name__ == "__main__":
    main()
