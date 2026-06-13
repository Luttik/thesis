# -*- coding: utf-8 -*-
"""Make the Executive Summary fit one page: set its paragraphs to single line
spacing (they were inheriting the document's double spacing) and lightly condense
the prose (~35 words) as tracked changes (author 'Claude')."""
from __future__ import annotations
import copy, shutil, sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

DOCX = Path(__file__).resolve().parents[1] / "Thesis Draft - Daan Luttik - MBA - copy from 2026-06-13.docx"
BACKUP = DOCX.with_suffix(".execsum2-backup.docx")
AUTHOR = "Claude"; DATE = "2026-06-13T19:00:00Z"
XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"
APO = "’"

def _used(doc): return {int(el.get(qn("w:id"), 0)) for el in doc.element.body.iter() if el.get(qn("w:id")) is not None}
def _nid(u): n = max(u, default=0) + 1; u.add(n); return n
def full_text(p): return "".join(t.text or "" for t in p.iter(qn("w:t")))
def find_elem(doc, m):
    for p in doc.paragraphs:
        if m in full_text(p._element): return p._element
    return None
def _mk(rPr, s):
    r = etree.Element(qn("w:r"))
    if rPr is not None: r.append(copy.deepcopy(rPr))
    t = etree.SubElement(r, qn("w:t")); t.set(XMLSPACE, "preserve"); t.text = s
    return r
def _ins(s, u):
    el = etree.Element(qn("w:ins")); el.set(qn("w:id"), str(_nid(u)))
    el.set(qn("w:author"), AUTHOR); el.set(qn("w:date"), DATE)
    r = etree.SubElement(el, qn("w:r")); t = etree.SubElement(r, qn("w:t"))
    t.set(XMLSPACE, "preserve"); t.text = s
    return el
def replace_in_inserted(p, old, new):
    """Condense text that lives inside an unaccepted <w:ins> by editing the run
    text in place (the whole exec summary is one pending Claude insertion)."""
    runs = [r for r in p.iter(qn("w:r")) if r.find(qn("w:t")) is not None]
    segs = []; total = ""
    for r in runs:
        t = r.find(qn("w:t")); txt = t.text or ""
        segs.append((t, txt)); total += txt
    start = total.find(old)
    if start < 0: return False
    end = start + len(old); pos = 0; placed = False
    for t, txt in segs:
        rs, re_ = pos, pos + len(txt); pos = re_
        if not txt or re_ <= start or rs >= end: continue
        ls = max(start, rs) - rs; le = min(end, re_) - rs
        t.text = txt[:ls] + (new if not placed else "") + txt[le:]
        t.set(XMLSPACE, "preserve")
        placed = True
    return placed

EDITS = [
    ("A newer category, agentic AI",
     "can pursue goals autonomously across complex, multi-step tasks and so promises to perform the work itself",
     "pursues goals autonomously across complex, multi-step tasks, promising to perform the work itself"),
    ("The study applies constructivist",
     "Seventeen in-depth, semi-structured interviews with marketing managers and AI experts were analyzed "
     "iteratively through initial, focused, and theoretical coding, with theoretical sampling continued until saturation.",
     "Seventeen in-depth interviews with marketing managers and AI experts were analyzed through initial, "
     "focused, and theoretical coding, with theoretical sampling continued to saturation."),
    ("Managers should treat adoption",
     "as a change program rather than a procurement: build the enabling conditions and lead the change; "
     "consciously choose a follow (parity) or a differentiate posture; start from the workflow and engineer "
     f"the configuration; govern AI with AI, calibrated to the cost of error; resource and source capability "
     f"strategically; and look outward to consumers{APO} own agents and to the junior-talent pipeline.",
     "as a change program, not a procurement: build the enabling conditions and lead the change; choose a "
     "follow or differentiate posture; start from the workflow and engineer the configuration; govern AI with "
     f"AI, calibrated to the cost of error; resource and source capability strategically; and look outward to "
     f"consumers{APO} own agents and the junior-talent pipeline."),
    ("This is one of the first grounded",
     "and offers a process model, together with an analytical vocabulary (reshaping, leveraging, and "
     "navigating; the harness; and value as a managed portfolio), for understanding the managerial work that "
     "turns agentic AI into value.",
     "and offers a process model and an analytical vocabulary for the managerial work that turns agentic AI into value."),
]

def main():
    check = "--check" in sys.argv
    doc = Document(str(DOCX)); u = _used(doc)
    # locate exec-summary range
    paras = doc.paragraphs
    start = end = None
    for i, p in enumerate(paras):
        t = full_text(p._element).strip()
        if t == "Executive Summary": start = i
        elif start is not None and t.startswith("1.") and "Introduction" in t: end = i; break
    problems = []
    if start is None or end is None: problems.append("exec-summary range not found")
    for marker, old, new in EDITS:
        el = find_elem(doc, marker)
        if el is None: problems.append(("NO PARA", marker[:30]))
        elif old not in full_text(el): problems.append(("NO FRAG", marker[:24]))
    print(f"exec range paras {start}..{end}; {len(EDITS)} condense edits")
    for x in problems: print("  PROBLEM:", x)
    if check:
        print("OK." if not problems else "FIX FIRST."); return
    if problems: print("NOT SAVED."); return
    # set single line spacing on the exec-summary body paragraphs
    n_spaced = 0
    for p in paras[start + 1:end]:
        p.paragraph_format.line_spacing = 1.0
        n_spaced += 1
    # condense the pending insertion in place
    for marker, old, new in EDITS:
        if not replace_in_inserted(find_elem(doc, marker), old, new):
            print("  WARN unmatched:", marker[:30])
    shutil.copy(DOCX, BACKUP); doc.save(str(DOCX))
    print(f"Set single spacing on {n_spaced} paragraphs; applied {len(EDITS)} trims. Backup: {BACKUP.name}")

if __name__ == "__main__":
    main()
