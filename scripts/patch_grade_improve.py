# -*- coding: utf-8 -*-
"""
Grade-improvement pass (tracked changes, author "Claude").

Produces two outputs from the same base file:
  python patch_grade_improve.py focused        -> ...- FOCUSED.docx
  python patch_grade_improve.py comprehensive  -> ...- COMPREHENSIVE.docx

ALWAYS applied: APA reference-list fixes + focused content additions grounded
in repository sources (Wharton/GBK 2025 adoption stat, Brynjolfsson 2025 in Ch2,
agentic managerial point in 2.2, LLM spell-out).
COMPREHENSIVE adds: practitioner perspective in 2.2 (Bughin 2024), a reader-
friendly ML mechanism sentence, and a social-desirability-bias note in 5.3.
"""
from __future__ import annotations
import copy, sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
SRC  = ROOT / "Thesis Draft - Daan Luttik - MBA - copy from 2026-06-14.docx"
AUTHOR = "Claude"
DATE   = "2026-06-15T12:00:00Z"
XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"

MODE = sys.argv[1] if len(sys.argv) > 1 else "focused"
assert MODE in ("focused", "comprehensive")
OUT = ROOT / ("Thesis Draft - Daan Luttik - MBA - %s.docx" % MODE.upper())

doc = Document(str(SRC))
BODY = doc.element.body

# --------------------------------------------------------------------------- #
# id + lookup helpers
# --------------------------------------------------------------------------- #
USED = {int(el.get(qn("w:id"), 0)) for el in BODY.iter() if el.get(qn("w:id")) is not None}
def nid():
    n = max(USED, default=0) + 1; USED.add(n); return n

def ft(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t")))

def find_p(needle, nth=0):
    hits = [p._element for p in doc.paragraphs if needle in ft(p._element)]
    if not hits:
        raise RuntimeError("NOT FOUND: %r" % needle)
    return hits[nth]

def run_text(r):
    return "".join(t.text or "" for t in r.findall(qn("w:t")))

def _ins_wrap(child_runs):
    ins = etree.Element(qn("w:ins"))
    ins.set(qn("w:id"), str(nid())); ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), DATE)
    for r in child_runs:
        ins.append(r)
    return ins

def _mk_run(text, rpr_template=None, italic=False, highlight=None):
    r = etree.Element(qn("w:r"))
    rpr = etree.SubElement(r, qn("w:rPr"))
    if rpr_template is not None:
        for ch in rpr_template:
            # copy non-conflicting formatting (lang etc.); skip i/highlight we set explicitly
            if ch.tag in (qn("w:i"), qn("w:highlight")):
                continue
            rpr.append(copy.deepcopy(ch))
    if highlight:
        etree.SubElement(rpr, qn("w:highlight")).set(qn("w:val"), highlight)
    if italic:
        etree.SubElement(rpr, qn("w:i"))
    if len(rpr) == 0:
        r.remove(rpr)
    t = etree.SubElement(r, qn("w:t")); t.set(XMLSPACE, "preserve"); t.text = text
    return r

# --------------------------------------------------------------------------- #
# tracked operations
# --------------------------------------------------------------------------- #
def tracked_replace_run(r, new_text):
    """Replace one run's text with a tracked del(old)+ins(new), preserving rPr."""
    parent = r.getparent(); pos = list(parent).index(r)
    # deletion
    d = etree.Element(qn("w:del"))
    d.set(qn("w:id"), str(nid())); d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), DATE)
    rc = copy.deepcopy(r)
    for t in rc.findall(qn("w:t")):
        dt = etree.Element(qn("w:delText")); dt.set(XMLSPACE, "preserve"); dt.text = t.text or ""
        t.getparent().replace(t, dt)
    d.append(rc)
    # insertion (clone rPr, new text)
    rn = copy.deepcopy(r)
    for t in rn.findall(qn("w:t")):
        rn.remove(t)
    t = etree.SubElement(rn, qn("w:t")); t.set(XMLSPACE, "preserve"); t.text = new_text
    ins = _ins_wrap([rn])
    parent.remove(r)
    parent.insert(pos, ins)
    parent.insert(pos, d)

def replace_exact_run(old, new, require_italic=False):
    cnt = 0
    for p in doc.paragraphs:
        for r in list(p._element.findall(qn("w:r"))):
            if run_text(r) == old:
                if require_italic:
                    rpr = r.find(qn("w:rPr"))
                    if rpr is None or rpr.find(qn("w:i")) is None:
                        continue
                tracked_replace_run(r, new); cnt += 1
    if cnt == 0:
        print("  WARN exact-run not found:", repr(old))
    return cnt

def replace_substr_in_run(run_marker, old_sub, new_sub):
    for p in doc.paragraphs:
        for r in list(p._element.findall(qn("w:r"))):
            txt = run_text(r)
            if run_marker in txt and old_sub in txt:
                tracked_replace_run(r, txt.replace(old_sub, new_sub, 1)); return True
    print("  WARN substr run not found:", repr(run_marker)); return False

def insert_inline_after(needle_run_marker, new_text):
    """Split the run containing the marker; insert a tracked run after the marker."""
    for p in doc.paragraphs:
        for r in list(p._element.findall(qn("w:r"))):
            txt = run_text(r)
            if needle_run_marker in txt:
                idx = txt.find(needle_run_marker) + len(needle_run_marker)
                head, tail = txt[:idx], txt[idx:]
                rpr = r.find(qn("w:rPr"))
                parent = r.getparent(); pos = list(parent).index(r)
                head_r = _mk_run(head, rpr)
                ins_r  = _ins_wrap([_mk_run(new_text, rpr)])
                tail_r = _mk_run(tail, rpr)
                parent.remove(r)
                for el in (tail_r, ins_r, head_r):
                    parent.insert(pos, el)
                return True
    print("  WARN inline marker not found:", repr(needle_run_marker)); return False

def insert_after_run_text(para_needle, exact_run, new_text):
    """Insert a tracked run immediately after the run whose full text == exact_run,
    within the paragraph containing para_needle."""
    p = find_p(para_needle)
    for r in list(p.findall(qn("w:r"))):
        if run_text(r) == exact_run:
            rpr = r.find(qn("w:rPr"))
            r.addnext(_ins_wrap([_mk_run(new_text, rpr)]))
            return True
    print("  WARN exact run %r not found in para" % exact_run); return False

def append_sentence(needle, new_text):
    """Append a tracked run as the last run of the paragraph containing needle."""
    p = find_p(needle)
    runs = p.findall(qn("w:r"))
    rpr = runs[-1].find(qn("w:rPr")) if runs else None
    ins_r = _ins_wrap([_mk_run(new_text, rpr)])
    # place after the last run (before any trailing bookmarks/paragraph mark stays in pPr)
    last = runs[-1]
    last.addnext(ins_r)
    return True

def insert_para_after(anchor_p, segments, ref=False):
    """Insert a new tracked paragraph after anchor_p. segments: list of (text, italic)."""
    new_p = etree.Element(qn("w:p"))
    # clone pPr from anchor for spacing/style consistency
    src_ppr = anchor_p.find(qn("w:pPr"))
    pPr = copy.deepcopy(src_ppr) if src_ppr is not None else etree.Element(qn("w:pPr"))
    new_p.append(pPr)
    rpr = pPr.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.SubElement(pPr, qn("w:rPr"))
    # mark the paragraph-mark itself as inserted
    insm = etree.Element(qn("w:ins"))
    insm.set(qn("w:id"), str(nid())); insm.set(qn("w:author"), AUTHOR); insm.set(qn("w:date"), DATE)
    rpr.insert(0, insm)
    runs = [_mk_run(txt, None, italic=ital, highlight=("white" if ref else None)) for txt, ital in segments]
    new_p.append(_ins_wrap(runs))
    anchor_p.addnext(new_p)
    return new_p

def rebuild_ref(needle, segments):
    """Replace all runs of a reference paragraph with tracked del(old)+ins(new white runs)."""
    p = find_p(needle)
    for r in list(p.findall(qn("w:r"))):
        # tracked delete in place
        pos = list(p).index(r)
        d = etree.Element(qn("w:del"))
        d.set(qn("w:id"), str(nid())); d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), DATE)
        rc = copy.deepcopy(r)
        for t in rc.findall(qn("w:t")):
            dt = etree.Element(qn("w:delText")); dt.set(XMLSPACE, "preserve"); dt.text = t.text or ""
            t.getparent().replace(t, dt)
        d.append(rc)
        p.remove(r); p.insert(pos, d)
    runs = [_mk_run(txt, None, italic=ital, highlight="white") for txt, ital in segments]
    p.append(_ins_wrap(runs))

# =========================================================================== #
# 1) APA REFERENCE FIXES  (both documents)
# =========================================================================== #
print("[%s] APA reference fixes" % MODE)
# journal-title casing (italic isolated runs -> title case)
CASING = {
    "Harvard business review": "Harvard Business Review",
    "Information systems frontiers": "Information Systems Frontiers",
    "Journal of consumer behaviour": "Journal of Consumer Behaviour",
    "Industrial marketing management": "Industrial Marketing Management",
    "European journal of operational research": "European Journal of Operational Research",
    "Data mining and knowledge discovery": "Data Mining and Knowledge Discovery",
    "Advances in neural information processing systems": "Advances in Neural Information Processing Systems",
    "Journal of marketing research": "Journal of Marketing Research",
    "Journal of marketing": "Journal of Marketing",
    "Education for information": "Education for Information",
    "The innovation": "The Innovation",
    "Science advances": "Science Advances",
    "nature": "Nature",
}
for old, new in CASING.items():
    n = replace_exact_run(old, new, require_italic=True)
    print("   casing %-50s x%d" % (old, n))

# title acronym fixes
replace_substr_in_run("Bert: Pre-training", "Bert:", "BERT:")
replace_substr_in_run("knowledge-intensive nlp tasks", "knowledge-intensive nlp", "knowledge-intensive NLP")

# whole-entry rebuilds
rebuild_ref("Charmaz, K. (2014). Constructing grounded theory", [
    ("Charmaz, K. (2014). ", False),
    ("Constructing grounded theory", True),
    (" (2nd ed.). Sage.", False),
])
rebuild_ref("Woodside, A. G., Golfetto, F., & Gibbert, M. (2008). Customer value", [
    ("Woodside, A. G., Golfetto, F., & Gibbert, M. (2008). Customer value: Theory, research, and practice. In ", False),
    ("Creating and managing superior customer value", True),
    (" (Advances in Business Marketing and Purchasing, Vol. 14, pp. 3–25). Emerald Group Publishing. https://doi.org/10.1016/S1069-0964(08)14001-7", False),
])
rebuild_ref("Radford, A., Narasimhan, K., Salimans, T., & Sutskever, I. (2018). Improving language", [
    ("Radford, A., Narasimhan, K., Salimans, T., & Sutskever, I. (2018). ", False),
    ("Improving language understanding by generative pre-training", True),
    (". OpenAI.", False),
])

# web-source entries (#246): drop redundant site-name where author == site; tidy
replace_substr_in_run(". Anthropic. ", ". Anthropic. ", ". ")               # Anthropic (2024)
replace_substr_in_run("Anthropic Engineering.  ", "Anthropic Engineering.  ", "Anthropic Engineering. ")  # (2025) double space
replace_substr_in_run(". Claude. https://claude.com/blog/skills", ". Claude. ", ". ")  # Claude (2025)
# ChatGPT (2026): move period outside italic title + drop redundant site
replace_exact_run("Skills in ChatGPT.", "Skills in ChatGPT", require_italic=True)
replace_substr_in_run(" ChatGPT. ", " ChatGPT. ", ". ")

# =========================================================================== #
# 2) FOCUSED CONTENT ADDITIONS (both documents)
# =========================================================================== #
print("[%s] focused content additions" % MODE)

# F1 - Introduction: adoption figure that makes the gap matter (#9/#14)
insert_inline_after(
    "unlock the large productivity gains mentioned by Mayer et al. (2025). ",
    "Adoption is no longer hypothetical: 46% of enterprise decision-makers report using generative AI "
    "daily and over 80% at least weekly, with marketing content creation among the most common use cases "
    "(Wharton Human-AI Research & GBK Collective, 2025). The gap therefore matters in practice, as managers "
    "are already investing in and deploying these systems faster than research can explain how they create value. ",
)

# F2 - 2.1.2: empirical grounding (Brynjolfsson) + theorized-vs-found signal (#221/#73)
vaid = find_p("do not describe what was needed to deploy these use cases")
insert_para_after(vaid, [(
    "Where Vaid et al. (2025) stop short of measuring outcomes, large-scale field evidence is beginning to "
    "appear outside marketing. Studying 5,172 customer-support agents, Brynjolfsson et al. (2025) found that "
    "access to a generative AI assistant raised productivity by roughly 14% on average, with the largest gains "
    "accruing to less-experienced workers. Such findings signal that the productivity potential of generative "
    "AI is increasingly measured rather than only theorized, although evidence from within marketing, and for "
    "agentic systems specifically, remains scarce.", False)])

# F3 - 2.2: the agentic managerial point (#64/#65)
mor = find_p("simplified in communications with interviewees to two key points")
insert_para_after(mor, [(
    "This goal-directed autonomy is what makes agentic AI distinct for managers. Where an earlier AI tool "
    "optimized a narrowly defined task, a single agentic system can be pointed at a wide range of loosely "
    "specified goals and determine the intermediate steps itself (Acharya et al., 2025; Moralles et al., 2026). "
    "For the manager, this shifts the task from specifying the exact output to specifying the objective, the "
    "tools the system may use, and the guardrails within which it may act.", False)])

# F4 - spell out LLM on first body use (#45)
replace_substr_in_run("showed how LLMs could combine reasoning",
                      "how LLMs could", "how large language models (LLMs) could")

# new reference entry: Wharton / GBK (alphabetical: after Wahid, before Woodside)
wahid = find_p("Written by ChatGPT, illustrated by Midjourney")
insert_para_after(wahid, [
    ("Wharton Human-AI Research, & GBK Collective. (2025). ", False),
    ("Gen AI fast-tracks into the enterprise: Year three full report", True),
    (". The Wharton School, University of Pennsylvania.", False),
], ref=True)

# =========================================================================== #
# 3) COMPREHENSIVE-ONLY ADDITIONS
# =========================================================================== #
if MODE == "comprehensive":
    print("[comprehensive] extra additions")

    # C1 - 2.2 practitioner perspective (#71/#72)
    calls = find_p("how marketing managers create value with agentic systems beyond content generation")
    insert_para_after(calls, [(
        "Practitioner studies of enterprise adoption reach a similar conclusion and stress that the binding "
        "constraint is rarely the technology itself. Analysing how firms capture value from AI, Bughin (2024) "
        "reports that the organizations pulling ahead are distinguished less by their models than by leadership "
        "commitment, talent, and the surrounding processes, while the Wharton Human-AI Research and GBK "
        "Collective (2025) document rapid but uneven enterprise adoption. This study examines that pattern "
        "empirically within marketing.", False)])

    # Bughin (2024) reference (alphabetical: after Brynjolfsson, before Burkhardt)
    bryn = find_p("Brynjolfsson, E., Li, D., & Raymond, L. (2025)")
    insert_para_after(bryn, [
        ("Bughin, J. (2024). Inside the successful make-up of ‘AI-first’ organisations. ", False),
        ("Journal of AI, Robotics & Workplace Automation", True),
        (", ", False),
        ("3", True),
        ("(3), 211–219.", False),
    ], ref=True)

    # C2 - 2.1 reader-friendly ML mechanism (#35/#37)
    append_sentence("product recommendation systems in e-commerce",
        " For readers less familiar with these techniques, the common thread is that the model learns patterns "
        "from historical data to make a prediction, for example, which version of a web page a visitor is most "
        "likely to act on or which product a given customer is most likely to buy.")

    # C3 - 5.3 name social desirability bias (#233)
    insert_after_run_text("others might not have been so", "forthcoming",
        " (a form of social desirability bias, where respondents portray themselves or their organizations "
        "more favorably than is warranted)")

doc.save(str(OUT))
print("SAVED ->", OUT.name)
