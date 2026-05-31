"""Insert Discussion chapter content into the thesis docx."""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
TEMP_DOCX = ROOT / ".cache" / "thesis_temp.docx"
MARKDOWN_PATH = ROOT / "thesis" / "5. Discussion.md"
UNPACKED_CUSTOM = ROOT / ".cache" / "thesis_unpacked" / "customXML"


def repair_docx_if_needed(path: Path) -> None:
    """Restore missing customXml parts referenced by the docx package."""
    try:
        Document(str(path))
        return
    except KeyError as exc:
        if "customXml" not in str(exc):
            raise

    backup = path.with_suffix(".docx.bak")
    shutil.copy2(path, backup)
    print(f"Repairing corrupted docx (backup at {backup.name})")

    with zipfile.ZipFile(path, "r") as zin:
        names = set(zin.namelist())
        entries = {name: zin.read(name) for name in names}

    custom_files = [
        "customXml/item1.xml",
        "customXml/itemProps1.xml",
        "customXml/_rels/item1.xml.rels",
    ]
    for rel_path in custom_files:
        if rel_path not in entries:
            src = UNPACKED_CUSTOM / Path(rel_path).name
            if rel_path.endswith(".rels"):
                src = UNPACKED_CUSTOM / "_rels" / "item1.xml.rels"
            if src.exists():
                entries[rel_path] = src.read_bytes()
                print(f"  restored {rel_path}")

    tmp = path.with_suffix(".docx.repairing")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in entries.items():
            zout.writestr(name, data)
    tmp.replace(path)
    Document(str(path))
    print("Repair complete")


def load_source_docx() -> Path:
    repair_docx_if_needed(DOCX_PATH)
    try:
        Document(str(DOCX_PATH))
        return DOCX_PATH
    except Exception:
        print("Main docx still unreadable; using thesis_temp.docx as source")
        return TEMP_DOCX


def insert_paragraph_after(
    paragraph: Paragraph, text: str, style: str = "normal"
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    style_map = {
        "h1": "Heading 1",
        "h2": "Heading 2",
        "h3": "Heading 3",
        "h4": "Heading 4",
        "normal": "Normal",
        "body": "Normal",
    }
    style_name = style_map.get(style, "Normal")
    try:
        new_para.style = style_name
    except (KeyError, ValueError):
        new_para.style = "Normal"
    if text:
        _add_formatted_text(new_para, text)
    return new_para


def _add_formatted_text(paragraph: Paragraph, text: str) -> None:
    """Add text with **bold** and *italic* markers as runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)")
    for part in pattern.findall(text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)


def remove_following_paragraphs_until(
    start: Paragraph, stop_predicate
) -> Paragraph:
    """Remove paragraphs after start until stop_predicate(p) is True. Return stop paragraph."""
    nxt = start._element.getnext()
    stop_para = start
    while nxt is not None:
        para = Paragraph(nxt, start._parent)
        if stop_predicate(para):
            stop_para = para
            break
        to_remove = nxt
        nxt = nxt.getnext()
        to_remove.getparent().remove(to_remove)
    return stop_para


def parse_markdown_blocks(md_text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            continue  # skip file title; doc already has H1
        if line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("#### "):
            blocks.append(("h4", line[5:].strip()))
        else:
            blocks.append(("body", line.strip()))
    return blocks


def normalize_heading(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())


def patch_document(doc: Document, blocks: list[tuple[str, str]]) -> None:
    discussion_h1 = None
    conclusion_h1 = None
    for p in doc.paragraphs:
        t = normalize_heading(p.text)
        if "5." in p.text and "discussion" in t and p.style.name.startswith("Heading"):
            discussion_h1 = p
        if "6." in p.text and "conclusion" in t and p.style.name.startswith("Heading"):
            conclusion_h1 = p
            break

    if discussion_h1 is None or conclusion_h1 is None:
        raise RuntimeError("Could not locate Discussion or Conclusion headings")

    remove_following_paragraphs_until(
        discussion_h1,
        lambda para: para._element is conclusion_h1._element,
    )

    current = discussion_h1
    for kind, text in blocks:
        current = insert_paragraph_after(current, text, kind if kind != "body" else "normal")


def main() -> None:
    source = load_source_docx()
    md_text = MARKDOWN_PATH.read_text(encoding="utf-8")
    blocks = parse_markdown_blocks(md_text)

    doc = Document(str(source))
    patch_document(doc, blocks)
    doc.save(str(DOCX_PATH))
    print(f"Patched Discussion in {DOCX_PATH} ({len(blocks)} blocks)")


if __name__ == "__main__":
    main()
