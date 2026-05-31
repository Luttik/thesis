"""Patch thesis docx with trustworthiness and reflexivity sections."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"


REFLEXIVITY_PARAS = [
    (
        "The researcher conducting this study has a professional background in marketing, "
        "data, and applied AI within commercial organizations. That background provides "
        "familiarity with martech stacks, campaign operations, and the language practitioners "
        "use when discussing automation, governance, and return on investment. Familiarity "
        "can accelerate rapport in interviews and help the researcher recognize when "
        "participants refer to workflows, vendors, or constraints that might otherwise "
        "require extensive clarification."
    ),
    (
        "The same background also introduces interpretive risks. Prior experience may "
        "sensitize the researcher toward efficiency narratives, tooling debates, or "
        "implementation patterns that resemble settings already known, while under-attending "
        "to accounts that emphasize brand risk, organizational politics, or skepticism "
        "toward AI claims. Access to participants through the researcher's professional "
        "network—primarily via LinkedIn—may further skew who agrees to participate and how "
        "frankly they speak about adoption difficulties."
    ),
    (
        "Constructivist grounded theory does not require the researcher to bracket prior "
        "knowledge entirely, but it does require reflexive transparency about how prior "
        "experience may shape what is noticed, questioned, and coded (Charmaz, 2014). "
        "Throughout data collection and analysis, analytic memos were used to record "
        "emerging interpretations separately from transcript evidence, to note surprises "
        "that challenged initial assumptions, and to document deliberate searches for "
        "disconfirming material. Interviewing two stakeholder groups—marketing managers "
        "and AI implementation experts—provided a further check against single-perspective "
        "blind spots. Section 3.4 describes how these practices were combined with "
        "additional trustworthiness strategies."
    ),
]

TRUSTWORTHINESS_BLOCKS: list[tuple[str, str]] = [
    ("h2", "3.4.\tEnsuring trustworthiness"),
    (
        "body",
        "Qualitative inquiry does not seek statistical generalization to a population; "
        "instead, readers assess whether findings are sufficiently grounded in data and "
        "whether the inquiry was conducted transparently. This study therefore adopts "
        "Lincoln and Guba's (1985) trustworthiness criteria—credibility, transferability, "
        "dependability, and confirmability—and implements them through the strategies "
        "proposed by Shenton (2004) for interview-based projects of limited duration. "
        "Shenton (2004) notes that tactics such as prolonged engagement and persistent "
        "observation, while valuable in long-term ethnography, are not realistic in "
        "shorter studies; the strategies below were selected accordingly.",
    ),
    ("h3", "3.4.1.\tCredibility"),
    (
        "body",
        "Credibility concerns whether the analysis is plausible and well grounded in "
        "participants' accounts. First, the study follows constructivist grounded theory "
        "(Charmaz, 2014), with iterative movement between interviewing, memo writing, and "
        "coding documented in Appendix C. Second, credibility was strengthened through "
        "methodological triangulation: marketing practitioners and AI implementation "
        "experts were interviewed about overlapping phenomena from different vantage "
        "points, which surfaced both convergence (for example, on implementation friction) "
        "and productive divergence (for example, on whether \"agentic AI\" denotes a "
        "distinct capability or mainly marketing language).",
    ),
    (
        "body",
        "Third, analysis proceeded iteratively. Initial coding identified line-by-line "
        "concepts; focused coding consolidated these into the categories summarized in "
        "Table 2 and developed in Chapter 4. Theoretical sampling guided later interviews "
        "and additions to the interview guide so that emerging categories could be "
        "compared across organizational contexts. Fourth, negative and paradoxical cases "
        "were retained rather than suppressed: participants who treated hallucination risk "
        "as a manageable design problem appear alongside those who regarded it as a "
        "reason to delay adoption; accounts of rapid prototyping coexist with accounts of "
        "\"unoperationalized tool effectiveness\" when experiments fail to embed in "
        "practice.",
    ),
    (
        "body",
        "Fifth, the researcher brought prior familiarity with marketing and AI practice "
        "(see Section 3.2.1). Shenton (2004) treats such familiarity as a credibility "
        "resource when paired with reflexive documentation. Sixth, emerging categories and "
        "chapter structure were discussed with the thesis supervisor during the analysis "
        "and writing phases, providing peer scrutiny of whether interpretations remained "
        "faithful to the data. Member checking—in which participants validate the "
        "researcher's analytic conclusions—was not conducted. In constructivist grounded "
        "theory, participants are experts on their experience but not on the analytic "
        "abstractions the researcher develops; returning full categories for approval can "
        "also constrain interpretation (Charmaz, 2014). Clarifying questions were asked "
        "in interviews when wording or context was ambiguous.",
    ),
    ("h3", "3.4.2.\tTransferability"),
    (
        "body",
        "Transferability replaces external validity in qualitative research: the reader "
        "judges whether findings might apply elsewhere (Lincoln & Guba, 1985; Shenton, "
        "2004). This study supports that judgment through thick description—contextual "
        "detail, process accounts, and extensive verbatim quotation in Chapter 4—and "
        "through transparent reporting of sampling logic in Section 3.3 and Table 1. "
        "Purposive and snowball sampling targeted informants who could speak to agentic AI "
        "in marketing-related roles rather than a statistically representative panel. "
        "Findings should therefore be read as analytically transferable insights into how "
        "marketing managers and AI experts make sense of agentic AI under comparable "
        "constraints, not as estimates of prevalence in a defined population. Scope "
        "conditions—including geographic concentration in the Netherlands and surrounding "
        "countries and the overweight of external AI advisors in the sample—are discussed "
        "in Section 5.3.",
    ),
    ("h3", "3.4.3.\tDependability"),
    (
        "body",
        "Dependability concerns whether the inquiry process is consistent and auditable. "
        "An audit trail links raw interview recordings and transcripts, cleaned transcript "
        "files, versioned qualitative analysis projects (QDPX exports dated between April "
        "and May 2026), analytic memos, the evolving code structure in Table 2, and the "
        "Findings narrative. The semi-structured interview guide and protocol (Appendix A) "
        "provided a baseline for data collection while allowing theoretically driven "
        "adaptation; backup and validation questions listed in Appendix A.1.3 illustrate "
        "how later interviews probed emerging themes. Coding decisions can be traced from "
        "quotation-level codings in the analysis software through focused categories to the "
        "aggregated concepts presented in Chapter 4. Code-application counts are reported "
        "descriptively where relevant; they indicate analytic emphasis, not statistical "
        "frequency.",
    ),
    ("h3", "3.4.4.\tConfirmability"),
    (
        "body",
        "Confirmability requires that interpretations remain traceable to data rather than "
        "to researcher preference alone (Lincoln & Guba, 1985; Shenton, 2004). Memos "
        "recorded analytic reasoning separately from transcript text. Findings claims are "
        "anchored in participant quotations, with supporting excerpts collated in "
        "Appendix D. Reflexive attention to how the researcher's marketing and AI "
        "background could shape attention and omission is documented in Section 3.2.1. "
        "Together, these practices aim to leave a decision trail that another researcher "
        "could follow and critique using the same materials.",
    ),
    ("h3", "3.4.5.\tTheoretical saturation"),
    (
        "body",
        "Theoretical saturation was assessed during iterative data collection and analysis "
        "(Charmaz, 2014). Interviews continued until later interviews chiefly densified "
        "properties of existing categories—adding nuance, boundary conditions, and "
        "examples—rather than introducing wholly new top-level concepts. The stopping "
        "point lay within the anticipated range of fifteen to twenty-five interviews and "
        "resulted in seventeen completed interviews (Table 1). Saturation pertained to "
        "the analytic categories used to organize Findings (observing change, affecting "
        "change, applying AI, value outcomes, and paradoxical tensions), not to closure "
        "of debate about a fast-moving technology. Section 4.6 summarizes how those "
        "categories relate to the coding development shown in Table 2.",
    ),
]

SECTION_46_BODY = (
    "Table 2 at the opening of this chapter summarizes how initial codes were consolidated "
    "into interpretive and aggregated categories. That hierarchy reflects focused coding "
    "and memo-based comparison across all seventeen interviews rather than a single "
    "pass of labeling. Trustworthiness of the analytic process—including credibility, "
    "transferability, dependability, and confirmability—is discussed in Section 3.4; "
    "this section records the code-development evidence on which Chapter 4 is built. "
    "Supporting quotations referenced in the Findings are collated in Appendix D."
)


def insert_paragraph_after(
    paragraph: Paragraph, text: str, style: str = "normal"
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    style_map = {
        "h1": "Heading 1",
        "h2": "Heading 2",
        "h3": "Heading 3",
        "normal": "normal",
        "body": "normal",
    }
    style_name = style_map.get(style, "normal")
    try:
        new_para.style = style_name
    except (KeyError, ValueError):
        new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para


def insert_blocks_after(anchor: Paragraph, blocks: list[tuple[str, str]]) -> Paragraph:
    current = anchor
    for kind, text in blocks:
        current = insert_paragraph_after(current, text, kind if kind != "body" else "normal")
    return current


def patch_document(doc: Document) -> None:
    paras = doc.paragraphs

    for p in paras:
        if "After reaching theoretical saturation (see section 4)" in p.text:
            p.text = (
                p.text.replace(
                    "After reaching theoretical saturation (see section 4)",
                    "After reaching theoretical saturation (see Section 3.4.5)",
                )
                .replace(" 3.2.1.\tResearcher's Bias", "")
                .replace(" 3.2.1.\tResearcher\u2019s Bias", "")
                .strip()
            )
            break

    todo_p = next(p for p in paras if p.text.strip() == "TODO")
    todo_p.text = "3.2.1.\tResearcher reflexivity"
    todo_p.style = "Heading 3"
    current = insert_paragraph_after(todo_p, REFLEXIVITY_PARAS[0], "normal")
    for extra in REFLEXIVITY_PARAS[1:]:
        current = insert_paragraph_after(current, extra, "normal")

    anchor = next(
        p for p in doc.paragraphs if "Overview of participants" in p.text
    )
    insert_blocks_after(anchor, TRUSTWORTHINESS_BLOCKS)

    for p in doc.paragraphs:
        if "4.x" in p.text and "Placeholder" in p.text:
            p.text = "4.6.\tCoding development and analytic density"
            p.style = "Heading 2"
            nxt = p._element.getnext()
            while nxt is not None:
                from docx.text.paragraph import Paragraph as P

                para = P(nxt, p._parent)
                if para.text.strip().startswith("5.") and "Discussion" in para.text:
                    break
                if "example" in para.text.lower() or "seek additional" in para.text.lower():
                    to_remove = nxt
                    nxt = nxt.getnext()
                    to_remove.getparent().remove(to_remove)
                    continue
                if not para.text.strip():
                    to_remove = nxt
                    nxt = nxt.getnext()
                    to_remove.getparent().remove(to_remove)
                    continue
                break
            insert_paragraph_after(p, SECTION_46_BODY, "normal")
            break


def main() -> None:
    doc = Document(str(DOCX_PATH))
    patch_document(doc)
    doc.save(str(DOCX_PATH))
    print(f"Patched {DOCX_PATH}")


if __name__ == "__main__":
    main()
