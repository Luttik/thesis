"""Split very long transcript lines into shorter paragraphs at sentence boundaries."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
FILES = [
    ROOT / "transcripts" / "Thesis interview Floris Reguoin.md",
    ROOT / "transcripts" / "Thesis interview Sylvia Vroklage.md",
]

# Lines longer than this get paragraph-split (speaker labels / short lines pass through).
MAX_LINE_CHARS = 580
# Target max characters per visual paragraph after splitting.
TARGET_PARAGRAPH_CHARS = 720


def _tiny_glued_sentence_fixes(text: str) -> str:
    """Fix missing space after period before capital (common ASR glitch)."""
    return re.sub(
        r"([a-zäöüà-ÿ])(\.)([A-ZÀ-Ö])",
        r"\1\2 \3",
        text,
    )


def _split_on_sentence_boundaries(text: str) -> list[str]:
    """Split into sentence-like chunks (rough, Dutch/English safe-ish)."""
    text = _tiny_glued_sentence_fixes(text.strip())
    if not text:
        return []
    # Split after . ! ? when followed by whitespace (keep delimiter on previous chunk).
    raw = re.split(r"(?<=[.!?])\s+", text)
    # Merge fragments that are too short (abbreviations / decimals).
    merged: list[str] = []
    for piece in raw:
        piece = piece.strip()
        if not piece:
            continue
        if merged and len(piece) < 12 and not piece[-1] in ".!?":
            merged[-1] = f"{merged[-1]} {piece}"
        elif merged and merged[-1] and merged[-1][-1] == "." and re.match(r"^[a-z]\.", piece):
            merged[-1] = f"{merged[-1]} {piece}"
        else:
            merged.append(piece)
    return merged


def _pack_into_paragraphs(sentences: list[str], target: int) -> list[str]:
    """Group sentences into paragraphs up to ~target characters."""
    if not sentences:
        return []
    out: list[str] = []
    buf: list[str] = []
    nlen = 0
    for s in sentences:
        add = len(s) + (1 if buf else 0)
        if buf and nlen + add > target and nlen >= target // 2:
            out.append(" ".join(buf))
            buf = [s]
            nlen = len(s)
        else:
            buf.append(s)
            nlen += add
    if buf:
        out.append(" ".join(buf))
    return out


def _split_plain_long_line(line: str) -> list[str]:
    sents = _split_on_sentence_boundaries(line)
    if len(sents) <= 1 and len(line) <= MAX_LINE_CHARS * 2:
        return [line]
    paras = _pack_into_paragraphs(sents, TARGET_PARAGRAPH_CHARS)
    if len(paras) <= 1:
        return [line]
    return paras


_ME_QUOTED = re.compile(r"^>\s*\*(.*)\s*\*\s*$")


def _split_me_quoted_line(line: str) -> list[str]:
    m = _ME_QUOTED.match(line.strip())
    if not m:
        return [line]
    inner = m.group(1).strip()
    if len(line) <= MAX_LINE_CHARS:
        return [line]
    sents = _split_on_sentence_boundaries(inner)
    paras = _pack_into_paragraphs(sents, TARGET_PARAGRAPH_CHARS)
    if len(paras) <= 1:
        return [line]
    return [f"> *{p} *" for p in paras]


def _is_speaker_label(line: str) -> bool:
    s = line.strip()
    if s in {"[Them]", "[Me]"}:
        return True
    if s in {"> *[Me]*", "> *[Them]*"}:
        return True
    return False


def process_file(path: Path) -> tuple[int, int]:
    """Returns (lines_in, lines_out)."""
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines(keepends=True)
    out: list[str] = []
    for line in lines:
        stripped = line.rstrip("\r\n")
        newline = "\n" if line.endswith("\n") else ""
        if len(stripped) <= MAX_LINE_CHARS or _is_speaker_label(stripped):
            out.append(line if line.endswith("\n") else stripped + newline)
            continue
        if stripped.startswith(">"):
            chunks = _split_me_quoted_line(stripped)
        else:
            chunks = _split_plain_long_line(stripped)
        if len(chunks) == 1:
            out.append(stripped + newline)
            continue
        for i, ch in enumerate(chunks):
            out.append(ch + "\n")
            if i < len(chunks) - 1:
                out.append("\n")
    new_text = "".join(out)
    path.write_text(new_text, encoding="utf-8")
    return len(lines), len(new_text.splitlines())


def export_docx(md_text: str, path: Path) -> None:
    doc = Document()
    for ln in md_text.splitlines():
        if ln.strip() == "":
            doc.add_paragraph("")
        elif ln.startswith("> "):
            p = doc.add_paragraph()
            p.add_run(ln[2:]).italic = True
        else:
            doc.add_paragraph(ln)
    doc.save(str(path))


def main() -> None:
    for path in FILES:
        before, after = process_file(path)
        md = path.read_text(encoding="utf-8")
        export_docx(md, path.with_suffix(".docx"))
        print(f"{path.name}: {before} -> {after} lines, docx refreshed")


if __name__ == "__main__":
    main()
