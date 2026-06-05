"""
Phase 1: Accept all tracked changes inside section 4.2 (between heading 4.2 and 4.3).
Phase 2: Insert italicised verb definitions as new tracked insertions.

Usage:
    cd C:\\workspace\\thesis
    poetry run python scripts/patch_sec42_verbs.py
"""

from __future__ import annotations
import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
AUTHOR = "Claude"
DATE = "2026-06-01T00:00:00Z"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def para_full_text(p_elem) -> str:
    """Return all text (including inside w:ins/w:del) for searching."""
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def para_visible_text(p_elem) -> str:
    """Text that would be visible after accepting changes (skip w:del content)."""
    parts = []
    for t in p_elem.iter(qn("w:t")):
        # skip if inside a w:del
        parent = t.getparent()
        in_del = False
        while parent is not None:
            if parent.tag == qn("w:del"):
                in_del = True
                break
            parent = parent.getparent()
        if not in_del:
            parts.append(t.text or "")
    return "".join(parts)


def next_rev_id(doc: Document) -> list[int]:
    """Counter that yields unique revision IDs."""
    used = {
        int(el.get(qn("w:id"), 0))
        for el in doc.element.body.iter()
        if el.get(qn("w:id")) is not None
    }
    current = [max(used, default=0) + 1]

    def get():
        val = current[0]
        current[0] += 1
        return val

    return get


# ---------------------------------------------------------------------------
# Phase 1: Accept tracked changes in a range of body elements
# ---------------------------------------------------------------------------

def accept_tracked_changes_in_range(elements: list) -> None:
    """
    For each element (paragraph or table) in the list:
    - w:ins  → unwrap: remove the <w:ins> tag, keep its children in place
    - w:del  → remove entirely
    Works recursively on all descendants.
    """
    for elem in elements:
        _accept_in_element(elem)


def _accept_in_element(elem) -> None:
    # Collect ins/del elements in document order (post-order to avoid mutation issues)
    # We do this iteratively to handle nested cases safely.
    changed = True
    while changed:
        changed = False
        for node in list(elem.iter()):
            if node.tag == qn("w:ins"):
                parent = node.getparent()
                if parent is None:
                    continue
                idx = list(parent).index(node)
                for i, child in enumerate(list(node)):
                    parent.insert(idx + i, child)
                parent.remove(node)
                changed = True
                break  # restart iteration after mutation
            elif node.tag == qn("w:del"):
                parent = node.getparent()
                if parent is None:
                    continue
                parent.remove(node)
                changed = True
                break


# ---------------------------------------------------------------------------
# Phase 2: Build definition paragraphs with italic verb names
# ---------------------------------------------------------------------------

def make_def_para(segments: list[tuple[str, bool]], rev_id_fn) -> etree._Element:
    """
    Build a <w:p> wrapped in <w:ins>.
    segments: list of (text, italic) tuples.
    """
    p = etree.Element(qn("w:p"))
    ins = etree.SubElement(p, qn("w:ins"))
    ins.set(qn("w:id"), str(rev_id_fn()))
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), DATE)

    for text, italic in segments:
        r = etree.SubElement(ins, qn("w:r"))
        if italic:
            rPr = etree.SubElement(r, qn("w:rPr"))
            etree.SubElement(rPr, qn("w:i"))
        t = etree.SubElement(r, qn("w:t"))
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text

    return p


def insert_after(body, anchor_elem, new_elem) -> None:
    idx = list(body).index(anchor_elem)
    body.insert(idx + 1, new_elem)


def insert_before(body, anchor_elem, new_elem) -> None:
    idx = list(body).index(anchor_elem)
    body.insert(idx, new_elem)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print(f"Opening: {DOCX_PATH}")
    doc = Document(str(DOCX_PATH))
    body = doc.element.body
    paras = doc.paragraphs
    rev_id = next_rev_id(doc)

    # ------------------------------------------------------------------
    # Identify section 4.2 range (body elements between 4.2 and 4.3)
    # ------------------------------------------------------------------
    sec42_start = None
    sec43_start = None
    for p in paras:
        txt = para_full_text(p._element)
        if sec42_start is None and "4.2" in txt and "Creating the context" in txt:
            sec42_start = p._element
        elif sec42_start is not None and sec43_start is None and "4.3" in txt and p.style.name in ("Heading 2", "heading 2"):
            sec43_start = p._element
            break

    if sec42_start is None:
        print("ERROR: Could not find section 4.2 heading.")
        return

    body_children = list(body)
    start_idx = body_children.index(sec42_start)
    end_idx = body_children.index(sec43_start) if sec43_start else len(body_children)
    sec42_elements = body_children[start_idx:end_idx]
    print(f"  Section 4.2 spans {len(sec42_elements)} body elements.")

    # ------------------------------------------------------------------
    # Phase 1: Accept all tracked changes in section 4.2
    # ------------------------------------------------------------------
    accept_tracked_changes_in_range(sec42_elements)
    print("  Phase 1: Accepted all tracked changes in section 4.2.")

    # Refresh paragraph list after XML mutations
    paras = doc.paragraphs

    # ------------------------------------------------------------------
    # Phase 2: Insert verb definitions
    # ------------------------------------------------------------------

    # 2a. Definitions for leveraging / navigating / reshaping
    # Find paragraph containing "managers leverage them"
    anchor_leverage = None
    for p in paras:
        if "managers leverage them" in para_visible_text(p._element):
            anchor_leverage = p._element
            break

    if anchor_leverage is not None:
        def_para = make_def_para([
            ("Leveraging", True),
            (" refers to using a favorable condition to amplify the organization’s ability to adopt and create value with agentic AI. ", False),
            ("Navigating", True),
            (" refers to adapting the adoption strategy to work within an unfavorable condition, accepting it as a given constraint. ", False),
            ("Reshaping", True),
            (" refers to actively working to change the condition itself.", False),
        ], rev_id)
        insert_after(body, anchor_leverage, def_para)
        print("  Inserted leveraging/navigating/reshaping definitions.")
    else:
        print("  WARNING: Could not find 'managers leverage them' paragraph.")

    # Refresh after insertion
    paras = doc.paragraphs

    # 2b. Definitions for experimenting with AI / educating and training
    anchor_experiment = None
    for p in paras:
        txt = para_visible_text(p._element)
        if "experimenting with AI (37 code" in txt or "experimenting with AI (3" in txt:
            anchor_experiment = p._element
            break

    if anchor_experiment is not None:
        def_para = make_def_para([
            ("Experimenting with AI", True),
            (" refers to running small-scale trials of agentic tools and workflows to generate practical knowledge and build organizational confidence. ", False),
            ("Educating and training", True),
            (" refers to the deliberate effort to raise AI literacy and capability across the organization.", False),
        ], rev_id)
        insert_before(body, anchor_experiment, def_para)
        print("  Inserted experimenting/educating definitions.")
    else:
        print("  WARNING: Could not find experimenting-with-AI paragraph.")

    # Refresh after insertion
    paras = doc.paragraphs

    # 2c. Definition for bringing people along
    anchor_bringing = None
    for p in paras:
        txt = para_visible_text(p._element)
        if "The process of bringing people along" in txt:
            anchor_bringing = p._element
            break

    if anchor_bringing is not None:
        def_para = make_def_para([
            ("Bringing people along", True),
            (" refers to the managerial process of building alignment, managing resistance, and developing shared commitment to AI adoption within the organization.", False),
        ], rev_id)
        insert_before(body, anchor_bringing, def_para)
        print("  Inserted bringing-people-along definition.")
    else:
        print("  WARNING: Could not find 'bringing people along' paragraph.")

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------
    doc.save(str(DOCX_PATH))
    print(f"\nSaved in-place: {DOCX_PATH}")


if __name__ == "__main__":
    main()
