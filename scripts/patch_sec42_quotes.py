"""
Add missing quotes and code counts to section 4.2.

Changes (all tracked, author "Claude"):
1. Add code count to Educating and training paragraph
2. Add code count to Experimenting with AI paragraph
3. Sharpen Leadership backing definition
4. Add I6 external-experts quote to navigate paragraph
5. Replace duplicate I6 vision quote in bringing-people-along with I14
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT      = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"

AUTHOR = "Claude"
DATE   = "2026-06-05T00:00:00Z"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _used_ids(doc: Document) -> set[int]:
    return {int(el.get(qn("w:id"), 0))
            for el in doc.element.body.iter()
            if el.get(qn("w:id")) is not None}

def _nid(used: set[int]) -> int:
    n = max(used, default=0) + 1; used.add(n); return n


def full_text(p_elem) -> str:
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def find_para(doc: Document, marker: str) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if marker in full_text(p._element):
            return i
    return None


def ins_run(text: str, used: set[int], bold: bool = False) -> etree._Element:
    ins = etree.Element(qn("w:ins"))
    ins.set(qn("w:id"), str(_nid(used)))
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), DATE)
    r = etree.SubElement(ins, qn("w:r"))
    if bold:
        rPr = etree.SubElement(r, qn("w:rPr"))
        etree.SubElement(rPr, qn("w:b"))
    t = etree.SubElement(r, qn("w:t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return ins


def del_run(r_elem: etree._Element, used: set[int]) -> etree._Element:
    """Wrap a single existing run in <w:del>."""
    d = etree.Element(qn("w:del"))
    d.set(qn("w:id"), str(_nid(used)))
    d.set(qn("w:author"), AUTHOR)
    d.set(qn("w:date"), DATE)
    # convert w:t -> w:delText inside a copy
    import copy
    rc = copy.deepcopy(r_elem)
    for t in rc.findall(qn("w:t")):
        dt = etree.Element(qn("w:delText"))
        dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        dt.text = t.text or ""
        t.getparent().replace(t, dt)
    d.append(rc)
    return d


def ins_para(text: str, style: str, used: set[int],
             block_quote: bool = False) -> etree._Element:
    p = etree.Element(qn("w:p"))
    pPr = etree.SubElement(p, qn("w:pPr"))
    if style and style.lower() != "normal":
        ps = etree.SubElement(pPr, qn("w:pStyle"))
        ps.set(qn("w:val"), style)
    if block_quote:
        ind = etree.SubElement(pPr, qn("w:ind"))
        ind.set(qn("w:left"), "720")
    ins = etree.SubElement(p, qn("w:ins"))
    ins.set(qn("w:id"), str(_nid(used)))
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), DATE)
    r = etree.SubElement(ins, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return p


def insert_after(anchor: etree._Element, paras: list) -> None:
    parent = anchor.getparent()
    idx = list(parent).index(anchor) + 1
    for i, p in enumerate(paras):
        parent.insert(idx + i, p)


# ---------------------------------------------------------------------------
# Inline text replacement as tracked del+ins within a paragraph
# ---------------------------------------------------------------------------

def replace_text_tracked(p_elem: etree._Element,
                         old: str, new: str,
                         used: set[int]) -> bool:
    """
    Replace `old` string in paragraph with tracked del+ins.
    Operates on the first <w:t> that contains the full old string
    (sufficient when old is entirely within one run, which is true for
    the code-count insertions).
    Returns True if replacement was made.
    """
    for r in p_elem.findall(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is None or not t.text:
            continue
        if old not in t.text:
            continue
        # Split into before / old / after
        before, _, after = t.text.partition(old)
        parent = r.getparent()
        idx = list(parent).index(r)

        # Build replacement sequence: ins(before) del(old) ins(new) ins(after)
        # Simpler: keep the run for before, del the old part, ins the new
        parts = []
        if before:
            rb = etree.Element(qn("w:r"))
            tb = etree.SubElement(rb, qn("w:t"))
            tb.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            tb.text = before
            parts.append(rb)

        d = etree.Element(qn("w:del"))
        d.set(qn("w:id"), str(_nid(used)))
        d.set(qn("w:author"), AUTHOR)
        d.set(qn("w:date"), DATE)
        rd = etree.Element(qn("w:r"))
        td = etree.SubElement(rd, qn("w:delText"))
        td.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        td.text = old
        d.append(rd)
        parts.append(d)

        i_new = ins_run(new, used)
        parts.append(i_new)

        if after:
            ra = etree.Element(qn("w:r"))
            ta = etree.SubElement(ra, qn("w:t"))
            ta.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            ta.text = after
            parts.append(ra)

        # Insert parts, remove original run
        for j, part in enumerate(parts):
            parent.insert(idx + j, part)
        parent.remove(r)
        return True
    return False


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    shutil.copy(DOCX_PATH, DOCX_PATH.with_suffix(".bak.docx"))
    doc = Document(str(DOCX_PATH))
    used = _used_ids(doc)

    # ------------------------------------------------------------------
    # 1. Add code count to Educating paragraph
    # ------------------------------------------------------------------
    idx = find_para(doc, "Educating and training is the primary response")
    assert idx, "Educating para not found"
    replaced = replace_text_tracked(
        doc.paragraphs[idx]._element,
        "Educating and training is",
        "Educating and training (28 code applications, 14 documents) is",
        used)
    print(f"  Educating code count: {'added' if replaced else 'FAILED'}")

    # ------------------------------------------------------------------
    # 2. Add code count to Experimenting paragraph
    # ------------------------------------------------------------------
    idx = find_para(doc, "Experimenting with AI primarily addresses resistance")
    assert idx, "Experimenting para not found"
    replaced = replace_text_tracked(
        doc.paragraphs[idx]._element,
        "Experimenting with AI primarily addresses",
        "Experimenting with AI (37 code applications, 17 documents) primarily addresses",
        used)
    print(f"  Experimenting code count: {'added' if replaced else 'FAILED'}")

    # ------------------------------------------------------------------
    # 3. Sharpen Leadership backing definition
    # ------------------------------------------------------------------
    idx = find_para(doc, "Leadership backing amplifies the champion effect")
    assert idx, "Leadership backing para not found"
    replaced = replace_text_tracked(
        doc.paragraphs[idx]._element,
        "Leadership backing amplifies the champion effect (14 code applications, 10 documents).",
        "Leadership backing refers to the visible endorsement, resource authorization, "
        "and organizational protection that senior leaders provide for AI initiatives "
        "(14 code applications, 10 documents). It amplifies the champion effect.",
        used)
    print(f"  Leadership backing definition: {'updated' if replaced else 'FAILED'}")

    # ------------------------------------------------------------------
    # 4. Add I6 external-experts quote to navigate paragraph
    # ------------------------------------------------------------------
    idx = find_para(doc,
        "some managers navigate around organizational capacity constraints")
    assert idx, "Navigate para not found"
    nav_elem = doc.paragraphs[idx]._element

    attribution = ("Interviewee 6 described this pattern: "
                   "“Another one is actually using experts externally, because if "
                   "we would’ve waited for people internally, I think it would’ve "
                   "been a major blocker.”")

    # Replace last sentence to add the attribution naturally
    replaced = replace_text_tracked(
        nav_elem,
        "This bypasses the bottleneck without fixing it, a pragmatic short-term "
        "strategy that can stall long-term capability development if it becomes the default.",
        "This bypasses the bottleneck without fixing it, a pragmatic short-term strategy "
        "that can stall long-term capability development if it becomes the default. "
        + attribution,
        used)
    print(f"  Navigate quote: {'added' if replaced else 'FAILED'}")

    # ------------------------------------------------------------------
    # 5. Replace duplicate I6 quote in bringing-people-along with I14
    # ------------------------------------------------------------------
    idx = find_para(doc,
        "Interviewee 6 captured the essential managerial move: "
        "“This is the vision. This is where we’re going. Help me get there.”")
    if idx is None:
        # Try with ASCII quotes
        idx = find_para(doc,
            "Interviewee 6 captured the essential managerial move")

    if idx is not None:
        ft = full_text(doc.paragraphs[idx]._element)
        if "Help me get there" in ft:
            replaced = replace_text_tracked(
                doc.paragraphs[idx]._element,
                "Interviewee 6 captured the essential managerial move: "
                "“This is the vision. This is where we’re going. "
                "Help me get there.” Interviewee",
                "Interviewee 14 framed the scale of the challenge: "
                "“Thirty percent of all AI trajectories is just technology; "
                "sixty to seventy percent is actually us as humans or as an organization.” "
                "Interviewee",
                used)
            print(f"  Bringing-people-along duplicate quote: {'replaced' if replaced else 'FAILED'}")
        else:
            print("  Bringing-people-along: I6 quote not found in expected paragraph.")
    else:
        print("  Bringing-people-along: paragraph not found.")

    doc.save(str(DOCX_PATH))
    print(f"\nSaved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
