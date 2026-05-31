"""Verify Table 3 placement in the thesis docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
OUT = ROOT / ".cache" / "place.txt"


def main() -> None:
    d = Document(str(DOCX))
    lines: list[str] = [f"tables: {len(d.tables)}"]
    title_p = next(p for p in d.paragraphs if p.text.strip() == "Trustworthiness assessment.")
    n = title_p._p.getnext()
    if n is None:
        lines.append("after title: NONE — table may be missing from flow")
    else:
        lines.append(f"after title: {n.tag.split('}')[-1]}")
        if n.tag.endswith("}tbl"):
            lines.append(f"table rows: {len(n.findall('.//' + qn('w:tr')))}")
        n2 = n.getnext()
        if n2 is not None:
            if n2.tag.endswith("}p"):
                t = "".join(x.text or "" for x in n2.iter(qn("w:t")))
                lines.append(f"after table/next: {t[:60]}")
    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(OUT.read_text(encoding="utf-8"))


if __name__ == "__main__":
    main()
