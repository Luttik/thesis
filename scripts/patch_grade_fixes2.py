# -*- coding: utf-8 -*-
"""Follow-up: Charmaz 2014 dedup (nbsp), Vidal DOI. Tracked, author Claude."""
from __future__ import annotations
import copy, shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Thesis Draft - Daan Luttik - MBA - copy from 2026-06-14.docx"
BACKUP = ROOT / "Thesis Draft - Daan Luttik - MBA - copy from 2026-06-14.gradefix2-backup.docx"
AUTHOR="Claude"; DATE="2026-06-15T00:00:00Z"
XMLSPACE="{http://www.w3.org/XML/1998/namespace}space"
def _used_ids(doc): return {int(el.get(qn("w:id"),0)) for el in doc.element.body.iter() if el.get(qn("w:id")) is not None}
def _nid(u): n=max(u,default=0)+1; u.add(n); return n
def _stamp(el,u): el.set(qn("w:id"),str(_nid(u))); el.set(qn("w:author"),AUTHOR); el.set(qn("w:date"),DATE)
def ftext(pe): return "".join(t.text or "" for t in pe.iter(qn("w:t")))
def _mk_run(text,rpr_t=None):
    r=etree.Element(qn("w:r"))
    if rpr_t is not None:
        rpr=copy.deepcopy(rpr_t)
        for b in rpr.findall(qn("w:ins"))+rpr.findall(qn("w:del")): rpr.remove(b)
    else: rpr=etree.Element(qn("w:rPr"))
    if len(rpr): r.append(rpr)
    t=etree.SubElement(r,qn("w:t")); t.set(XMLSPACE,"preserve"); t.text=text; return r
def _wrap_ins(ch,u):
    ins=etree.Element(qn("w:ins")); _stamp(ins,u)
    for c in ch: ins.append(c)
    return ins
def _wrap_del(txt,rpr_t,u):
    d=etree.Element(qn("w:del")); _stamp(d,u); r=etree.SubElement(d,qn("w:r"))
    if rpr_t is not None:
        rpr=copy.deepcopy(rpr_t)
        for b in rpr.findall(qn("w:ins"))+rpr.findall(qn("w:del")): rpr.remove(b)
        if len(rpr): r.append(rpr)
    dt=etree.SubElement(r,qn("w:delText")); dt.set(XMLSPACE,"preserve"); dt.text=txt; return d
def replace_tracked(p,old,new,u):
    runs=p.findall(qn("w:r")); spans=[]; full=""
    for r in runs:
        t=r.find(qn("w:t")); tx=(t.text if t is not None and t.text else "")
        spans.append((r,len(full),len(full)+len(tx),tx)); full+=tx
    idx=full.find(old)
    if idx==-1: return False
    end=idx+len(old); ins_done=False
    for r,rs,re_,tx in spans:
        if re_<=idx or rs>=end: continue
        rpr=r.find(qn("w:rPr")); before=tx[:idx-rs] if rs<idx else ""; after=tx[end-rs:] if re_>end else ""
        deleted=tx[max(idx,rs)-rs:min(end,re_)-rs]
        parent=r.getparent(); i=list(parent).index(r); parts=[]
        if before: parts.append(_mk_run(before,rpr))
        if deleted: parts.append(_wrap_del(deleted,rpr,u))
        if not ins_done and new: parts.append(_wrap_ins([_mk_run(new,rpr)],u)); ins_done=True
        if after: parts.append(_mk_run(after,rpr))
        for j,pt in enumerate(parts): parent.insert(i+j,pt)
        parent.remove(r)
    return True
def find_start(doc,pre):
    for p in doc.paragraphs:
        if ftext(p._element).strip().startswith(pre): return p._element
    return None

def main():
    shutil.copy(DOCX,BACKUP)
    doc=Document(str(DOCX)); u=_used_ids(doc)
    # Charmaz nbsp dedup
    ch=find_start(doc,"Charmaz, K. (2014)")
    if ch is not None:
        ok=replace_tracked(ch,"(introducing qualitative methods series).\xa0Constructing grounded theory (2nd ed.)","(2nd ed.)",u)
        print(f"  {'ok  ' if ok else 'MISS'} Charmaz 2014 dedup (nbsp)")
    else: print("  FAIL Charmaz not found")
    # Vidal DOI: locate entry by Perotti, report + fix if needed
    vidal=None
    for p in doc.paragraphs:
        tx=ftext(p._element)
        if tx.strip().startswith("Vidal") and "Perotti" in tx and "Journal of Business Research" in tx:
            vidal=p._element; break
    if vidal is not None:
        tx=ftext(vidal); print("  VIDAL entry DOI tail:", repr(tx[-40:]))
        if "07.0200" in tx:
            ok=replace_tracked(vidal,"jbusres.2022.07.0200","jbusres.2022.07.020",u)
            print(f"  {'ok  ' if ok else 'MISS'} Vidal DOI 0200->020")
        else:
            print("  ok   Vidal DOI already clean (no '07.0200')")
    else:
        print("  note: Vidal reference entry not located by (Perotti + JBR); skipped")
    doc.save(str(DOCX)); print("Saved + backup", BACKUP.name)

if __name__=="__main__": main()
