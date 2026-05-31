"""Scan section 4 paragraph indices in the thesis docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
OUT = ROOT / ".cache" / "sec4_scan.txt"


def main() -> None:
    doc = Document(str(DOCX))
    lines: list[str] = []
    lines.append(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        if 120 <= i <= 300 or "4." in t[:6]:
            lines.append(f"{i:3}|{p.style.name[:14]:14}|{t[:100]}")
    for ti, table in enumerate(doc.tables):
        if table.rows:
            lines.append(f"TABLE{ti} R0C0={table.rows[0].cells[0].text[:60]}")
    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(f"wrote {OUT} ({len(lines)} lines)")


if __name__ == "__main__":
    main()
