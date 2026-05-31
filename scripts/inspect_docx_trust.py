"""Inspect §3.4 and trustworthiness table state in thesis docx files."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
FILES = [ROOT / "Thesis Draft - Daan Luttik - MBA.docx"]


def inspect(path: Path) -> list[str]:
    if not path.exists():
        return [f"{path.name}: NOT FOUND"]
    try:
        d = Document(str(path))
    except Exception as exc:
        return [f"{path.name}: CANNOT OPEN — {exc}"]
    lines = [f"\n=== {path.name} ({path.stat().st_size} bytes) ===", f"tables: {len(d.tables)}"]
    trust_tables = []
    for i, t in enumerate(d.tables):
        if t.rows and "trustworthiness" in t.rows[0].cells[0].text.lower():
            trust_tables.append(i)
            lines.append(f"  table[{i}]: {len(t.rows)} rows — {t.rows[0].cells[0].text[:50]}")
    if not trust_tables:
        lines.append("  NO trustworthiness table found")

    lines.append("§3.4 area paragraphs:")
    in_34 = False
    for p in d.paragraphs:
        t = p.text.strip()
        if t.startswith("3.4."):
            in_34 = True
        if in_34:
            if t.startswith("4.") and "Findings" in t:
                lines.append(f"  -> {t[:80]}")
                break
            if t:
                lines.append(f"  | {t[:100]}")
        elif "3.4.1" in t or "3.4.2" in t:
            lines.append(f"  (old subsection?) {t[:80]}")

    # Body order around 3.4
    lines.append("Body sequence (3.4 .. Findings):")
    capturing = False
    for child in d.element.body:
        if child.tag.endswith("}p"):
            text = "".join(x.text or "" for x in child.iter(qn("w:t"))).strip()
            if text.startswith("3.4."):
                capturing = True
            if not capturing:
                continue
            lines.append(f"  P: {text[:90]}")
            if text.startswith("4.") and "Findings" in text:
                break
        elif child.tag.endswith("}tbl") and capturing:
            hdr = "".join(x.text or "" for x in child.iter(qn("w:t")))[:70]
            lines.append(f"  TBL: {hdr}")
    return lines


def main() -> None:
    out = ROOT / ".cache" / "inspect-trust.txt"
    all_lines: list[str] = []
    for f in FILES:
        all_lines.extend(inspect(f))
    text = "\n".join(all_lines)
    out.write_text(text, encoding="utf-8")
    print(text)


if __name__ == "__main__":
    main()
