# -*- coding: utf-8 -*-
"""Verify the Chapter 3 tense/sampling tracked patch.

- counts w:ins / w:del and their authors
- simulates accept-all and reject-all (reload as Document, read paras[109:143])
- asserts the accepted wording is correct and reject-all == pre-patch backup
"""
from __future__ import annotations

import copy
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT   = Path(__file__).resolve().parents[1]
DOCX   = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
BACKUP = ROOT / "Thesis Draft - Daan Luttik - MBA.ch3-backup.docx"
XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"

CH3 = range(109, 143)  # methodology body paragraphs


def counts(doc):
    body = doc.element.body
    ins = body.findall(".//" + qn("w:ins"))
    dele = body.findall(".//" + qn("w:del"))
    authors = sorted({e.get(qn("w:author")) for e in ins + dele})
    return len(ins), len(dele), authors


def accept_all(body):
    for el in list(body.findall(".//" + qn("w:del"))):
        el.getparent().remove(el)
    for ins in list(body.findall(".//" + qn("w:ins"))):
        parent = ins.getparent()
        idx = list(parent).index(ins)
        for j, child in enumerate(list(ins)):
            parent.insert(idx + j, child)
        parent.remove(ins)


def reject_all(body):
    for ins in list(body.findall(".//" + qn("w:ins"))):
        ins.getparent().remove(ins)
    for d in list(body.findall(".//" + qn("w:del"))):
        parent = d.getparent()
        idx = list(parent).index(d)
        for child in list(d):
            for dt in child.findall(qn("w:delText")):
                t = etree.Element(qn("w:t"))
                t.set(XMLSPACE, "preserve")
                t.text = dt.text
                dt.getparent().replace(dt, t)
        for j, child in enumerate(list(d)):
            parent.insert(idx + j, child)
        parent.remove(d)


def resolved_paras(src, mode):
    """Apply accept/reject to a copy, save+reload, return paragraph texts."""
    doc = Document(str(src))
    (accept_all if mode == "accept" else reject_all)(doc.element.body)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = tmp.name
    doc.save(path)
    rl = Document(path)
    Path(path).unlink(missing_ok=True)
    return [p.text for p in rl.paragraphs], rl


# --- tracked-change counts -------------------------------------------------- #
live = Document(str(DOCX))
ni, nd, authors = counts(live)
print(f"tracked: {ni} w:ins, {nd} w:del   authors={authors}")
assert authors == ["Claude"], f"unexpected authors: {authors}"

# --- accept-all ------------------------------------------------------------- #
acc, _ = resolved_paras(DOCX, "accept")
print("\n=== ACCEPT-ALL : Chapter 3 (3.1-3.3) ===")
for i in CH3:
    if i < len(acc) and acc[i].strip():
        print(f"{i}: {acc[i]}")

ch3_acc = "\n".join(acc[i] for i in CH3 if i < len(acc))

checks = {
    "no 'purposive' anywhere in Ch3": "purposive" not in ch3_acc.lower(),
    "heading retitled": any("Sample and sample size" in acc[i] and "Intended" not in acc[i] for i in CH3 if i < len(acc)),
    "no 'will be collected' (111)": "will be collected" not in ch3_acc,
    "no 'will be used to spark' (113)": "will be used to spark" not in ch3_acc,
    "no 'will be augmented' (119)": "will be augmented" not in ch3_acc,
    "no 'will be utilized' (119)": "will be utilized" not in ch3_acc,
    "theoretical sampling (Charmaz, 2014) in 122": "theoretical sampling (Charmaz, 2014)" in acc[122],
    "no Patton in 122": "Patton" not in acc[122],
    "'were interviewed' (121)": "were interviewed" in acc[121],
    "'was used to identify' (122)": "was used to identify" in acc[122],
    "'were noted' (122)": "are noted" not in acc[122] and "were noted" in acc[122],
}
print("\n=== CHECKS (accept-all) ===")
ok = True
for k, v in checks.items():
    print(f"  [{'PASS' if v else 'FAIL'}] {k}")
    ok = ok and v

# --- reject-all == pre-patch backup ----------------------------------------- #
rej, _ = resolved_paras(DOCX, "reject")
base = [p.text for p in Document(str(BACKUP)).paragraphs]
mism = [i for i in CH3 if i < len(rej) and i < len(base) and rej[i] != base[i]]
print("\n=== REJECT-ALL vs pre-patch backup (Ch3) ===")
print(f"  paragraph count: reject={len(rej)} backup={len(base)}")
print(f"  mismatched Ch3 paragraphs: {mism if mism else 'NONE (reject-all == original)'}")
ok = ok and not mism and len(rej) == len(base)

# residual unresolved markers after each pass should be zero
acc_doc = Document(str(DOCX)); accept_all(acc_doc.element.body)
rej_doc = Document(str(DOCX)); reject_all(rej_doc.element.body)
res_acc = len(acc_doc.element.body.findall(".//" + qn("w:ins"))) + len(acc_doc.element.body.findall(".//" + qn("w:del")))
res_rej = len(rej_doc.element.body.findall(".//" + qn("w:ins"))) + len(rej_doc.element.body.findall(".//" + qn("w:del")))
print(f"\n  residual markers after accept: {res_acc} | after reject: {res_rej}")
ok = ok and res_acc == 0 and res_rej == 0

print("\nRESULT:", "ALL GOOD" if ok else "PROBLEMS FOUND")
