"""Collapse excessive uh/uhm chains and 3+ same-word comma stutters in thesis transcripts."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
FILES = [
    ROOT / "transcripts" / "Thesis interview Floris Reguoin.md",
    ROOT / "transcripts" / "Thesis interview Sylvia Vroklage.md",
]

# Same word repeated 3+ times with commas (Dutch dysfluency); keep "word, word" pairs.
_STUTTER_WORDS = (
    "en",
    "dat",
    "de",
    "ik",
    "dus",
    "het",
    "in",
    "is",
    "maar",
    "dan",
    "daar",
    "wat",
    "dit",
    "toch",
    "zo",
    "nog",
    "voor",
    "ja",
    "nee",
    "toen",
    "een",
)


def _collapse_triple_comma_word(text: str, word: str) -> str:
    w = re.escape(word)
    # At least 3 occurrences: w, w, w
    pat = re.compile(rf"(?i)\b({w})(,\s*\1){{2,}}\b")
    return pat.sub(r"\1", text)


def _collapse_filler_token_runs(text: str, token: str) -> str:
    t = re.escape(token)
    for _ in range(80):
        n = re.sub(rf"(?i)\b{t}\b\s*,\s*\b{t}\b", token, text)
        n = re.sub(rf"(?i)\b{t}\b\s+\b{t}\b", token, n)
        if n == text:
            return n
        text = n
    return text


def _collapse_double_comma_word(text: str, word: str) -> str:
    """Collapse ``word, word`` dysfluency (two repeats only)."""
    w = re.escape(word)
    pat = re.compile(rf"(?i)\b({w}),\s*\1\b")
    for _ in range(40):
        n = pat.sub(r"\1", text)
        if n == text:
            return n
        text = n
    return text


def _fix_glued_typos(text: str) -> str:
    return (
        text.replace("dieDie", "die wij")
        .replace("die wij wij ", "die wij ")
        .replace("die wij wij,", "die wij,")
        .replace("uhmMaar", "uhm. Maar")
        .replace("ChatGPT of Cloud of een", "ChatGPT of Claude of een")
        .replace("we-werknemer", "werknemer")
        .replace("uhOf", "uh. Of")
        .replace(
            "validiteit van, uh, van, van echt heeft",
            "validiteit van het echte heeft",
        )
        .replace("en om, om op", "en om op")
    )


def _fix_letter_stutters(text: str) -> str:
    text = re.sub(r"(?i)\bo-o-,\s*", "", text)
    text = re.sub(r"(?i)\ba-als\b", "als", text)
    return text


def dedupe(text: str) -> str:
    text = _fix_glued_typos(text)
    text = _fix_letter_stutters(text)
    text = re.sub(r"(?i)\buh,\s*uhm\b", "uhm", text)
    text = re.sub(r"(?i)\buhm,\s*uh\b", "uhm", text)
    for tok in ("uh", "uhm", "eh"):
        text = _collapse_filler_token_runs(text, tok)
    for w in _STUTTER_WORDS:
        text = _collapse_triple_comma_word(text, w)
    for w in ("en", "dat", "dus", "het", "toen", "in", "is"):
        text = _collapse_double_comma_word(text, w)
    text = re.sub(r"(?i)\b(wij)(\s*,\s*\1){2,}\b", r"\1", text)
    # Clean double spaces left in prose (not line starts)
    text = re.sub(r"([^*\n])  +", r"\1 ", text)
    return text


def export_docx(md_text: str, path: Path) -> None:
    doc = Document()
    for line in md_text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.add_run(line[2:]).italic = True
        else:
            doc.add_paragraph(line)
    doc.save(str(path))


def main() -> None:
    for path in FILES:
        raw = path.read_text(encoding="utf-8")
        out = dedupe(raw)
        path.write_text(out, encoding="utf-8")
        export_docx(out, path.with_suffix(".docx"))
        print(f"Updated {path.name} + .docx ({len(raw)} -> {len(out)} chars)")


if __name__ == "__main__":
    main()
