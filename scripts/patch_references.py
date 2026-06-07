# -*- coding: utf-8 -*-
"""
Reference reconciliation pass (all tracked changes, author "Claude").

Part 1: add 20 new reference-list entries (17 orphans + Lepak/Brynjolfsson/Kumar),
        alphabetical, matching the existing hanging-indent + highlight-white style,
        journal name + volume italic. Charmaz 2006 entry converted to the 2014 2nd ed.
Part 2: in-text fixes (year/spelling/author) as tracked del+ins.
Part 3: insert new citations at the supervisor's "citation needed" anchors
        (located precisely via each comment's commentRangeEnd).
"""
from __future__ import annotations
import copy, shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT      = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
BACKUP    = ROOT / "Thesis Draft - Daan Luttik - MBA.refs-backup.docx"
AUTHOR    = "Claude"
DATE      = "2026-06-07T00:00:00Z"
XMLSPACE  = "{http://www.w3.org/XML/1998/namespace}space"
EN        = "–"   # en dash for page ranges


def _used_ids(doc):
    return {int(el.get(qn("w:id"), 0)) for el in doc.element.body.iter()
            if el.get(qn("w:id")) is not None}
def _nid(used):
    n = max(used, default=0) + 1; used.add(n); return n
def ftext(p): return "".join(t.text or "" for t in p.iter(qn("w:t")))
def present(doc, s): return any(s in ftext(p._element) for p in doc.paragraphs)
def find_para(doc, marker):
    for p in doc.paragraphs:
        if marker in ftext(p._element):
            return p._element
    return None


# ---------- run builders ----------
def _mk_run(text, rpr_template=None, italic=False):
    r = etree.Element(qn("w:r"))
    if rpr_template is not None:
        rpr = copy.deepcopy(rpr_template)
        # strip any ins/del tracking that might be on the template
        for bad in rpr.findall(qn("w:ins")) + rpr.findall(qn("w:del")):
            rpr.remove(bad)
    else:
        rpr = etree.Element(qn("w:rPr"))
    if italic and rpr.find(qn("w:i")) is None:
        rpr.insert(0, etree.Element(qn("w:iCs")))
        rpr.insert(0, etree.Element(qn("w:i")))
    if len(rpr):
        r.append(rpr)
    t = etree.SubElement(r, qn("w:t")); t.set(XMLSPACE, "preserve"); t.text = text
    return r

def _wrap_ins(children, used):
    ins = etree.Element(qn("w:ins"))
    ins.set(qn("w:id"), str(_nid(used))); ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), DATE)
    for c in children: ins.append(c)
    return ins

def _wrap_del(deleted_text, rpr_template, used):
    d = etree.Element(qn("w:del"))
    d.set(qn("w:id"), str(_nid(used))); d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), DATE)
    r = etree.SubElement(d, qn("w:r"))
    if rpr_template is not None:
        rpr = copy.deepcopy(rpr_template)
        for bad in rpr.findall(qn("w:ins")) + rpr.findall(qn("w:del")):
            rpr.remove(bad)
        if len(rpr): r.append(rpr)
    dt = etree.SubElement(r, qn("w:delText")); dt.set(XMLSPACE, "preserve"); dt.text = deleted_text
    return d


# ---------- reference paragraph (fully tracked insert) ----------
def ref_para(segments, used):
    """segments: list of (text, italic). Builds a tracked reference paragraph."""
    p = etree.Element(qn("w:p"))
    pPr = etree.SubElement(p, qn("w:pPr"))
    ind = etree.SubElement(pPr, qn("w:ind")); ind.set(qn("w:left"), "720"); ind.set(qn("w:hanging"), "720")
    rprmark = etree.SubElement(pPr, qn("w:rPr"))
    insmark = etree.SubElement(rprmark, qn("w:ins"))
    insmark.set(qn("w:id"), str(_nid(used))); insmark.set(qn("w:author"), AUTHOR); insmark.set(qn("w:date"), DATE)
    hl = etree.SubElement(rprmark, qn("w:highlight")); hl.set(qn("w:val"), "white")
    runs = []
    for text, italic in segments:
        r = etree.Element(qn("w:r"))
        rpr = etree.SubElement(r, qn("w:rPr"))
        if italic:
            etree.SubElement(rpr, qn("w:i")); etree.SubElement(rpr, qn("w:iCs"))
        h = etree.SubElement(rpr, qn("w:highlight")); h.set(qn("w:val"), "white")
        t = etree.SubElement(r, qn("w:t")); t.set(XMLSPACE, "preserve"); t.text = text
        runs.append(r)
    p.append(_wrap_ins(runs, used))
    return p


# ---------- run-spanning tracked replace ----------
def replace_tracked(p, old, new, used):
    runs = p.findall(qn("w:r"))
    spans = []
    full = ""
    for r in runs:
        t = r.find(qn("w:t"))
        txt = (t.text if t is not None and t.text else "")
        spans.append((r, len(full), len(full) + len(txt), txt))
        full += txt
    idx = full.find(old)
    if idx == -1:
        return False
    end = idx + len(old)
    inserted = False
    for r, rs, re_, txt in spans:
        if re_ <= idx or rs >= end:
            continue
        rpr = r.find(qn("w:rPr"))
        before = txt[:idx - rs] if rs < idx else ""
        after = txt[end - rs:] if re_ > end else ""
        os_, oe_ = max(idx, rs) - rs, min(end, re_) - rs
        deleted = txt[os_:oe_]
        parent = r.getparent(); i = list(parent).index(r)
        parts = []
        if before: parts.append(_mk_run(before, rpr))
        if deleted: parts.append(_wrap_del(deleted, rpr, used))
        if not inserted:
            parts.append(_wrap_ins([_mk_run(new, rpr)], used)); inserted = True
        if after: parts.append(_mk_run(after, rpr))
        for j, part in enumerate(parts): parent.insert(i + j, part)
        parent.remove(r)
    return True

def replace_in(doc, para_marker, old, new, used):
    p = find_para(doc, para_marker)
    if p is None:
        print(f"   !! paragraph not found for marker {para_marker!r}"); return False
    ok = replace_tracked(p, old, new, used)
    print(f"   {'ok ' if ok else 'FAIL'}: {old!r} -> {new!r}")
    return ok


# ---------- citation insertion at a comment's end anchor ----------
def _comment_end(body, cid):
    for el in body.iter(qn("w:commentRangeEnd")):
        if el.get(qn("w:id")) == str(cid):
            return el
    return None

def _nearest_prev_run_rpr(el):
    sib = el.getprevious()
    while sib is not None:
        if sib.tag == qn("w:r"):
            return sib.find(qn("w:rPr"))
        sib = sib.getprevious()
    return None

def insert_citation_at_comment(body, cid, text, used):
    end = _comment_end(body, cid)
    if end is None:
        print(f"   !! commentRangeEnd {cid} not found"); return False
    rpr = _nearest_prev_run_rpr(end)
    ins = _wrap_ins([_mk_run(text, rpr)], used)
    # insert after the commentReference run if it directly follows, else after end
    anchor = end
    nxt = end.getnext()
    if nxt is not None and nxt.tag == qn("w:r") and nxt.find(qn("w:commentReference")) is not None:
        anchor = nxt
    anchor.addnext(ins)
    print(f"   ok: comment {cid} -> inserted {text!r}")
    return True

def insert_at_offset(doc, phrase, ins_text, used):
    """Insert ins_text immediately AFTER `phrase` (untracked surrounding text)."""
    for p in doc.paragraphs:
        el = p._element
        runs = el.findall(qn("w:r"))
        full = ""; spans = []
        for r in runs:
            tt = r.find(qn("w:t")); txt = tt.text if (tt is not None and tt.text) else ""
            spans.append((r, len(full), len(full) + len(txt), txt)); full += txt
        i = full.find(phrase)
        if i == -1:
            continue
        pos = i + len(phrase)
        for r, rs, re_, txt in spans:
            if txt == "" or not (rs < pos <= re_):
                continue
            cut = pos - rs
            before, after = txt[:cut], txt[cut:]
            rpr = r.find(qn("w:rPr"))
            parent = r.getparent(); idx = list(parent).index(r)
            parts = []
            if before: parts.append(_mk_run(before, rpr))
            parts.append(_wrap_ins([_mk_run(ins_text, rpr)], used))
            if after: parts.append(_mk_run(after, rpr))
            for j, part in enumerate(parts): parent.insert(idx + j, part)
            parent.remove(r)
            return True
    return False


def insert_citation(doc, body, cid, text, fallback_phrase, used):
    """Insert citation after comment `cid`'s anchor; fall back to text anchor if the
    comment was resolved/deleted in Word."""
    if present(doc, text.strip()):
        print(f"   skip comment {cid} (citation already present): {text.strip()!r}"); return True
    if _comment_end(body, cid) is not None:
        return insert_citation_at_comment(body, cid, text, used)
    ok = insert_at_offset(doc, fallback_phrase, text, used)
    print(f"   {'ok ' if ok else 'FAIL'}: comment {cid} gone -> text anchor {fallback_phrase!r}")
    return ok


def insert_sentence_before(doc, anchor_marker, sentence, used):
    """Insert a tracked sentence run immediately before the run containing anchor_marker."""
    for p in doc.paragraphs:
        if anchor_marker in ftext(p._element):
            for r in p._element.findall(qn("w:r")):
                t = r.find(qn("w:t"))
                if t is not None and t.text and anchor_marker in t.text:
                    rpr = r.find(qn("w:rPr"))
                    ins = _wrap_ins([_mk_run(sentence, rpr)], used)
                    r.addprevious(ins)
                    print(f"   ok: RAG definition inserted before {anchor_marker!r}")
                    return True
    print(f"   !! anchor not found: {anchor_marker!r}"); return False


# ---------- reference entries ----------
REFS = {  # anchor_substring (insert BEFORE this existing entry) -> list of segment-lists
 "Introducing the Model Context Protocol": [[
    ("Ancillai, C., Sabatini, A., Gatti, M., & Perna, A. (2023). Digital technology and business model innovation: A systematic literature review and future research agenda. ", False),
    ("Technological Forecasting and Social Change, 188", True),
    (", 122307. https://doi.org/10.1016/j.techfore.2022.122307", False)]],
 "Blumer, H. (1969)": [[
    ("Bessen, J. E. (2018). ", False),
    ("AI and jobs: The role of demand", True),
    (" (NBER Working Paper No. 24235). National Bureau of Economic Research. https://doi.org/10.3386/w24235", False)]],
 "Burkhardt, S., & Rieder": [[
    ("Brynjolfsson, E., Li, D., & Raymond, L. (2025). Generative AI at work. ", False),
    ("The Quarterly Journal of Economics, 140", True),
    (f"(2), 889{EN}942. https://doi.org/10.1093/qje/qjae044", False)]],
 "Duarte, V., Zuniga-Jara": [[
    ("Doshi, A. R., & Hauser, O. P. (2024). Generative AI enhances individual creativity but reduces the collective diversity of novel content. ", False),
    ("Science Advances, 10", True),
    ("(28), eadn5290. https://doi.org/10.1126/sciadv.adn5290", False)]],
 "Gao, Y., Xiong": [
    [("Ellström, D., Holtström, J., Berg, E., & Josefsson, C. (2022). Dynamic capabilities for digital transformation. ", False),
     ("Journal of Strategy and Management, 15", True),
     (f"(2), 272{EN}286. https://doi.org/10.1108/JSMA-04-2021-0089", False)],
    [("Enholm, I. M., Papagiannidis, E., Mikalef, P., & Krogstie, J. (2022). Artificial intelligence and business value: A literature review. ", False),
     ("Information Systems Frontiers, 24", True),
     (f"(5), 1709{EN}1734. https://doi.org/10.1007/s10796-021-10186-w", False)]],
 "Holmström, J. (2022)": [
    [("Hanelt, A., Bohnsack, R., Marz, D., & Antunes Marante, C. (2021). A systematic review of the literature on digital transformation: Insights and implications for strategy and organizational change. ", False),
     ("Journal of Management Studies, 58", True),
     (f"(5), 1159{EN}1197. https://doi.org/10.1111/joms.12639", False)],
    [("Heimbach, I., Kostyra, D. S., & Hinz, O. (2015). Marketing automation. ", False),
     ("Business & Information Systems Engineering, 57", True),
     (f"(2), 129{EN}133. https://doi.org/10.1007/s12599-015-0370-8", False)]],
 "Kastner, J. K., & Hong": [[
    ("Kaartemo, V., & Helkkula, A. (2018). A systematic review of artificial intelligence and robots in value co-creation: Current status and future research avenues. ", False),
    ("Journal of Creating Value, 4", True),
    (f"(2), 211{EN}228. https://doi.org/10.1177/2394964318805625", False)]],
 "Krizhevsky, Alex": [[
    ("Kitsios, F., & Kamariotou, M. (2021). Artificial intelligence and business strategy towards digital transformation: A research agenda. ", False),
    ("Sustainability, 13", True),
    ("(4), 2025. https://doi.org/10.3390/su13042025", False)]],
 "LeCun, Y., Bengio": [[
    ("Kumar, V., Kotler, P., Gupta, S., & Rajan, B. (2025). Generative AI in marketing: Promises, perils, and public policy implications. ", False),
    ("Journal of Public Policy & Marketing, 44", True),
    ("(3). https://doi.org/10.1177/07439156241286499", False)]],
 "Lewis, P., Perez": [[
    ("Lepak, D. P., Smith, K. G., & Taylor, M. S. (2007). Value creation and value capture: A multilevel perspective. ", False),
    ("Academy of Management Review, 32", True),
    (f"(1), 180{EN}194. https://doi.org/10.5465/amr.2007.23464011", False)]],
 "Little, J. D. (1979)": [[
    ("Lincoln, Y. S., & Guba, E. G. (1985). ", False),
    ("Naturalistic inquiry", True),
    (". Sage.", False)]],
 "Ritala, P., Aaltonen": [
    [("Prasad Agrawal, K. (2023). Towards adoption of generative AI in organizational settings. ", False),
     ("Journal of Computer Information Systems, 64", True),
     (f"(5), 636{EN}651. https://doi.org/10.1080/08874417.2023.2240744", False)],
    [("Radford, A., Narasimhan, K., Salimans, T., & Sutskever, I. (2018). ", False),
     ("Improving language understanding by generative pre-training", True),
     (" [Technical report]. OpenAI. https://cdn.openai.com/research-covers/language-unsupervised/language_understanding_paper.pdf", False)]],
 "Vidal, J. F., Perotti": [[
    ("Verhoef, P. C., Broekhuizen, T., Bart, Y., Bhattacharya, A., Qi Dong, J., Fabian, N., & Haenlein, M. (2021). Digital transformation: A multidisciplinary reflection and research agenda. ", False),
    ("Journal of Business Research, 122", True),
    (f", 889{EN}901. https://doi.org/10.1016/j.jbusres.2019.09.022", False)]],
 "Weber, M., Engert": [[
    ("Warner, K. S. R., & Wäger, M. (2019). Building dynamic capabilities for digital transformation: An ongoing process of strategic renewal. ", False),
    ("Long Range Planning, 52", True),
    (f"(3), 326{EN}349. https://doi.org/10.1016/j.lrp.2018.12.001", False)]],
 "Woodside, A. G., Golfetto": [[
    ("Wessel, L., Baiyere, A., Ologeanu-Taddei, R., Cha, J., & Blegind-Jensen, T. (2021). Unpacking the difference between digital transformation and IT-enabled organizational transformation. ", False),
    ("Journal of the Association for Information Systems, 22", True),
    (f"(1), 102{EN}129. https://doi.org/10.17705/1jais.00655", False)]],
}

YAO = [("Yao, S., Zhao, J., Yu, D., Du, N., Shafran, I., Narasimhan, K., & Cao, Y. (2023). ReAct: Synergizing reasoning and acting in language models. ", False),
       ("International Conference on Learning Representations (ICLR)", True),
       (". https://doi.org/10.48550/arXiv.2210.03629", False)]


def main():
    shutil.copy(DOCX_PATH, BACKUP)
    doc = Document(str(DOCX_PATH))
    body = doc.element.body
    used = _used_ids(doc)

    # ---------- PART 1: reference entries ----------
    print("PART 1: reference-list entries")
    added = 0
    for anchor, entries in REFS.items():
        target = None
        for p in doc.paragraphs:
            if anchor in ftext(p._element):
                target = p._element; break
        if target is None:
            print(f"   !! anchor entry not found: {anchor!r}"); continue
        parent = target.getparent(); base = list(parent).index(target)
        offset = 0
        for segs in entries:
            sig = segs[0][0][:22]
            if present(doc, sig):
                print(f"   skip (already present): {sig!r}"); continue
            parent.insert(base + offset, ref_para(segs, used)); offset += 1; added += 1
        print(f"   ok: inserted {offset} before {anchor!r}")
    # Yao at end of list (after Xu)
    if present(doc, "Yao, S., Zhao"):
        print("   skip Yao (already present)")
    else:
        xu = None
        for p in doc.paragraphs:
            if "Xu, Y., Liu, X." in ftext(p._element):
                xu = p._element; break
        if xu is not None:
            xu.addnext(ref_para(YAO, used)); added += 1
            print("   ok: Yao inserted after Xu")
        else:
            print("   !! Xu anchor not found")
    print(f"   -> {added} reference entries added")

    # Charmaz 2006 -> 2014 (2nd ed.)
    print("PART 1b: Charmaz entry -> 2014 (2nd ed.)")
    charmaz = None
    for p in doc.paragraphs:
        tx = ftext(p._element)
        if "Charmaz, K. (2006). Constructing" in tx:
            charmaz = p._element; break
    if charmaz is not None:
        replace_tracked(charmaz, "(2006)", "(2014)", used)
        if "(2nd ed.)" not in ftext(charmaz):
            replace_tracked(charmaz, "qualitative analysis", "qualitative analysis (2nd ed.)", used)
        print("   ok: Charmaz converted to 2014 2nd ed.")
    else:
        print("   skip: Charmaz 2006 entry not found (already converted?)")

    # ---------- PART 2: in-text fixes ----------
    print("PART 2: in-text citation fixes")
    replace_in(doc, "Attention is All You Need", "Vaswani et al., 2016", "Vaswani et al., 2017", used)
    replace_in(doc, "made the public keenly aware", "Grewal et al., 2024", "Grewal et al., 2025", used)
    replace_in(doc, "made the public keenly aware", "Abou Elgeit", "Abou Elgheit", used)
    replace_in(doc, "identify sensing, seizing", "Ellström et al. (2021)", "Ellström et al. (2022)", used)
    replace_in(doc, "ongoing strategic renewal", "Warner and Wager (2019)", "Warner and Wäger (2019)", used)
    replace_in(doc, "competitive pressure and institutional expectations", "Agrawal, 2023", "Prasad Agrawal, 2023", used)
    replace_in(doc, "Leone et al. (2020) similarly show", "Leone et al. (2020)", "Leone et al. (2021)", used)
    # §5.1 discussion occurrences
    replace_in(doc, "corresponds closely to digital sensing", "Ellström et al., 2021", "Ellström et al., 2022", used)
    replace_in(doc, "passive instrument", "Leone et al., 2020", "Leone et al., 2021", used)
    # Charmaz normalize to 2014 (the §5.1 2006 in-text citation was already removed in Word)
    replace_in(doc, "surprises and sparks of ideas", "Charmaz, 2024", "Charmaz, 2014", used)

    # ---------- PART 3: insert new citations (comment anchor, text fallback) ----------
    print("PART 3: new citations at 'citation needed' anchors")
    insert_citation(doc, body, 28,  " (Chintalapati & Pandey, 2022)", "practitioner discourse", used)
    insert_citation(doc, body, 30,  " (Hanelt et al., 2021; Verhoef et al., 2021)", "GenAI and agentic systems", used)
    insert_citation(doc, body, 47,  " (Cottier et al., 2024)", "data and computation", used)
    insert_citation(doc, body, 99,  " (Huang & Rust, 2020)", "commercial performance", used)
    insert_citation(doc, body, 104, " (Brynjolfsson et al., 2025)", "improved employee productivity", used)
    insert_citation(doc, body, 107, " (Wessel et al., 2021)", "and sometimes organizational identity", used)
    insert_citation(doc, body, 112, " (Enholm et al., 2022)", "changes the work system", used)
    insert_citation(doc, body, 123, " (Lepak et al., 2007)", "a distinction between value creation and value capture", used)
    insert_citation(doc, body, 125, " (Kumar et al., 2025)", "increase content homogeneity", used)
    # RAG definition sentence (before the business-use sentence)
    RAG_DEF = ("In other words, RAG supplements a model’s input with relevant text retrieved from an "
               "external knowledge source at query time, grounding its output in that material rather than "
               "in its training data alone (Gao et al., 2023). ")
    if present(doc, "RAG supplements a model"):
        print("   skip RAG definition (already present)")
    else:
        insert_sentence_before(doc, "This development was immediately relevant", RAG_DEF, used)

    doc.save(str(DOCX_PATH))
    print(f"\nSaved: {DOCX_PATH}\nBackup: {BACKUP}")


if __name__ == "__main__":
    main()
