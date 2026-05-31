from docx import Document
from pathlib import Path

root = Path(__file__).resolve().parents[1]
out = []
for name in [
    "Thesis Draft - Daan Luttik - MBA.docx",
    "Thesis Draft - Daan Luttik - MBA - trustworthiness-table.docx",
]:
    path = root / name
    if not path.exists():
        out.append(f"MISSING {name}")
        continue
    d = Document(str(path))
    out.append(f"=== {name} tables={len(d.tables)} ===")
    for i, p in enumerate(d.paragraphs):
        if any(k in p.text for k in ("3.4", "Table 3", "Trustworthiness", "3.4.1")):
            out.append(f"{i}|{p.style.name}|{p.text[:80]}")
Path(__file__).with_name("check-docs.txt").write_text("\n".join(out), encoding="utf-8")
