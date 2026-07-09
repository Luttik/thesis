# -*- coding: utf-8 -*-
"""
Chapter 3 (Methodology) extensive proofread pass - tracked changes + comments,
author "Claude". Runs against the CURRENT state of the doc (after the user
accepted the previous round's tracked changes and resolved most comments).

Two phases on the same in-memory tree:
  1. ~22 tracked text fixes (grammar, terminology, tense, redundancy,
     mechanical punctuation), each a single-run substring replace.
  2. 6 review comments for judgment calls (vague referents, mixed tense,
     structural redundancy) - flagged, not force-edited.

Paragraphs are located via python-docx's `doc.paragraphs[idx]` (index is
stable and unambiguous, unlike walking raw body <w:p> which would also hit
paragraphs nested inside the chapter's 4 data tables). All edits within a
paragraph are applied before any comment in that paragraph is anchored, so
comment searches see the post-edit run structure.

Comments.xml / commentsExtended.xml / commentsIds.xml are NOT modeled by
python-docx, so: (a) peek at their current max id/paraId/durableId via raw
zipfile BEFORE opening with python-docx, (b) do the python-docx edit+save
pass, (c) re-open the saved file as a zip and graft in the new comment
entries using the ids reserved in (a).
"""
from __future__ import annotations

import copy
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT   = Path(__file__).resolve().parents[1]
DOCX   = ROOT / "Thesis Draft - Daan Luttik - MBA.docx"
BACKUP = ROOT / "Thesis Draft - Daan Luttik - MBA.ch3-proofread-backup.docx"
TEMP   = ROOT / "Thesis Draft - Daan Luttik - MBA.ch3-proofread-tmp.docx"

W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
W16 = "http://schemas.microsoft.com/office/word/2016/wordml/cid"
XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"

AUTHOR, INITIALS, DATE = "Claude", "C", "2026-07-09T00:00:00Z"


def w(t):   return f"{{{W}}}{t}"
def w14(t): return f"{{{W14}}}{t}"
def w15(t): return f"{{{W15}}}{t}"
def w16(t): return f"{{{W16}}}{t}"


# --------------------------------------------------------------------------- #
# tracked-change helpers (single-run substring replace, proven pattern)
# --------------------------------------------------------------------------- #
def _used_ids(doc) -> set[int]:
    return {int(el.get(qn("w:id"), 0))
            for el in doc.element.body.iter()
            if el.get(qn("w:id")) is not None}


def _nid(used: set[int]) -> int:
    n = max(used, default=0) + 1
    used.add(n)
    return n


def clone_rpr(r_elem):
    rpr = r_elem.find(qn("w:rPr"))
    return copy.deepcopy(rpr) if rpr is not None else None


def replace_text_tracked(p_elem, old: str, new: str, used: set[int]) -> bool:
    for r in p_elem.findall(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is None or not t.text or old not in t.text:
            continue
        rpr = clone_rpr(r)
        before, _, after = t.text.partition(old)
        parent = r.getparent()
        idx = list(parent).index(r)
        parts = []
        if before:
            rb = etree.Element(qn("w:r"))
            if rpr is not None:
                rb.append(copy.deepcopy(rpr))
            tb = etree.SubElement(rb, qn("w:t"))
            tb.set(XMLSPACE, "preserve")
            tb.text = before
            parts.append(rb)
        d = etree.Element(qn("w:del"))
        d.set(qn("w:id"), str(_nid(used)))
        d.set(qn("w:author"), AUTHOR)
        d.set(qn("w:date"), DATE)
        rd = etree.SubElement(d, qn("w:r"))
        if rpr is not None:
            rd.append(copy.deepcopy(rpr))
        td = etree.SubElement(rd, qn("w:delText"))
        td.set(XMLSPACE, "preserve")
        td.text = old
        parts.append(d)
        ins = etree.Element(qn("w:ins"))
        ins.set(qn("w:id"), str(_nid(used)))
        ins.set(qn("w:author"), AUTHOR)
        ins.set(qn("w:date"), DATE)
        ri = etree.SubElement(ins, qn("w:r"))
        if rpr is not None:
            ri.append(copy.deepcopy(rpr))
        ti = etree.SubElement(ri, qn("w:t"))
        ti.set(XMLSPACE, "preserve")
        ti.text = new
        parts.append(ins)
        if after:
            ra = etree.Element(qn("w:r"))
            if rpr is not None:
                ra.append(copy.deepcopy(rpr))
            ta = etree.SubElement(ra, qn("w:t"))
            ta.set(XMLSPACE, "preserve")
            ta.text = after
            parts.append(ra)
        for j, part in enumerate(parts):
            parent.insert(idx + j, part)
        parent.remove(r)
        return True
    return False


# --------------------------------------------------------------------------- #
# comment-range helpers (proven pattern from patch_ch3_comments.py)
# --------------------------------------------------------------------------- #
def crs(cid):
    e = etree.Element(qn("w:commentRangeStart")); e.set(qn("w:id"), str(cid)); return e

def cre(cid):
    e = etree.Element(qn("w:commentRangeEnd")); e.set(qn("w:id"), str(cid)); return e

def cref(cid):
    r = etree.Element(qn("w:r"))
    rpr = etree.SubElement(r, qn("w:rPr"))
    etree.SubElement(rpr, qn("w:rStyle")).set(qn("w:val"), "CommentReference")
    etree.SubElement(r, qn("w:commentReference")).set(qn("w:id"), str(cid))
    return r


def find_ins_by_text(p_elem, text):
    for ins in p_elem.iter(qn("w:ins")):
        ts = ins.findall(".//" + qn("w:t"))
        if len(ts) == 1 and (ts[0].text or "") == text:
            return ins
    return None


def comment_on_plain_substr(p_elem, substr, cid) -> bool:
    """Wrap `substr` (must live inside a single plain, non-ins/del run) with a comment range."""
    for r in p_elem.findall(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is None or not t.text or substr not in t.text:
            continue
        rpr = clone_rpr(r)
        before, _, after = t.text.partition(substr)
        parent = r.getparent()
        idx = list(parent).index(r)
        seq = []
        if before:
            rb = etree.Element(qn("w:r"))
            if rpr is not None:
                rb.append(copy.deepcopy(rpr))
            tb = etree.SubElement(rb, qn("w:t")); tb.set(XMLSPACE, "preserve"); tb.text = before
            seq.append(rb)
        seq.append(crs(cid))
        rm = etree.Element(qn("w:r"))
        if rpr is not None:
            rm.append(copy.deepcopy(rpr))
        tm = etree.SubElement(rm, qn("w:t")); tm.set(XMLSPACE, "preserve"); tm.text = substr
        seq.append(rm)
        seq.append(cre(cid))
        seq.append(cref(cid))
        if after:
            ra = etree.Element(qn("w:r"))
            if rpr is not None:
                ra.append(copy.deepcopy(rpr))
            ta = etree.SubElement(ra, qn("w:t")); ta.set(XMLSPACE, "preserve"); ta.text = after
            seq.append(ra)
        for j, el in enumerate(seq):
            parent.insert(idx + j, el)
        parent.remove(r)
        return True
    return False


def comment_on_ins(p_elem, ins_text, cid) -> bool:
    """Attach a comment range around an existing <w:ins> whose full text equals ins_text."""
    ins = find_ins_by_text(p_elem, ins_text)
    if ins is None:
        return False
    parent = ins.getparent()
    idx = list(parent).index(ins)
    parent.insert(idx, crs(cid))
    parent.insert(idx + 2, cre(cid))
    parent.insert(idx + 3, cref(cid))
    return True


def add_comment_entry(com_t, cex_t, cid_t, cid, paraid, durable, text):
    c = etree.SubElement(com_t, w("comment"))
    c.set(w("id"), str(cid)); c.set(w("author"), AUTHOR); c.set(w("date"), DATE); c.set(w("initials"), INITIALS)
    p = etree.SubElement(c, w("p")); p.set(w14("paraId"), paraid); p.set(w14("textId"), "77777777")
    etree.SubElement(etree.SubElement(p, w("pPr")), w("pStyle")).set(w("val"), "CommentText")
    r1 = etree.SubElement(p, w("r"))
    etree.SubElement(etree.SubElement(r1, w("rPr")), w("rStyle")).set(w("val"), "CommentReference")
    etree.SubElement(r1, w("annotationRef"))
    r2 = etree.SubElement(p, w("r"))
    t = etree.SubElement(r2, w("t")); t.set(XMLSPACE, "preserve"); t.text = text
    etree.SubElement(cex_t, w15("commentEx")).attrib.update({w15("paraId"): paraid, w15("done"): "0"})
    etree.SubElement(cid_t, w16("commentId")).attrib.update({w16("paraId"): paraid, w16("durableId"): durable})


# --------------------------------------------------------------------------- #
# EDITS  (para_index, old, new) - grouped by paragraph, left-to-right order
# --------------------------------------------------------------------------- #
EDITS = [
    # 3.2 Data collection
    (115, ". The process combines", ", combining"),
    (115, ")", ")."),
    (116, "employs", "employed"),                                    # re-fix: reverted after merge/accept
    (117, "constructive grounded theory", "constructivist grounded theory"),
    (117, "have been created", "were created"),
    (118, "constructive grounded theory", "constructivist grounded theory"),
    (118, "starts with open, intensive interviews", "started with open, intensive interviews"),
    (119, "their emerging categories", "the emerging categories"),
    # 3.3 Sample and sample size
    (121, "To find relevant candidates for the first group, the marketing managers",
          "For the first group, the marketing managers"),
    (121, "was chosen", "was used to identify relevant candidates"),
    (122, "is ", "was "),                                            # tense fix (see comment)
    (122, "stopped when", "continued until"),
    (122, "insights. ", "insights; throughout, "),
    (122, "During data gathering we ", "the researcher "),
    (122, "the following ", "subsequent "),
    # 3.4 Trustworthiness / table captions
    (128, "3.4", "3.4."),
    (129, "(see Table 2)", "(see Table 2)."),
    (130, "professional background in marketing and AI", "professional experience in marketing and AI"),
    (135, " 3", " 3."),
    (136, " development", " development."),
    (139, "Table 4", "Table 4."),
    (140, "Definitions of the theoretical categories", "Definitions of the theoretical categories."),
]

T_C1 = ("The referent of 'these insights' is unclear: the previous sentence describes "
        "participants' experiences, meanings, and actions, not insights, and strictly you collect "
        "data/experiences, not insights (insights are what analysis produces). Consider cutting this "
        "sentence (the previous one already establishes the theory-building purpose) or rewording to "
        "'These experiences, meanings, and actions were collected and analyzed to develop the "
        "resulting theory.'")
T_C2 = ("Shifts from 'the researcher' (singular, used throughout 3.1-3.4 for you) to 'the researchers' "
        "(plural) here. Fine if this is a general claim about the field; if it means this study, change "
        "to 'the researcher does not yet have.'")
T_C3 = ("'Forms the basis' / 'forms the link... (Charmaz, 2014)' are left present tense as general, "
        "citation-backed claims about how coding works in grounded theory (same convention as 'the "
        "point where new data no longer yields new insights' in 3.2). 'Was analyzed' / 'were "
        "categorized' / 'started with' describe this study specifically and are past tense. Flagging in "
        "case you'd rather the whole paragraph read in one tense.")
T_C4A = ("This was 'access might be limited' before the last edit; fixed the stray present tense to "
         "'was limited' for consistency with 'was already narrow' just before it, but that also drops "
         "the hedge ('might'). Confirm whether you want the firmer claim or 'access might have been "
         "limited.'")
T_C4B = ("Merged with the previous sentence and changed 'During data gathering we relied on...' to "
         "'...; throughout, the researcher relied on...' for two reasons: (1) 'we' is the only "
         "first-person use in this chapter (elsewhere it's 'the researcher' / passive; Chapter 5 does "
         "use 'we', but that's that section's own convention); (2) the two sentences both opened with "
         "'data gathering', which read as repetitive. Please check the merged sentence still says what "
         "you intended.")
T_C5 = ("The two stakeholder groups (marketing managers, AI experts) are introduced in 3.2 and then "
        "described again here in 3.3. Roughly the split is 3.2 = who was interviewed, 3.3 = how they "
        "were found, but worth checking it doesn't read as redundant.")

# (para_index, kind, anchor_arg, text) -- kind: 'substr' (plain run) | 'ins' (exact <w:ins> text)
COMMENTS = [
    (111, "substr", "In turn, these insights were collected and converted into theory", T_C1),
    (112, "substr", "the researchers do not yet have a robust understanding", T_C2),
    (118, "substr", "forms the link between the interviews and the resulting theory", T_C3),
    (122, "substr", "the sampling universe was already narrow and access", T_C4A),
    (122, "ins",    "the researcher ", T_C4B),
    (121, "substr", "consists of AI experts", T_C5),
]


def main():
    # ---- peek current max comment id / paraId / durableId BEFORE python-docx touches the file ----
    with zipfile.ZipFile(DOCX) as z:
        com0 = etree.fromstring(z.read("word/comments.xml"))
        cex0 = etree.fromstring(z.read("word/commentsExtended.xml"))
        cid0 = etree.fromstring(z.read("word/commentsIds.xml"))
    next_cid = max(int(c.get(w("id"))) for c in com0.findall(w("comment"))) + 1
    used_paraids = {p.get(w14("paraId"), "").upper()
                    for c in com0.findall(w("comment")) for p in c.iter(w("p"))}
    used_paraids |= {e.get(w15("paraId"), "").upper() for e in cex0}
    max_dur = max(int(x.get(w16("durableId")), 16) for x in cid0)

    def fresh_paraid(seed):
        v = seed
        while ("%08X" % v) in used_paraids:
            v += 1
        used_paraids.add("%08X" % v)
        return "%08X" % v

    reserved = []  # (cid, paraid, durable, text) in COMMENTS order
    for i, (_, _, _, text) in enumerate(COMMENTS):
        reserved.append((next_cid + i, fresh_paraid(0x0C200001 + i), "%08X" % (max_dur + 1 + i), text))

    # ---- phase 1: python-docx pass (text edits + comment ranges on the same tree) ----
    doc = Document(str(DOCX))
    paras = doc.paragraphs
    used_rev = _used_ids(doc)

    print("=== tracked edits ===")
    failures = []
    for idx, old, new in EDITS:
        ok = replace_text_tracked(paras[idx]._element, old, new, used_rev)
        print(f"[{'OK ' if ok else 'MISS'}] para {idx}: {old!r} -> {new!r}")
        if not ok:
            failures.append((idx, old))

    print("\n=== comments ===")
    for (idx, kind, arg, _), (cid, paraid, durable, text) in zip(COMMENTS, reserved):
        p_el = paras[idx]._element
        ok = comment_on_plain_substr(p_el, arg, cid) if kind == "substr" else comment_on_ins(p_el, arg, cid)
        print(f"[{'OK ' if ok else 'MISS'}] comment {cid} -> para {idx} ({kind}): {arg!r}")
        if not ok:
            failures.append((idx, arg))

    if failures:
        print(f"\nABORT: {len(failures)} anchor(s) not found; NOT saving.")
        for idx, what in failures:
            print(f"   miss para {idx}: {what!r}")
        return

    shutil.copy(DOCX, BACKUP)
    doc.save(str(TEMP))

    # ---- phase 2: graft new comment entries into the saved file ----
    with zipfile.ZipFile(TEMP) as z:
        names = z.namelist()
        blobs = {n: z.read(n) for n in names}

    com_t = etree.fromstring(blobs["word/comments.xml"])
    cex_t = etree.fromstring(blobs["word/commentsExtended.xml"])
    cid_t = etree.fromstring(blobs["word/commentsIds.xml"])

    for cid, paraid, durable, text in reserved:
        add_comment_entry(com_t, cex_t, cid_t, cid, paraid, durable, text)

    def ser(tree):
        return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    blobs["word/comments.xml"]         = ser(com_t)
    blobs["word/commentsExtended.xml"] = ser(cex_t)
    blobs["word/commentsIds.xml"]      = ser(cid_t)

    with zipfile.ZipFile(DOCX, "w", zipfile.ZIP_DEFLATED) as z:
        for n in names:
            z.writestr(n, blobs[n])
    TEMP.unlink(missing_ok=True)

    print(f"\nBacked up to: {BACKUP.name}")
    print(f"Applied {len(EDITS)} tracked edits + {len(COMMENTS)} comments, saved: {DOCX.name}")


if __name__ == "__main__":
    main()
